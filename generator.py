"""
Module de génération de documents avec Groq - Prompts stricts
"""

import os
from dotenv import load_dotenv
from groq import Groq
from config import (
    SYSTEM_PROMPT,
    TEMPLATES,
    GROQ_MODEL,
    GROQ_TEMPERATURE,
    GROQ_MAX_TOKENS,
    valider_donnees,
)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Charger la clé API depuis .env
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    raise ValueError("Veuillez configurer la variable d'environnement GROQ_API_KEY dans le fichier .env")

client = Groq(api_key=GROQ_API_KEY)


def generer_avec_groq(system_prompt, user_prompt):
    """Génère du contenu avec Groq en suivant strictement le prompt"""
    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=GROQ_TEMPERATURE,
            max_tokens=GROQ_MAX_TOKENS,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Erreur: {str(e)}"


def generer_document(template_name, donnees):
    """
    Génère un document Word basé sur le template et les données fournies.
    Valide d'abord les données puis utilise Groq pour générer le contenu.
    """
    # 1. Valider les données requises
    valide, message = valider_donnees(template_name, donnees)
    if not valide:
        return None, message

    template = TEMPLATES[template_name]

    # 2. Construire le prompt utilisateur avec les données
    user_prompt = f"""
Génère un document "{template["nom"]}" pour un établissement de santé au Sénégal.

DONNÉES FOURNIES (utilise EXACTEMENT ces données, n'invente rien):
{format_donnees_prompt(donnees)}

TEMPLATE DU DOCUMENT:
{template["description"]}

STRUCTURE À SUIVRE:
{" | ".join(template["sections"])}

INSTRUCTIONS:
1. Rédige en français administratif professionnel
2. Utilise EXACTEMENT les données fournies
3. Si une donnée n'est pas disponible, écrit "Non spécifié"
4. Ne fais jamais de suppositions
5. Respecte la structure définie
"""

    # 3. Générer le contenu avec Groq
    contenu = generer_avec_groq(SYSTEM_PROMPT, user_prompt)

    # 4. Créer le document Word
    doc = creer_document_word(template, donnees, contenu)

    return doc, "OK"


def format_donnees_prompt(donnees):
    """Formate les données pour le prompt"""
    lines = []
    for key, value in donnees.items():
        if value is not None and value != "":
            lines.append(f"- {key}: {value}")
    return "\n".join(lines)


def creer_document_word(template, donnees, contenu):
    """Crée un document Word formaté"""
    doc = Document()

    # Titre principal
    titre = doc.add_heading(template["nom"], level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter les métadonnées
    if "nom_etablissement" in donnees:
        p = doc.add_paragraph()
        p.add_run(f"Établissement: {donnees['nom_etablissement']}").bold = True

    if "periode" in donnees:
        doc.add_paragraph(f"Période: {donnees['periode']}")

    if "date" in donnees:
        doc.add_paragraph(f"Date: {donnees['date']}")

    doc.add_paragraph()  # Espace

    # Ajouter le contenu généré (par段)
    for paragraph in contenu.split("\n\n"):
        if paragraph.strip():
            # Essayer de détecter les titres
            if any(section in paragraph.upper() for section in template["sections"]):
                doc.add_heading(paragraph, level=2)
            else:
                doc.add_paragraph(paragraph)

    return doc


def generer_rapport_evaluation(etablissement, donnees_scores):
    """
    Génère le rapport d'auto-évaluation qualité basé sur les scores ESPC.
    """
    # Construire le prompt avec les scores
    scores_text = "\n".join(
        [f"- {cat}: {score}%" for cat, score in donnees_scores.items()]
    )

    user_prompt = f"""
Génère un rapport d'auto-évaluation qualité pour un Centre de Santé au Sénégal.

Établissement: {etablissement["nom"]}
Type: {etablissement.get("type", "CSR")}
Région: {etablissement.get("region", "Non spécifiée")}
District: {etablissement.get("district", "Non spécifié")}
Population: {etablissement.get("population", "Non spécifiée")}

SCORES OBTENUS (utilise ces scores exacts):
{scores_text}

STRUCTURE DU RAPPORT:
I. INTRODUCTION
- Contexte de l'établissement
- Objectif de l'évaluation
- Périmètre de l'évaluation

II. MÉTHODOLOGIE
- Outils utilisés (Grille ESPC vf STASS)
- Sources de données

III. ÉVALUATION DU MANAGEMENT (Catégorie A)
- Analyser le score fourni
- Identifier les points forts et axes d'amélioration

IV. ÉVALUATION DE LA QUALITÉ DES SOINS (Catégorie B)
- Analyser le score fourni
- Identifier les points forts et axes d'amélioration

V. ÉVALUATION DE LA SATISFACTION DES USAGERS (Catégorie C)
- Analyser le score fourni

VI. ÉVALUATION DES INTERVENTIONS COMMUNAUTAIRES (Catégorie D)
- Analyser le score fourni

VII. SYNTHÈSE GLOBALE
- Score global
- Points forts globaux
- Problèmes prioritaires identifiés

VIII. RECOMMANDATIONS ET PLAN D'AMÉLIORATION

INSTRUCTIONS CRITIQUES:
1. Utilise EXACTEMENT les scores fournis (n'invente pas de scores)
2. Si un score n'est pas fourni, indique "Évaluation non effectuée"
3. Fais des analyses factuelles basées sur les scores
4. Propose des recommandations réalistes
5. Utilise un langage administratif professionnel
"""

    contenu = generer_avec_groq(SYSTEM_PROMPT, user_prompt)

    # Créer le document
    doc = Document()

    # Titre
    titre = doc.add_heading("RAPPORT D'AUTO-ÉVALUATION DE LA QUALITÉ", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sous-titre
    doc.add_heading(etablissement["nom"], level=1)
    doc.add_paragraph()

    # Métadonnées
    doc.add_paragraph(
        f"Période d'évaluation: {donnees_scores.get('periode', 'Non spécifiée')}"
    )
    doc.add_paragraph(
        f"Date du rapport: {donnees_scores.get('date_rapport', 'Non spécifiée')}"
    )
    doc.add_paragraph()

    # Ajouter le contenu
    for paragraph in contenu.split("\n\n"):
        if paragraph.strip():
            if any(
                x in paragraph.upper()
                for x in ["I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII."]
            ):
                doc.add_heading(paragraph, level=1)
            elif any(
                x in paragraph.upper()
                for x in ["INTRODUCTION", "MÉTHODOLOGIE", "SYNTHÈSE", "RECOMMANDATIONS"]
            ):
                doc.add_heading(paragraph, level=2)
            else:
                doc.add_paragraph(paragraph)

    return doc
