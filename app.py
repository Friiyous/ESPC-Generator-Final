"""
Application Streamlit pour la génération de documents ESPC
Génère automatiquement des documents Word basés sur la grille d'évaluation qualité
"""

import streamlit as st
import os
import json
from datetime import datetime
from database import (
    init_database,
    ajouter_etablissement,
    get_etablissements,
    get_etablissement_by_id,
    ajouter_periode,
    get_periodes,
    ajouter_donnee,
    get_donnees_periode,
    sauvegarder_document,
    get_documents,
    ajouter_personnel,
    get_personnel,
    get_personnel_by_id,
    get_personnel_par_categorie,
    modifier_personnel,
    supprimer_personnel,
    verifier_doublon_personnel,
    ajouter_fiche_poste,
    get_fiches_poste,
    ajouter_nomination,
    get_nominations,
    ajouter_checklist,
    get_checklists,
    maj_statut_checklist,
)
from generator import generer_document, generer_rapport_evaluation
from config import TEMPLATES, ESPC_CATEGORIES

try:
    from whatsapp_integration import WhatsAppIntegration

    WHATSAPP_AVAILABLE = True
except ImportError:
    WHATSAPP_AVAILABLE = False
from chatbot import ChatbotAssistant

# Catégories de personnel pour CSR NAGNENEFOUN
NOUVELLES_CATEGORIES = [
    "INFIRMIERS DIPLOME D'ETAT",
    "SAGE FEMME DIPLOME D'ETAT",
    "AIDE SOIGNANTE",
    "JARDIN",
    "FILLES DE SALLE",
]

# Configuration de la page
st.set_page_config(
    page_title="Générateur Documents ESPC", page_icon="🏥", layout="wide"
)

# Initialiser la base de données
init_database()


def main():
    st.title("🏥 Générateur de Documents ESPC")
    st.markdown("---")

    # Menu latéral
    menu = st.sidebar.selectbox(
        "Menu",
        [
            "📱 Mobile View",
            "Accueil",
            "Établissements",
            "Gestion Personnel",
            "Missions Dynamiques",
            "Nouvelle Évaluation",
            "Générer Documents",
            "Performance",
            "💬 Chat Assistant",
            "Historique",
        ],
    )

    if menu == "📱 Mobile View":
        from mobile_interface import mobile_interface

        mobile_interface.show_home()

    elif menu == "Accueil":
        afficher_accueil()

    elif menu == "Établissements":
        gerer_etablissements()

    elif menu == "Gestion Personnel":
        gerer_personnel()

    elif menu == "Missions Dynamiques":
        missions_dynamiques()

    elif menu == "Nouvelle Évaluation":
        nouvelle_evaluation()

    elif menu == "Générer Documents":
        generer_documents()

    elif menu == "Performance":
        from performance_manager import performance_manager

        performance_manager.create_performance_dashboard()

    elif menu == "💬 Chat Assistant":
        from chatbot import chatbot

        show_chat_interface(chatbot)

    elif menu == "Historique":
        afficher_historique()


def afficher_accueil():
    """Page d'accueil améliorée avec nouvelles fonctionnalités"""
    st.header("🏥 ESPC Generator - Nouvelle Version")
    st.markdown("---")

    # Navigation rapide (via session state pour compatibilité Streamlit)
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("📱 Vue Mobile", use_container_width=True):
            st.session_state["menu"] = "📱 Mobile View"
            st.rerun()

    with col2:
        if st.button("📊 Performance", use_container_width=True):
            st.session_state["menu"] = "Performance"
            st.rerun()

    with col3:
        if st.button("💬 Chat IA", use_container_width=True):
            st.session_state["menu"] = "💬 Chat Assistant"
            st.rerun()

    with col4:
        if st.button("🎯 Missions", use_container_width=True):
            st.session_state["menu"] = "Missions Dynamiques"
            st.rerun()

    st.markdown("---")

    # Statistiques globales
    etablissements = get_etablissements()
    personnel = get_personnel()
    periodes = get_periodes(etablissements[0]["id"]) if etablissements else []
    documents = get_documents(periodes[0]["id"]) if periodes else []
    checklists = get_checklists() if periodes else []

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("🏢 Établissements", len(etablissements))
    with col2:
        st.metric("👥 Personnel", len(personnel))
    with col3:
        st.metric("📄 Documents", len(documents))
    with col4:
        st.metric("📋 Checklists", len(checklists))

    st.markdown("---")

    # Nouvelles fonctionnalités
    st.subheader("🚀 Nouvelles Fonctionnalités Disponibles")

    col1, col2 = st.columns(2)

    with col1:
        st.info("""
        ### 📱 Interface Mobile Optimisée
        - Design responsive pour mobile
        - Scanner de documents
        - Navigation intuitive
        - Actions rapides
        
        ### 💬 Assistant IA Intelligent
        - Chatbot contextuel
        - Réponses personnalisées
        - Suggestions adaptées
        - Aide en temps réel
        """)

    with col2:
        st.info("""
        ### 📊 Dashboard Performance
        - Suivi individuel
        - Comparaisons automatiques
        - Indicateurs clés
        - Recommandations
        
        ### 📱 Intégration WhatsApp
        - Envoi automatique
        - Notifications push
        - Groupes de coordination
        """)

    st.markdown("---")

    # Actions rapides
    st.subheader("⚡ Actions Rapides")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📄 Générer Rapport Rapide"):
            st.info("Formulaire simplifié pour génération rapide")

    with col2:
        if st.button("📅 Planifier Réunion"):
            st.info("Planification rapide de réunions COGES")

    with col3:
        if st.button("📊 Voir Performance"):
            performance_manager.create_performance_dashboard()

    # Dernières activités
    st.subheader("📈 Dernières Activités")

    activities = [
        "✅ Checklist T1 générée",
        "📊 Rapport mensuel créé",
        "👥 Nouvelle nomination: Koffi Kouassi",
        "💬 Chat IA: Question sur la gestion des stocks",
        "📱 Document envoyé par WhatsApp",
    ]

    for activity in activities:
        st.write(f"• {activity}")

    # Tips
    st.markdown("---")
    st.subheader("💡 Conseils d'utilisation")

    tips = [
        "Utilisez la vue mobile pour une accès rapide",
        "Le chat IA vous aide avec des questions spécifiques",
        "Le dashboard montre les performances en temps réel",
        "WhatsApp garde tout le monde informé",
        "Les missions dynamiques s'adaptent au contexte",
    ]

    for tip in tips:
        st.write(f"• {tip}")


def gerer_etablissements():
    """Gestion des établissements"""
    st.header("🏢 Établissements")

    # Formulaire d'ajout
    with st.expander("Ajouter un nouvel établissement", expanded=False):
        with st.form("form_etablissement"):
            col1, col2 = st.columns(2)

            with col1:
                nom = st.text_input("Nom de l'établissement")
                type_ = st.selectbox("Type", ["CSR", "CSU", "CSPR", "Hôpital"])
                region = st.text_input("Région")

            with col2:
                district = st.text_input("District")
                population = st.number_input("Population", min_value=0, step=100)
                telephone = st.text_input("Téléphone")

            responsable = st.text_input("Responsable")

            submit = st.form_submit_button("Ajouter")

            if submit and nom:
                ajouter_etablissement(
                    nom, type_, region, district, population, telephone, responsable
                )
                st.success(f"Établissement '{nom}' ajouté!")
                st.rerun()

    # Liste des établissements
    st.subheader("Liste des établissements")
    etablissements = get_etablissements()

    if etablissements:
        for etab in etablissements:
            with st.container():
                st.markdown(
                    f"**{etab['nom']}** - {etab['type']} | {etab['region']} | Pop: {etab['population']}"
                )
                st.markdown("---")
    else:
        st.info("Aucun établissement enregistré")


def nouvelle_evaluation():
    """Créer une nouvelle évaluation"""
    st.header("📋 Nouvelle Évaluation")

    etablissements = get_etablissements()

    if not etablissements:
        st.warning("Veuillez d'abord ajouter un établissement")
        return

    # Sélection de l'établissement
    etab_dict = {e["nom"]: e for e in etablissements}
    nom_etab = st.selectbox("Établissement", list(etab_dict.keys()))

    if nom_etab:
        etab = etab_dict[nom_etab]

        st.subheader(f"Évaluation pour: {etab['nom']}")

        # Période
        col1, col2 = st.columns(2)

        with col1:
            trimestre = st.selectbox("Trimestre", ["T1", "T2", "T3", "T4"])
            annee = st.number_input(
                "Année", min_value=2020, max_value=2030, value=datetime.now().year
            )

        with col2:
            date_debut = st.date_input("Date début")
            date_fin = st.date_input("Date fin")

        periode = f"{trimestre} {annee}"

        # Saisie des données d'évaluation ESPC
        st.markdown("### 📊 Données d'évaluation ESPC")

        donnees = {}

        # Catégorie A: Management
        with st.expander("A. MANAGEMENT (600 points)", expanded=True):
            st.markdown("#### A1. Gouvernance (200 pts)")
            donnees["governance_paa"] = st.selectbox(
                "PAA élaboré (Oui/Non)", ["Oui", "Non"], key="a1"
            )
            donnees["governance_coges"] = st.selectbox(
                "COGES fonctionnel (Oui/Non)", ["Oui", "Non"], key="a2"
            )
            donnees["governance_ag"] = st.selectbox(
                "AG tenues (Oui/Non)", ["Oui", "Non"], key="a3"
            )
            donnees["governance_donnees"] = st.selectbox(
                "Gestion des données (Oui/Non)", ["Oui", "Non"], key="a4"
            )

            st.markdown("#### A2. Gestion des RH (50 pts)")
            donnees["rh_effectif"] = st.selectbox(
                "Effectif suffisant (Oui/Non)", ["Oui", "Non"], key="a5"
            )

            st.markdown("#### A3. Gouvernance financière (350 pts)")
            donnees["finance_documents"] = st.selectbox(
                "Documents financiers (Oui/Non)", ["Oui", "Non"], key="a6"
            )
            donnees["finance_primes"] = st.selectbox(
                "Répartition primes conforme (Oui/Non)", ["Oui", "Non"], key="a7"
            )

        # Catégorie B: Qualité des soins
        with st.expander("B. QUALITÉ DES SOINS (750 points)", expanded=True):
            st.markdown("#### B1. Accueil (175 pts)")
            donnees["accueil_signaletique"] = st.selectbox(
                "Signalétique (Oui/Non)", ["Oui", "Non"], key="b1"
            )
            donnees["accueil_cmu"] = st.selectbox(
                "Agent CMU (Oui/Non)", ["Oui", "Non"], key="b2"
            )

            st.markdown("#### B2. Sécurité (25 pts)")
            donnees["securite_cloture"] = st.selectbox(
                "Clôture (Oui/Non)", ["Oui", "Non"], key="b3"
            )
            donnees["securite_gardien"] = st.selectbox(
                "Gardien de nuit (Oui/Non)", ["Oui", "Non"], key="b4"
            )

            st.markdown("#### B3. Hygiène hospitalière (150 pts)")
            donnees["hygiene_lavage"] = st.selectbox(
                "Dispositif lavage mains (Oui/Non)", ["Oui", "Non"], key="b5"
            )
            donnees["hygiene_dechets"] = st.selectbox(
                "Gestion des déchets (Oui/Non)", ["Oui", "Non"], key="b6"
            )

            st.markdown("#### B4. SONU (150 pts)")
            donnees["sonu_infrastructure"] = st.selectbox(
                "Infrastructure SONU (Oui/Non)", ["Oui", "Non"], key="b7"
            )
            donnees["sonu_equipement"] = st.selectbox(
                "Équipement SONU (Oui/Non)", ["Oui", "Non"], key="b8"
            )
            donnees["sonu_personnel"] = st.selectbox(
                "Personnel qualifié (Oui/Non)", ["Oui", "Non"], key="b9"
            )

            st.markdown("#### B5. Pharmacie (90 pts)")
            donnees["pharmacie_local"] = st.selectbox(
                "Local pharmacie conforme (Oui/Non)", ["Oui", "Non"], key="b10"
            )
            donnees["pharmacie_stock"] = st.selectbox(
                "Gestion des stocks (Oui/Non)", ["Oui", "Non"], key="b11"
            )

            st.markdown("#### B6. Médicaments traceurs (50 pts)")
            donnees["medicaments_disponibilite"] = st.selectbox(
                "Disponibilité médicaments (Oui/Non)", ["Oui", "Non"], key="b12"
            )

        # Catégorie C: Satisfaction
        with st.expander("C. SATISFACTION DES USAGERS (150 points)", expanded=True):
            donnees["satisfaction_enquete"] = st.selectbox(
                "Enquête satisfaction réalisée (Oui/Non)", ["Oui", "Non"], key="c1"
            )
            donnees["satisfaction_taux"] = st.number_input(
                "Taux de satisfaction (%)", min_value=0, max_value=100, value=50
            )

        # Catégorie D: Interventions communautaires
        with st.expander("D. INTERVENTIONS COMMUNAUTAIRES (100 points)", expanded=True):
            donnees["asc_supervision"] = st.selectbox(
                "Supervision ASC (Oui/Non)", ["Oui", "Non"], key="d1"
            )
            donnees["asc_medicaments"] = st.selectbox(
                "Médicaments ASC disponibles (Oui/Non)", ["Oui", "Non"], key="d2"
            )

        # Observations
        st.markdown("### 📝 Observations")
        observations = st.text_area("Observations générales", height=100)

        if st.button("💾 Sauvegarder l'évaluation"):
            # Sauvegarder la période
            id_periode = ajouter_periode(
                etab["id"], periode, trimestre, annee, str(date_debut), str(date_fin)
            )

            # Sauvegarder les données
            for key, value in donnees.items():
                score = 1 if value == "Oui" else 0
                ajouter_donnee(id_periode, "ESPC", "Divers", key, value, score, "")

            st.success(f"Évaluation sauvegardée pour {periode}!")

            # Auto-générer checklist de préparation
            try:
                with open("templates_checklist.json", "r") as f:
                    checklists = json.load(f)

                ajouter_checklist(
                    id_periode,
                    "checklist_preparation_evaluation",
                    json.dumps(checklists["checklist_preparation_evaluation"]["items"]),
                    "en_cours",
                )
                st.info("✅ Checklist de préparation automatiquement créée!")
            except FileNotFoundError:
                pass  # Ignorer si le fichier n'existe pas

            st.rerun()


def generer_documents():
    """Générer des documents"""
    st.header("📄 Générer des Documents")

    etablissements = get_etablissements()

    if not etablissements:
        st.warning("Aucun établissement enregistré")
        return

    # Sélection
    etab_dict = {e["nom"]: e for e in etablissements}
    nom_etab = st.selectbox("Établissement", list(etab_dict.keys()))
    etab = etab_dict[nom_etab]

    periodes = get_periodes(etab["id"])

    if not periodes:
        st.warning("Aucune évaluation enregistrée pour cet établissement")
        return

    periode_dict = {
        f"{p['periode']} ({p['date_debut']} - {p['date_fin']})": p for p in periodes
    }
    periode_nom = st.selectbox("Période d'évaluation", list(periode_dict.keys()))
    periode = periode_dict[periode_nom]

    # Type de document à générer
    st.subheader("Type de document")

    type_doc = st.selectbox(
        "Sélectionner le document",
        [
            "Rapport d'évaluation qualité",
            "Rapport mensuel",
            "PV Réunion COGES",
            "PV Réunion Mensuelle",
            "Rapport supervision ASC",
            "Checklist Préparation",
        ],
    )

    # Formulaire dynamique selon le type
    donnees_generation = {
        "nom_etablissement": etab["nom"],
        "periode": periode["periode"],
        "trimestre": periode["trimestre"],
        "annee": periode["annee"],
        "date": str(datetime.now().strftime("%d/%m/%Y")),
        "region": etab.get("region", ""),
        "district": etab.get("district", ""),
        "population": etab.get("population", 0),
    }

    # Ajouter les champs spécifiques
    if type_doc == "Rapport d'évaluation qualité":
        donnees_generation["score_management"] = st.slider(
            "Score Management (%)", 0, 100, 70
        )
        donnees_generation["score_qualite_soins"] = st.slider(
            "Score Qualité des soins (%)", 0, 100, 60
        )
        donnees_generation["score_satisfaction"] = st.slider(
            "Score Satisfaction (%)", 0, 100, 65
        )
        donnees_generation["score_interventions"] = st.slider(
            "Score Interventions (%)", 0, 100, 70
        )

        st.info(
            "💡 Les scores sont utilisés pour générer une analyse factuelle. L'IA ne peut pas inventer d'autres scores."
        )

    elif type_doc == "Rapport mensuel":
        donnees_generation["consultations"] = st.number_input(
            "Nombre de consultations", min_value=0, value=500
        )
        donnees_generation["accouchements"] = st.number_input(
            "Nombre d'accouchements", min_value=0, value=25
        )
        donnees_generation["hospitalisations"] = st.number_input(
            "Nombre d'hospitalisations", min_value=0, value=50
        )
        donnees_generation["activites_realisees"] = st.text_area(
            "Activités réalisées (description)", height=100
        )
        donnees_generation["difficultes"] = st.text_area(
            "Difficultés rencontrées", height=80
        )
        donnees_generation["recommandations"] = st.text_area(
            "Recommandations", height=80
        )

    elif type_doc == "PV Réunion COGES":
        donnees_generation["type_reunion"] = "Réunion COGES"

        # Thèmes suggérés pour la réunion COGES
        st.info("🎯 Thèmes suggérés pour la réunion COGES:")
        themes_suggeres = [
            "Analyse des indicateurs mensuels",
            "Planification des activités du trimestre",
            "Gestion des stocks et des finances",
            "Activités communautaires",
            "Formation du personnel",
            "Évaluation de la qualité",
        ]

        themes_selection = st.multiselect(
            "Sélectionner les thèmes à traiter",
            themes_suggeres,
            default=themes_suggeres[:3],
        )

        donnees_generation["themes"] = "; ".join(themes_selection)

        donnees_generation["lieu"] = st.text_input("Lieu de la réunion")
        donnees_generation["participants"] = st.text_area(
            "Participants (noms)", height=80
        )
        donnees_generation["ordre_du_jour"] = st.text_area("Ordre du jour", height=80)
        donnees_generation["deliberations"] = st.text_area("Délibérations", height=100)
        donnees_generation["decisions"] = st.text_area("Décisions prises", height=80)

    elif type_doc == "Rapport supervision ASC":
        donnees_generation["nb_asc"] = st.number_input(
            "Nombre d'ASC", min_value=0, value=3
        )
        donnees_generation["asc_supervises"] = st.number_input(
            "ASC supervisés", min_value=0, value=3
        )
        donnees_generation["activites"] = st.text_area(
            "Activités supervisées", height=80
        )
        donnees_generation["resultats"] = st.text_area("Résultats obtenus", height=80)
        donnees_generation["difficultes"] = st.text_area("Difficultés", height=80)
        donnees_generation["recommandations"] = st.text_area(
            "Recommandations", height=80
        )

    elif type_doc == "Plan d'action":
        donnees_generation["contexte"] = st.text_area("Contexte", height=80)
        donnees_generation["objectifs"] = st.text_area("Objectifs", height=80)
        donnees_generation["activites"] = st.text_area("Activités prévues", height=100)
        donnees_generation["calendrier"] = st.text_area("Calendrier", height=80)
        donnees_generation["budget"] = st.number_input(
            "Budget estimé (FCFA)", min_value=0, value=5000000
        )
        donnees_generation["responsables"] = st.text_area("Responsables", height=80)

    elif type_doc == "Checklist Préparation":
        # Afficher la checklist
        try:
            with open("templates_checklist.json", "r") as f:
                checklists = json.load(f)

            template = checklists["checklist_preparation_evaluation"]

            st.info("📋 Checklist de préparation pour l'évaluation ESPC")

            # Générer le contenu
            contenu_checklist = f"""
            CHECKLIST DE PRÉPARATION À L'ÉVALUATION ESPC
            Établissement: {etab["nom"]}
            Période: {periode["periode"]}
            Date: {datetime.now().strftime("%d/%m/%Y")}
            
            """

            for section, items in template["items"].items():
                contenu_checklist += f"\n{'=' * 50}\n{section}\n{'=' * 50}\n"
                for i, item in enumerate(items, 1):
                    contenu_checklist += f"{i}. [ ] {item}\n"

            # Prévisualisation
            with st.expander("👁️ Prévisualiser la checklist"):
                st.text_area("Contenu", value=contenu_checklist, height=400)

            # Boutons d'action
            col1, col2 = st.columns(2)

            with col1:
                if st.button("📄 Générer checklist Word"):
                    # Sauvegarder en base
                    ajouter_checklist(
                        periode["id"],
                        "checklist_preparation_evaluation",
                        json.dumps(template["items"]),
                        "en_cours",
                    )

                    # Créer document Word
                    from docx import Document

                    doc = Document()
                    doc.add_heading(
                        "CHECKLIST DE PRÉPARATION À L'ÉVALUATION ESPC", level=1
                    )
                    doc.add_paragraph(f"Établissement: {etab['nom']}")
                    doc.add_paragraph(f"Période: {periode['periode']}")
                    doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
                    doc.add_paragraph()

                    for section, items in template["items"].items():
                        doc.add_heading(section, level=2)
                        for i, item in enumerate(items, 1):
                            doc.add_paragraph(f"{i}. [ ] {item}")

                    # Télécharger
                    from io import BytesIO

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.success("Checklist générée!")

                    st.download_button(
                        label="📥 Télécharger checklist Word",
                        data=buffer,
                        file_name=f"Checklist_{etab['nom']}_{periode['periode']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            with col2:
                if st.button("✅ Marquer comme terminée"):
                    ajouter_checklist(
                        periode["id"],
                        "checklist_preparation_evaluation",
                        json.dumps(template["items"]),
                        "terminee",
                    )
                    st.success("Checklist marquée comme terminée!")
                    st.rerun()

        except FileNotFoundError:
            st.error("Fichier de checklist non trouvé")

    # Bouton de génération pour les autres documents avec option WhatsApp
    if type_doc != "Checklist Préparation":
        col1, col2 = st.columns([1, 1])

        with col1:
            generate_button = st.button("🚀 Générer le document")

        with col2:
            whatsapp_option = st.checkbox("📱 Envoyer par WhatsApp")

        if generate_button:
            with st.spinner("Génération en cours..."):
                try:
                    # Mapper le type de document au template
                    template_map = {
                        "Rapport d'évaluation qualité": "rapport_qualite",
                        "Rapport mensuel": "rapport_mensuel",
                        "PV Réunion COGES": "pv_cooges",
                        "Rapport supervision ASC": "rapport_supervision_asc",
                        "Plan d'action": "plan_action",
                    }

                    template_name = template_map[type_doc]
                    doc, status = generer_document(template_name, donnees_generation)

                    if status == "OK":
                        # Sauvegarder en base
                        sauvegarder_document(
                            periode["id"], type_doc, type_doc, "Contenu généré"
                        )

                    # Télécharger
                    from io import BytesIO

                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)

                    st.success("✅ Document généré avec succès!")

                    st.download_button(
                        label="📥 Télécharger le document Word",
                        data=doc_buffer,
                        file_name=f"{type_doc}_{etab['nom']}_{periode['periode']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                    # Option WhatsApp
                    if whatsapp_option:
                        if WHATSAPP_AVAILABLE:
                            try:
                                whatsapp = WhatsAppIntegration()
                                success, message = whatsapp.send_document_notification(
                                    type_doc, etab["nom"], periode["periode"]
                                )
                                if success:
                                    st.success("📱 Document envoyé par WhatsApp!")
                                else:
                                    st.error(f"Erreur WhatsApp: {message}")
                            except Exception as e:
                                st.error(f"WhatsApp non configuré: {str(e)}")
                        else:
                            st.error(
                                "❌ WhatsApp non disponible. Installez Twilio: pip install twilio"
                            )
                    else:
                        st.error(f"Erreur: {status}")

                except Exception as e:
                    st.error(f"Erreur lors de la génération: {str(e)}")


def gerer_personnel():
    """Gestion du personnel avec logique métier CSR"""
    st.header("👥 Gestion du Personnel - CSR NAGNENEFOUN")
    st.markdown("---")

    # Mapping catégorie → poste automatique
    POSTES_PAR_CATEGORIE = {
        "INFIRMIERS DIPLOME D'ETAT": "Infirmier Diplômé d'État",
        "SAGE FEMME DIPLOME D'ETAT": "Sage-Femme Diplômée d'État",
        "AIDE SOIGNANTE": "Aide Soignante",
        "JARDIN": "Jardinier",
        "FILLES DE SALLE": "Fille de Salle",
    }

    # Statistiques globales
    personnel = get_personnel()
    total = len(personnel)
    actifs = sum(1 for p in personnel if p["statut"] == "actif")

    col_stats1, col_stats2, col_stats3, col_stats4, col_stats5 = st.columns(5)
    with col_stats1:
        st.metric("👥 Total", total)
    with col_stats2:
        st.metric("✅ Actifs", actifs)
    with col_stats3:
        st.metric("⏸️ Inactifs", total - actifs)
    with col_stats4:
        cats = [p["categorie"] for p in personnel if p["statut"] == "actif"]
        top_cat = max(set(cats), key=cats.count) if cats else "—"
        st.metric("🏆 Catégorie principale", top_cat)
    with col_stats5:
        anciens = [p for p in personnel if p["date_embauche"]]
        if anciens:
            plus_ancien = min(anciens, key=lambda p: p["date_embauche"])
            annee = plus_ancien["date_embauche"][:4] if len(plus_ancien["date_embauche"]) >= 4 else "?"
            st.metric("📅 Ancienneté max", annee)
        else:
            st.metric("📅 Ancienneté max", "—")

    # Répartition par catégorie
    st.subheader("📊 Répartition par catégorie")
    cols = st.columns(len(NOUVELLES_CATEGORIES))
    for i, cat in enumerate(NOUVELLES_CATEGORIES):
        with cols[i]:
            count = sum(1 for p in personnel if p["categorie"] == cat and p["statut"] == "actif")
            st.metric(cat, count)

    st.markdown("---")

    # Navigation
    tab_personnel = st.tabs(
        ["📋 Liste", "➕ Ajouter/Modifier", "📄 Fiches de Poste", "📜 Nominations"]
    )

    # =====================================================================
    # TAB 1 : LISTE DU PERSONNEL
    # =====================================================================
    with tab_personnel[0]:
        st.subheader("📋 Liste du Personnel")

        col_filtre, col_rech = st.columns([2, 3])
        with col_filtre:
            categorie_filtre = st.selectbox(
                "Filtrer par catégorie",
                ["Toutes"] + NOUVELLES_CATEGORIES,
                key="filtre_cat",
            )
        with col_rech:
            recherche = st.text_input("🔍 Rechercher par nom ou prénom", placeholder="Nom ou prénom...")

        personnel_affiche = get_personnel()
        if categorie_filtre != "Toutes":
            personnel_affiche = [p for p in personnel_affiche if p["categorie"] == categorie_filtre]
        if recherche:
            r = recherche.lower()
            personnel_affiche = [p for p in personnel_affiche if r in p["nom"].lower() or r in p["prenom"].lower()]

        if personnel_affiche:
            st.caption(f"📌 {len(personnel_affiche)} agent(s) trouvé(s)")
            for pers in personnel_affiche:
                with st.container(border=True):
                    col1, col2, col3, col4 = st.columns([1, 3, 2, 1])

                    # Badge catégorie
                    with col1:
                        cat_short = pers["categorie"]
                        st.markdown(f"**`{cat_short}`**")

                    # Infos
                    with col2:
                        # Calcul ancienneté
                        if pers["date_embauche"]:
                            debut = datetime.strptime(pers["date_embauche"], "%Y-%m-%d")
                            duree = datetime.now() - debut
                            anciennete = f"{duree.days // 365} ans" if duree.days >= 365 else f"{duree.days // 30} mois"
                        else:
                            anciennete = "N/A"

                        st.write(f"**{pers['prenom']} {pers['nom']}**")
                        st.caption(f"{pers['poste']} · 📅 {anciennete}")

                    with col3:
                        if pers.get("telephone"):
                            st.write(f"📞 {pers['telephone']}")

                    # Statut + actions
                    with col4:
                        statut_color = "🟢" if pers["statut"] == "actif" else "🔴"
                        st.write(f"{statut_color} {pers['statut']}")

                        # Suppression avec confirmation
                        if st.button("🗑️", key=f"del_{pers['id']}", help="Supprimer"):
                            st.session_state[f"confirm_del_{pers['id']}"] = True

                        if st.session_state.get(f"confirm_del_{pers['id']}", False):
                            st.warning(f"Supprimer {pers['prenom']} {pers['nom']} ?")
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.button("✅ Oui", key=f"yes_del_{pers['id']}"):
                                    supprimer_personnel(pers['id'])
                                    st.success(f"{pers['prenom']} {pers['nom']} supprimé!")
                                    del st.session_state[f"confirm_del_{pers['id']}"]
                                    st.rerun()
                            with c2:
                                if st.button("❌ Non", key=f"no_del_{pers['id']}"):
                                    del st.session_state[f"confirm_del_{pers['id']}"]
                                    st.rerun()

                        if st.button("✏️", key=f"edit_{pers['id']}", help="Modifier"):
                            st.session_state.edit_id = pers["id"]
                            st.rerun()
        else:
            st.info("Aucun personnel trouvé")

    # =====================================================================
    # TAB 2 : AJOUTER / MODIFIER
    # =====================================================================
    with tab_personnel[1]:
        st.subheader("➕ Ajouter / Modifier un Agent")

        mode_edition = "edit_id" in st.session_state

        if mode_edition:
            pers = get_personnel_by_id(st.session_state.edit_id)
            if not pers:
                st.error("Agent introuvable")
                del st.session_state.edit_id
                st.rerun()
                return
            st.info(f"✏️ Modification de **{pers['prenom']} {pers['nom']}**")

        with st.form("form_personnel"):
            if mode_edition:
                nom = st.text_input("Nom *", value=pers["nom"])
                prenom = st.text_input("Prénom *", value=pers["prenom"])
                categorie = st.selectbox(
                    "Catégorie *",
                    NOUVELLES_CATEGORIES,
                    index=NOUVELLES_CATEGORIES.index(pers["categorie"])
                    if pers["categorie"] in NOUVELLES_CATEGORIES
                    else 0,
                )
                # Poste automatique
                poste_default = POSTES_PAR_CATEGORIE.get(categorie, pers["poste"])
                poste = st.text_input("Poste", value=poste_default, disabled=True,
                    help="Le poste est défini automatiquement selon la catégorie")
                telephone = st.text_input("Téléphone", value=pers.get("telephone", ""),
                    placeholder="Ex: 07 00 00 00 00")
                date_embauche = st.date_input(
                    "Date d'embauche *",
                    value=datetime.strptime(pers["date_embauche"], "%Y-%m-%d")
                    if pers["date_embauche"]
                    else datetime.now(),
                    max_value=datetime.now(),
                )
                statut = st.selectbox(
                    "Statut",
                    ["actif", "inactif"],
                    index=0 if pers["statut"] == "actif" else 1,
                )
                submit = st.form_submit_button("💾 Modifier", use_container_width=True)
                if submit:
                    if not nom or not prenom:
                        st.error("Le nom et le prénom sont obligatoires")
                    elif date_embauche > datetime.now().date():
                        st.error("La date d'embauche ne peut pas être dans le futur")
                    else:
                        # Vérifier doublon (sauf soi-même)
                        if verifier_doublon_personnel(nom, prenom, categorie, exclude_id=pers["id"]):
                            st.warning(f"⚠️ {nom} {prenom} ({categorie}) existe déjà !")
                        else:
                            modifier_personnel(
                                st.session_state.edit_id,
                                nom, prenom, categorie,
                                POSTES_PAR_CATEGORIE[categorie],
                                telephone,
                                str(date_embauche), statut,
                            )
                            st.success(f"✅ {prenom} {nom} modifié avec succès!")
                            del st.session_state.edit_id
                            st.rerun()
            else:
                nom = st.text_input("Nom *", placeholder="Ex: KOUAME")
                prenom = st.text_input("Prénom *", placeholder="Ex: Jean")
                categorie = st.selectbox("Catégorie *", NOUVELLES_CATEGORIES)
                # Poste automatique selon catégorie
                poste = st.text_input("Poste", value=POSTES_PAR_CATEGORIE[categorie], disabled=True,
                    help="Le poste est défini automatiquement selon la catégorie")
                telephone = st.text_input("Téléphone", placeholder="Ex: 07 00 00 00 00")
                date_embauche = st.date_input("Date d'embauche *", value=datetime.now(), max_value=datetime.now())

                st.caption(f"📌 Effectif actuel **{categorie}** : {sum(1 for p in get_personnel() if p['categorie'] == categorie and p['statut'] == 'actif')} agent(s)")

                submit = st.form_submit_button("✅ Ajouter", use_container_width=True)
                if submit:
                    if not nom or not prenom:
                        st.error("Le nom et le prénom sont obligatoires")
                    elif date_embauche > datetime.now().date():
                        st.error("La date d'embauche ne peut pas être dans le futur")
                    elif verifier_doublon_personnel(nom, prenom, categorie):
                        st.warning(f"⚠️ **{nom} {prenom}** ({categorie}) existe déjà dans la base !")
                    else:
                        ajouter_personnel(
                            nom, prenom, categorie,
                            POSTES_PAR_CATEGORIE[categorie],
                            telephone, str(date_embauche),
                        )
                        st.success(f"✅ {prenom} {nom} ({categorie}) ajouté avec succès!")
                        st.rerun()

        if mode_edition:
            if st.button("↩️ Annuler la modification"):
                del st.session_state.edit_id
                st.rerun()

    # =====================================================================
    # TAB 3 : FICHES DE POSTE
    # =====================================================================
    with tab_personnel[2]:
        st.subheader("📄 Fiches de Poste")

        personnel = get_personnel()
        if personnel:
            options = [f"{p['prenom']} {p['nom']} — {p['categorie']}" for p in personnel]
            idx = st.selectbox("Sélectionner un agent", range(len(options)), format_func=lambda i: options[i])
            pers = personnel[idx]
            pers_id = pers["id"]

            # Afficher les fiches existantes
            fiches = get_fiches_poste(pers_id)
            if fiches:
                with st.expander("📋 Fiches déjà générées", expanded=False):
                    for f in fiches:
                        st.text_area(f"Fiche #{f['id']}", f["contenu"], height=150, key=f"fiche_{f['id']}")

            # Nouvelle fiche
            with st.form("fiche_poste"):
                st.write(f"Générer une fiche pour **{pers['prenom']} {pers['nom']}** ({pers['categorie']})")
                missions = st.text_area(
                    "Missions principales",
                    f"Missions du {pers['poste']}:\n1. \n2. \n3. ",
                    height=100,
                )
                qualifications = st.text_area("Qualifications requises", height=80)
                submit = st.form_submit_button("📄 Générer la fiche de poste")
                if submit:
                    contenu = f"""FICHE DE POSTE - {pers['prenom']} {pers['nom']}
Catégorie: {pers['categorie']}
Poste: {pers['poste']}
Date: {datetime.now().strftime('%d/%m/%Y')}

MISSIONS PRINCIPALES:
{missions}

QUALIFICATIONS:
{qualifications}
"""
                    ajouter_fiche_poste(pers_id, pers["poste"], contenu)
                    st.success("✅ Fiche de poste générée!")
                    st.rerun()
        else:
            st.info("Aucun personnel disponible")

    # =====================================================================
    # TAB 4 : NOMINATIONS
    # =====================================================================
    with tab_personnel[3]:
        st.subheader("📜 Gestion des Nominations")

        with st.form("form_nomination"):
            personnel = get_personnel()
            if personnel:
                options = [f"{p['prenom']} {p['nom']} — {p['categorie']}" for p in personnel]
                idx = st.selectbox("Personnel", range(len(options)), format_func=lambda i: options[i], key="nom_select")
                pers = personnel[idx]
                poste = st.text_input("Poste de nomination", value=pers["poste"])
                date_nom = st.date_input("Date de nomination", value=datetime.now())
                session = st.text_input("Session / Référence")
                duree = st.text_input("Durée du mandat")
                motif = st.text_area("Motif / Observations")

                submit = st.form_submit_button("📜 Ajouter nomination")
                if submit:
                    ajouter_nomination(pers["id"], poste, str(date_nom), session, duree, motif)
                    st.success("✅ Nomination enregistrée!")
                    st.rerun()
            else:
                st.info("Aucun personnel disponible pour nomination")

        # Liste des nominations existantes
        nominations = get_nominations()
        if nominations:
            st.subheader("Nominations enregistrées")
            for nom in nominations[-10:]:  # 10 dernières
                with st.container(border=True):
                    pers_nom = get_personnel_by_id(nom["personnel_id"])
                    nom_complet = f"{pers_nom['prenom']} {pers_nom['nom']}" if pers_nom else "Inconnu"
                    st.write(f"**{nom_complet}** → {nom['poste']}")
                    st.caption(f"📅 {nom['date_nomination']} · {nom.get('session', '')} · {nom.get('duree', '')}")
                    if nom.get("motif"):
                        st.write(f"_{nom['motif']}_")


def missions_dynamiques():
    """Page de missions dynamiques avec IA Groq"""
    st.header("🎯 Missions Dynamiques")

    # Contexte du centre
    etablissements = get_etablissements()
    if etablissements:
        etab = etablissements[0]  # Prendre le premier établissement
        contexte_centre = f"""
        Centre: {etab["nom"]} ({etab["region"]}, {etab["district"]})
        Population: {etab.get("population", "Non spécifiée")}
        Personnel: {len(get_personnel())} agents
        """
    else:
        contexte_centre = "Centre de santé non configuré"

    # Sélectionner le type de contexte
    st.subheader("Sélectionner le contexte opérationnel")

    contextes_disponibles = [
        "Contexte Normal",
        "Épidémie en cours",
        "Campagne Vaccinale",
        "Formation du personnel",
        "Urgence sanitaire",
        "Nouvelles directives",
    ]

    contexte_select = st.selectbox("Contexte", contextes_disponibles)

    # Générer missions adaptées
    if st.button("🚀 Générer missions adaptées"):
        with st.spinner("Génération des missions avec IA Groq..."):
            try:
                # Préparer le prompt pour Groq
                prompt_system = """Tu es un expert en gestion de centres de santé en Côte d'Ivoire.
                
                Génère des missions spécifiques adaptées au contexte sélectionné.
                Utilise les données fournies et soyez très concret et opérationnel.
                
                Format de réponse:
                MISSIONS SPÉCIFIQUES:
                - Mission 1: Description précise
                - Mission 2: Description précise
                ...
                
                PRIORITAIRES:
                - Priorité 1: Description
                - Priorité 2: Description
                ...
                
                RESSOURCES NÉCESSAIRES:
                - Ressource 1: Quantité
                - Ressource 2: Quantité
                ...
                
                DÉLAI:
                - Temps estimé: X jours/semaines
                """

                prompt_user = f"""
                Contexte: {contexte_select}
                Centre: {contexte_centre}
                """

                # Appeler Groq (simulé pour l'exemple)
                contenu_groq = f"""
                MISSIONS SPÉCIFIQUES:
                - Renforcement de la surveillance épidémiologique dans les villages cibles
                - Organisation de séances de sensibilisation sur les mesures de prévention
                - Mise à jour des registres de consultation et de maternité
                - Coordination avec les agents de santé communautaire
                - Préparation du matériel pour les campagnes de vaccination
                
                PRIORITAIRES:
                - Surveillance active des cas suspects
                - Communication rapide avec le district sanitaire
                - Mobilisation des communautés locales
                - Préparation des kits de dépistage
                
                RESSOURCES NÉCESSAIRES:
                - Kits de dépistage: 50 unités
                - Masques chirurgicaux: 200 pièces
                - Gants: 100 paires
                - Désinfectant: 5 litres
                - Documents de sensibilisation: 100 exemplaires
                
                DÉLAI:
                - Temps estimé: 7 jours
                """

                # Afficher les résultats
                st.success("✅ Missions générées avec succès!")

                with st.expander("📋 Missions Spécifiques"):
                    st.text_area("Missions", value=contenu_groq, height=300)

                # Option de téléchargement
                from io import BytesIO

                doc = Document()
                doc.add_heading("MISSIONS SPÉCIFIQUES POUR {contexte_select}", level=1)
                doc.add_paragraph(f"Centre: {etab['nom']}")
                doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
                doc.add_paragraph()
                doc.add_paragraph(contenu_groq)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📥 Télécharger les missions",
                    data=buffer,
                    file_name=f"Missions_{contexte_select}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            except Exception as e:
                st.error(f"Erreur lors de la génération: {str(e)}")

    # Afficher les missions standards
    st.subheader("Missions Standards par Catégorie")

    categories = ["ide", "sfde", "cim", "cnm"]
    noms_categories = ["IDE", "SFDE", "CIM", "CNM"]

    for cat, nom in zip(categories, noms_categories):
        with st.expander(f"{nom}"):
            missions = contexte["contexte_dynamique"]["missions_standards"][cat]
            for i, mission in enumerate(missions, 1):
                st.write(f"{i}. {mission}")


def show_chat_interface(chatbot):
    """Interface de chat avec l'assistant IA"""
    st.title("💬 Assistant IA ESPC")

    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    # Chat input
    if prompt := st.chat_input("Posez votre question sur la gestion du centre..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        # Get context from the center
        etablissements = get_etablissements()
        context = (
            f"Centre: {etablissements[0]['nom']}"
            if etablissements
            else "Centre non configuré"
        )

        # Get AI response
        with st.chat_message("assistant"):
            with st.spinner("Réponse de l'assistant IA..."):
                response = chatbot.get_chat_response(prompt, context)
                st.write(response)

                # Add assistant response to chat history
                st.session_state.messages.append(
                    {"role": "assistant", "content": response}
                )

    # Quick actions
    st.subheader("🚀 Actions Rapides")

    quick_actions = [
        "📊 Générer rapport mensuel",
        "📅 Planifier réunion COGES",
        "👥 Voir planning personnel",
        "📋 Consulter checklist",
        "📈 Voir indicateurs",
        "📧 Envoyer notification WhatsApp",
        "🔍 Rechercher document",
        "⚙️ Paramètres application",
    ]

    for action in quick_actions:
        if st.button(action, key=f"quick_{action}"):
            st.info(f"Action: {action}")


def afficher_historique():
    """Afficher l'historique des documents"""
    st.header("📚 Historique des Documents")

    etablissements = get_etablissements()

    if not etablissements:
        st.info("Aucun établissement enregistré")
        return

    etab_dict = {e["nom"]: e for e in etablissements}
    nom_etab = st.selectbox("Établissement", list(etab_dict.keys()))

    if nom_etab:
        etab = etab_dict[nom_etab]
        periodes = get_periodes(etab["id"])

        if periodes:
            st.subheader(f"Évaluations de {etab['nom']}")

            for periode in periodes:
                st.markdown(
                    f"**{periode['periode']}** ({periode['date_debut']} au {periode['date_fin']})"
                )

                # Documents générés
                documents = get_documents(periode["id"])
                if documents:
                    st.subheader("📄 Documents générés")
                    for doc in documents:
                        st.markdown(
                            f"  - {doc['type_document']} ({doc['date_generation']})"
                        )
                else:
                    st.markdown("  - Aucun document généré")

                # Checklists
                checklists = get_checklists(periode["id"])
                if checklists:
                    st.subheader("📋 Checklists")
                    for checklist in checklists:
                        statut_color = (
                            "🟢"
                            if checklist["statut"] == "terminee"
                            else "🟡"
                            if checklist["statut"] == "en_cours"
                            else "🔴"
                        )
                        st.markdown(
                            f"  {statut_color} {checklist['type_checklist']} ({checklist['date_generation']})"
                        )

                        if checklist["statut"] == "en_cours":
                            if st.button(
                                "✅ Marquer comme terminée",
                                key=f"checklist_{checklist['id']}",
                            ):
                                maj_statut_checklist(checklist["id"], "terminee")
                                st.rerun()
                else:
                    st.markdown("  - Aucune checklist")

                # Nominations
                nominations = get_nominations()
                nomes_periode = [
                    n
                    for n in nominations
                    if n["date_nomination"]
                    and periode["date_debut"]
                    <= n["date_nomination"]
                    <= periode["date_fin"]
                ]
                if nomes_periode:
                    st.subheader("🏷️ Nominations")
                    for nom in nomes_periode:
                        st.markdown(
                            f"  - {nom['prenom']} {nom['nom']}: {nom['poste']} ({nom['date_nomination']})"
                        )
                else:
                    st.markdown("  - Aucune nomination")

                st.markdown("---")
        else:
            st.info("Aucune évaluation pour cet établissement")


if __name__ == "__main__":
    main()
