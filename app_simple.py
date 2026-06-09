"""
Application Streamlit - Générateur Documents ESPC
Conforme à la grille ESPC
"""

import streamlit as st
import os
import json
import calendar
import pandas as pd
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# =============================================================================
# CONFIGURATION
# =============================================================================

st.set_page_config(
    page_title="Générateur Documents ESPC", page_icon="🏥", layout="centered"
)

st.markdown(
    """
<style>
@media (max-width: 768px) {
    section[data-testid="stSidebar"] {
        width: 280px !important;
        min-width: 280px !important;
    }
    section[data-testid="stSidebar"] .stRadio > div {
        gap: 0.3rem !important;
    }
    section[data-testid="stSidebar"] .stRadio label {
        font-size: 0.9rem !important;
        padding: 6px 8px !important;
    }
    .block-container {
        padding-top: 1rem !important;
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    h1 { font-size: 1.5rem !important; }
    h2 { font-size: 1.2rem !important; }
    .stButton > button {
        width: 100% !important;
    }
}
@media (max-width: 480px) {
    h1 { font-size: 1.3rem !important; }
    section[data-testid="stSidebar"] {
        width: 100% !important;
        min-width: 100% !important;
    }
}
</style>
""",
    unsafe_allow_html=True,
)


# Configuration de la clé API Groq
import os
from dotenv import load_dotenv

# Déterminer le répertoire racine du projet (fonctionne même dans Streamlit)
if "__file__" in dir():
    _BASE_DIR = os.path.dirname(os.path.abspath(__file__))
else:
    _BASE_DIR = os.getcwd()

# Charger les variables d'environnement depuis .env (fallback local)
_env_path = os.path.join(_BASE_DIR, ".env")
load_dotenv(_env_path)

# Configuration de la clé API Groq via st.secrets (Streamlit Cloud) ou .env (local)
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except Exception:
    GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

if not GROQ_API_KEY:
    st.warning("🔑 Clé API Groq non configurée. La génération IA sera désactivée.")
    st.info("""
    **Pour configurer:**
    1. Fichier `.env` à la racine : `GROQ_API_KEY=votre_clé`
    2. Redémarre l'application
    """)
else:
    try:
        _ = Groq(api_key=GROQ_API_KEY)
    except TypeError:
        pass

# =============================================================================
# CONTEXTE DU CSR NAGNENEFOUN
# =============================================================================


def get_contexte_csr():
    """Génère le contexte du CSR NAGNENEFOUN"""
    contexte = """
CONTEXTE DU CSR NAGNENEFOUN (PORO, District KORHOGO 1):
- Population totale: 34055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau
- Activités principales: Consultations, CPN, Accouchements, PEV, Paludisme, VIH, Nutrition, IEC/CCC
- Maladies principales: Paludisme (76%), IRA (19%), Diarrhées (5%)
"""
    return contexte


# =============================================================================
# CHARGEMENT DES TEMPLATES
# =============================================================================


@st.cache_data
def charger_templates():
    """Charge les templates depuis le fichier JSON"""
    try:
        with open("templates.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return None


def sauvegarder_templates(templates):
    """Sauvegarde les templates dans le fichier JSON"""
    with open("templates.json", "w", encoding="utf-8") as f:
        json.dump(templates, f, ensure_ascii=False, indent=2)


templates = charger_templates()


def get_sections_template(doc_key):
    """Retourne les sections d'un template"""
    if templates and doc_key in templates:
        return templates[doc_key]["sections"]
    return []


# =============================================================================
# PROMPTS STRICTS - CONFORMES À LA GRILLE ESPC
# =============================================================================

PROMPTS = {
    "pv_reunion_mensuelle": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu rédiges un PROCÈS-VERBAL DE RÉUNION MENSUELLE conforme à la grille ESPC.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Personnel: Infirmiers, Sages-femmes, Aides-soignants, ASC, Filles de salle, Secrétaire, Gardiens
- Fréquence: Réunion mensuelle (tous les mois)

RÈGLES TRÈS IMPORTANTES:
1. Utilise les données réelles du CSR (personnel, statistiques, activités mensuelles)
2. Génère un contenu COMPLET sans placeholders ni [À compléter]
3. N'inclus JAMAIS de section "LISTE DE PRÉSENCE" ou "SIGNATURES" - ces parties sont saisies manuellement
4. Le PV doit contenir un TABLEAU DES DÉLIBÉRATIONS avec 4 colonnes: N° | Point discuté | Décisions prises | Responsable/Délai""",
        "user": """Génère PV RÉUNION MENSUELLE pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

THÈMES À TRAITER:
{themes}

SECTIONS À INCLURE:
{sections}

## I. OUVERTURE
- Date: __/__/{periode}
- Lieu: Salle de réunion du CSR {nom_etablissement}
- Heure: 08h30
- Président de séance: Chef de Centre
- Secrétaire de séance: Major
- Objet: Réunion mensuelle de {mois} {periode}

## II. MOT D'OUVERTURE
Le Chef de Centre souhaite la bienvenue à tout le personnel et rappelle l'importance de ces réunions mensuelles pour évaluer les activités du service, partager les informations et prendre des décisions collégiales pour l'amélioration de la qualité des soins au CSR {nom_etablissement}.

## III. LECTURE ET ADOPTION DU PV PRÉCÉDENT
Lecture du procès-verbal de la réunion du mois précédent.
- Observations/Suggestions: [RAS]
- Adoption: Le PV est adopté à l'unanimité

## IV. ORDRE DU JOUR
{themes}

## V. DÉLIBÉRATIONS (TABLEAU)

Génère un tableau avec 4 colonnes: N° | Point discuté | Décisions prises | Responsable/Délai
|---|---|---|---
| 1 | [Point 1] | [Décision] | [Responsable] - [Délai] |
| 2 | [Point 2] | [Décision] | [Responsable] - [Délai] |
| 3 | [Point 3] | [Décision] | [Responsable] - [Délai] |

Pour chaque thème de l'ordre du jour, génère un point avec décision et responsable.

## VI. DIVERS
- Questions diverses soulevées par le personnel
- Informations administratives
- Prochaine réunion

## VII. CLÔTURE
- Heure de clôture: 10h30
- Prochaine réunion: __/__/20..
- Le Chef de Centre remercie les participants et lève la séance

Génère le PV complet avec des décisions spécifiques et contextualisées pour {mois} {periode} en lien avec le thème: {themes}.""",
    },
    "pv_coges": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu rédiges un PROCÈS-VERBAL DE RÉUNION DU COMITÉ DE GESTION (COGES) conforme à la grille ESPC - Norme 1.02.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Fréquence: Réunion trimestrielle (T1-T4)
- Membres COGES: Président, Vice-Président, Secrétaire, Trésorier, Commissaire aux Comptes, Membres

RÈGLES TRÈS IMPORTANTES:
1. Le terme correct est COMITÉ DE GESTION ou COGES - JAMAIS "Conseil de Gestion"
2. Utilise les données réelles du CSR (personnel, statistiques, activités)
3. N'inclus JAMAIS de section "LISTE DE PRÉSENCE" ou "SIGNATURES" - ces parties sont saisies manuellement
4. Inclus un TABLEAU DES DÉLIBÉRATIONS avec 4 colonnes: N° | Point discuté | Décisions prises | Responsable/Délai""",
        "user": """Génère un PROCÈS-VERBAL DE RÉUNION DU COMITÉ DE GESTION (COGES) pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

THÈMES À TRAITER:
{themes}

SECTIONS À INCLURE:
{sections}

## I. OUVERTURE
- Date: __/__/{periode}
- Lieu: Salle de réunion du CSR {nom_etablissement}
- Heure: 09h00
- Président de séance: [Président du COGES]
- Secrétaire de séance: [Secrétaire du COGES]
- Objet: Réunion trimestrielle du COGES du {trimestre} {periode}

## II. MOT D'OUVERTURE
Le Président du COGES souhaite la bienvenue à tous les membres et remercie le Chef de Centre pour l'organisation. Il rappelle le rôle du COGES dans la gestion participative du centre de santé et l'importance de ces réunions trimestrielles.

## III. LECTURE ET ADOPTION DU PV PRÉCÉDENT
Lecture du procès-verbal de la réunion COGES du trimestre précédent.
- Observations/Suggestions:
- Adoption: Le PV est adopté à l'unanimité

## IV. ORDRE DU JOUR
{themes}

## V. DÉLIBÉRATIONS (TABLEAU)

Génère un tableau avec 4 colonnes: N° | Point discuté | Décisions prises | Responsable/Délai
|---|---|---|---
| 1 | [Point COGES 1] | [Décision] | [Responsable] - [Délai] |
| 2 | [Point COGES 2] | [Décision] | [Responsable] - [Délai] |
| 3 | [Point COGES 3] | [Décision] | [Responsable] - [Délai] |

Points typiques COGES:
- Gestion financière et recettes du centre
- États des dépenses et budget
- État du personnel et des équipements
- Projets de développement du CSR
- Relations avec le District Sanitaire
- Difficultés et recommandations

## VI. DÉCISIONS ADOPTÉES
- Décision N°1: [Résumé de la décision]
- Décision N°2: [Résumé de la décision]
- Décision N°3: [Résumé de la décision]

## VII. QUESTIONS DIVERSES
- Prochaine réunion COGES: [Date]
- Recommandations au Chef de Centre

## VIII. CLÔTURE
- Heure de clôture: 11h00
- Prochaine réunion: __/__/20..
- Le Président remercie les participants et lève la séance

Génère le PV complet adapté au CSR {nom_etablissement} pour le {trimestre} {periode} en lien avec les thèmes COGES.""",
    },
    "pv_ag": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu rédiges un PROCÈS-VERBAL D'ASSEMBLÉE GÉNÉRALE conforme à la grille ESPC - Norme 1.03.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Fréquence: Assemblée Générale annuelle
- Participants: COGES, Personnel du centre, Représentants de la communauté, Autorités locales

RÈGLES:
1. Utilise les données réelles du CSR
2. N'inclus JAMAIS de section "LISTE DE PRÉSENCE" ou "SIGNATURES" - ces parties sont saisies manuellement
3. Inclus un TABLEAU DES POINTS ABORDÉS avec 4 colonnes: N° | Point discuté | Décisions/Votes | Observations""",
        "user": """Génère un PROCÈS-VERBAL D'ASSEMBLÉE GÉNÉRALE pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

THÈMES À TRAITER:
{themes}

SECTIONS À INCLURE:
{sections}

## I. OUVERTURE
- Date: __/__/{periode}
- Lieu: CSR {nom_etablissement}
- Heure: 09h00
- Président de séance: Chef de Centre / Président du COGES
- Secrétaire de séance: Major
- Objet: Assemblée Générale Annuelle du CSR {nom_etablissement}

## II. MOT D'OUVERTURE
Le Chef de Centre ouvre l'Assemblée Générale et souhaite la bienvenue à tous les participants (COGES, Personnel, Communauté, Autorités). Il présente le rapport moral de l'année {periode} et remercie les partenaires.

## III. ÉLECTION DU BUREAU DE SÉANCE
- Président de séance: [Élu(e)]
- Secrétaire: [Élu(e)]
- Assesseurs: [Élu(e)s]

## IV. LECTURE ET ADOPTION DU PV PRÉCÉDENT
Lecture du procès-verbal de l'Assemblée Générale de [Année N-1].
- Adoption à l'unanimité

## V. ADOPTION DE L'ORDRE DU JOUR
{themes}

## VI. DÉLIBÉRATIONS ET DÉCISIONS (TABLEAU)

Génère un tableau avec 4 colonnes: N° | Point discuté | Décisions/Votes adoptés | Observations
|---|---|---|---
| 1 | [Point AG 1] | [Décision] | [Obs] |
| 2 | [Point AG 2] | [Décision] | [Obs] |
| 3 | [Point AG 3] | [Décision] | [Obs] |

Points typiques AG annuelle:
- Rapport d'activités de l'année écoulée
- Rapport financier et budget
- Projets et perspectives
- Élection/Renouvellement des membres COGES
- Questions diverses

## VII. RECOMMANDATIONS ADOPTÉES
1. [Recommandation 1]
2. [Recommandation 2]
3. [Recommandation 3]

## VIII. CLÔTURE
- Heure de clôture: 12h00
- Prochaine Assemblée Générale: Année [N+1]
- Le Président remercie les participants et lève la séance

Génère le PV complet adapté au CSR {nom_etablissement} pour l'année {periode}.""",
    },
    "rapport_supervision_asc": {
        "system": """Tu es un assistant spécialisé dans les rapports de supervision ASC pour les CSR en Côte d'Ivoire.

RÈGLES:
1. Utilise les données réelles du CSR (3 ASC, villages, activités communautaires)
2. N'inclus JAMAIS de "SIGNATURES" ou "LISTE PRÉSENCE" - ces parties sont saisies manuellement""",
        "user": """Génère RAPPORT SUPERVISION ASC pour {nom_etablissement}.
{contexte}

STRUCTURE (Norme 14.01 - sans SIGNATURES):
I. INFOS GÉNÉRALES
II. PLAN SUPERVISION ANNUEL
III. EFFECTIFS SUPERVISÉS (3 ASC)
IV. LISTE ASC
V. GRILLE SUPERVISION
VI. ACTIVITÉS (sensibilisation, dépistage, référence)
VII. RÉSULTATS
VIII. DISPO MÉDICAMENTS
IX. DIFFICULTÉS
X. RECOMMANDATIONS
XI. TRANSMISSION DISTRICT

Génère un rapport contextualisé.""",
    },
    "rapport_plaintes": {
        "system": """Tu es un assistant spécialisé dans les rapports de boîte à suggestions pour CSR en Côte d'Ivoire.
RÈGLES: Génère un contenu COMPLET contextualisé, sans SIGNATURES.""",
        "user": """Génère RAPPORT BOÎTE À SUGGESTIONS pour {nom_etablissement}.
{contexte}

INFORMATIONS:
- Nombre de suggestions: {nb_suggestions}
- Types de suggestions: {types_suggestions}
- Actions menées: {actions}

STRUCTURE DU DOCUMENT:
{sections}
I. CONTEXTE
II. STATISTIQUES
III. TYPES SUGGESTIONS
IV. ACTIONS MENÉES
V. RÉSULTATS
VI. RECOMMANDATIONS

Génère un rapport contextualisé.""",
    },
    # =============================================================================
    # FICHE DE POSTE - Templates par catégorie - Conforme Grille ESPC 2.01
    # =============================================================================
    "fiche_poste": {
        "system": """Tu es un responsable des ressources humaines dans un centre de santé (CSR) en Côte d'Ivoire. Tu rédiges une FICHE DE POSTE officielle conforme à la grille ESPC - Norme 2.01 a,b,c.

CONTEXTE DU CSR:
- Établissement: CSR NAGNENEFOUN
- District Sanitaire: KORHOGO 1 (Région Sanitaire du PORO)
- Population desservie: 34 055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau

RÈGLES ABSOLUES:
1. Génère un document professionnel officiel avec en-tête administrative complète
2. Adapte le contenu (missions, qualifications, compétences) à la catégorie de personnel choisie
3. Sois précis et réaliste - chaque catégorie a des missions spécifiques
4. N'inclus JAMAIS de zone de signature ou cadre de signature (signé manuellement)
5. La section des missions doit être détaillée avec au moins 5-8 missions concrètes
6. Les qualifications doivent être précises (diplômes, années d'expérience, inscriptions ordinales)
7. Inclus les relations fonctionnelles internes et externes adaptées à chaque catégorie""",
        "user": """Génère une FICHE DE POSTE officielle pour {nom_etablissement}.

----------------------------------------------------
INFORMATIONS GÉNÉRALES
----------------------------------------------------
CATÉGORIE: {categorie_poste}
TITRE DU POSTE: {titre_poste}
NOM DU TITULAIRE: {nom_titulaire}
SUPÉRIEUR HIÉRARCHIQUE: {superieur}
RÉGIME DE TRAVAIL: {regime_travail}

CONTEXTE DU CENTRE:
{contexte}

----------------------------------------------------
STRUCTURE À GÉNÉRER (conserver impérativement):
----------------------------------------------------
{sections}

## I. IDENTIFICATION DU POSTE
- Établissement: {nom_etablissement}
- Intitulé du poste: {titre_poste}
- Catégorie: {categorie_poste}
- Supérieur hiérarchique direct: {superieur}
- Lieu d'affectation: CSR NAGNENEFOUN, District Sanitaire de KORHOGO 1, Région Sanitaire du PORO
- Régime de travail: {regime_travail}

## II. MISSIONS PRINCIPALES
{missions_poste}

Rédige chaque mission sous forme d'un paragraphe détaillé expliquant concrètement ce que la personne fait au quotidien. Minimum 7 missions détaillées.

## III. QUALIFICATIONS REQUISES
{qualifications_poste}

Détaille par catégorie :
- Formation académique requise (diplôme précis)
- Inscription ordinale / autorisation d'exercer
- Expérience professionnelle minimale
- Formations complémentaires souhaitables

## IV. COMPÉTENCES ET APTITUDES
{competences_poste}

- Compétences techniques spécifiques
- Compétences relationnelles
- Aptitudes personnelles (rigueur, discrétion, etc.)
- Capacités d'adaptation et d'initiative

## V. RESPONSABILITÉS
{responsabilites_poste}

Détail des responsabilités par domaine : clinique, administratif, managérial, formation

## VI. MOYENS ET RESSOURCES MISES À DISPOSITION
{moyens_poste}

Équipement, matériel, locaux, ressources humaines encadrées

## VII. RELATIONS FONCTIONNELLES
- Interne: Direction du centre, Personnel soignant et administratif, COGES
- Externe: District Sanitaire de KORHOGO 1, Communauté, Partenaires techniques (PSI, UNICEF, etc.)

## VIII. OBSERVATIONS PARTICULIÈRES
- Conditions d'exercice spécifiques
- Astreintes et gardes (si applicable)
- Formation continue obligatoire

Génère la fiche de poste complète et professionnelle. Sois TRÈS DÉTAILLÉ et précis dans chaque section.""",
    },
    # =============================================================================
    # FICHE DE NOMINATION - 12 types conformes Grille ESPC
    # =============================================================================
    "fiche_nomination": {
        "system": """Tu es un responsable administratif de centre de santé (CSR) en Côte d'Ivoire. Tu rédiges des documents administratifs officiels.

FORMAT OFFICIEL (EN-TÊTE OBLIGATOIRE):
RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail

MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA COUVERTURE MALADIE UNIVERSELLE
RÉGION SANITAIRE DU PORO
DISTRICT SANITAIRE DE KORHOGO 1

CSR NAGNENEFOUN

CONTEXTE GÉNÉRAL:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau

RÈGLES ABSOLUES:
1. L'en-tête républicaine officielle DOIT être incluse
2. Adapte le format (Note de service / Arrêté / Fiche de poste / Liste) selon le TYPE DE DOCUMENT demandé
3. Génère un document officiel complet et réaliste
4. N'inclus JAMAIS de zone de signature (les originaux sont signés manuellement)
5. Le document doit être PRÊT À IMPRIMER sur papier officiel
6. Pour les "Notes de service" : inclure un cachet "Vu et approuvé" sans cadre de signature
7. Pour les "Arrêtés" : format juridique avec "Considérant que..." et articles
8. Pour les "Fiches de poste signées" : format fiche technique détaillée""",
        "user": """Génère un DOCUMENT ADMINISTRATIF OFFICIEL pour {nom_etablissement}.

----------------------------------------------------
PARAMÈTRES DU DOCUMENT
----------------------------------------------------
TYPE DE DOCUMENT: {type_nomination}
OBJET / FONCTION: {objet_nomination}
NORME ESPC: {reference_norme}
SIGNATAIRE REQUIS: {signataire}
NOM DU BÉNÉFICIAIRE: {nom_beneficiaire}
FONCTION DU BÉNÉFICIAIRE: {fonction_beneficiaire}
DATE PRISE D'EFFET: {date_effet}
NUMÉRO D'ORDRE: {numero_ordre}

CONTEXTE DU CENTRE:
{contexte}

----------------------------------------------------
STRUCTURE À GÉNÉRER:
----------------------------------------------------
{sections}

I. EN-TÊTE OFFICIEL
RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA COUVERTURE MALADIE UNIVERSELLE
RÉGION SANITAIRE DU PORO
DISTRICT SANITAIRE DE KORHOGO 1
{nom_etablissement}
---

{type_nomination} N° {numero_ordre}

II. PRÉAMBULE / VU
(Inclure les textes légaux : Code de la santé publique, Loi N°... portant organisation sanitaire, etc.)

III. DISPOSITIONS / DÉSIGNATION
- Article 1: {nom_beneficiaire} est nommé(e) en qualité de {fonction_beneficiaire}
- Article 2: {objet_nomination}
- Missions et attributions détaillées

IV. ATTRIBUTIONS / MISSIONS DÉTAILLÉES
(Description complète des missions confiées au bénéficiaire)

V. DISPOSITIONS FINALES
- SIGNATAIRE: {signataire}
- DATE D'EFFET: {date_effet}
- DIFFUSION: {nom_etablissement}, District Sanitaire KORHOGO 1, Intéressé(e), Archives

---
Cachet "Vu et approuvé" - Signature manuelle requise

Génère le document officiel complet, prêt à imprimer. Adapte le style et le format selon le TYPE DE DOCUMENT.""",
    },
    "programme_reunions_trimestrielles": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères un PROGRAMME DE RÉUNIONS TRIMESTRIELLES stratégique.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau
- Activités: Consultations, CPN, Accouchements, PEV, Paludisme, VIH, Nutrition, IEC/CCC
- Maladies principales: Paludisme (76%), IRA (19%), Diarrhées (5%)

FORMAT OBLIGATOIRE DU TABLEAU (markdown):
Génère le calendrier sous forme de tableau markdown avec 5 colonnes et 4 lignes (T1 à T4).
Le séparateur doit être: |---|---|----|----|----|
Chaque cellule doit contenir du texte EXPLICITE, précis et développé (pas de simples listes à puces).""",
        "user": """Génère PROGRAMME RÉUNIONS TRIMESTRIELLES pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. INFORMATIONS GÉNÉRALES
- Établissement: {nom_etablissement}
- Année: {periode}
- Population: 34 055 habitants
- District: KORHOGO 1 (PORO)

## II. CALENDRIER STRATÉGIQUE DES RÉUNIONS TRIMESTRIELLES

Génère un tableau markdown avec 5 colonnes: Trimestre | Mois | Thème de la réunion | Activités clés à mener | Responsable

Exemple de structure (à remplir avec du contenu détaillé et contextualisé pour le CSR):
| Trimestre | Mois | Thème de la réunion | Activités clés à mener | Responsable |
|---|---|---|---|---|
| T1 | Janvier - Mars | Planification annuelle et stratégies PEV | Définir le plan d'action annuel 2026, organiser la campagne de vaccination PEV de routine, planifier les séances de CPN et accouchements, budgétiser les intrants et consommables, former le personnel sur les nouveaux protocoles paludisme | Chef de Centre, Major, Équipe encadrement |
| T2 | Avril - Juin | Suivi des activités et campagne paludisme | Évaluer les indicateurs CPN et accouchements du 1er trimestre, organiser la campagne de chimio-prévention du paludisme saisonnier (CPS), assurer le PEV de rattrapage des enfants perdus de vue, intensifier les séances IEC/CCC sur le paludisme et la nutrition | IDE, Sage-femme, ASC, Chargé PEV |
| T3 | Juillet - Septembre | Nutrition et surveillance épidémiologique | Analyser les données nutritionnelles et les indicateurs VIH/PTME, renforcer la surveillance des IRA et diarrhées (saison des pluies), organiser les séances de dépistage VIH et malnutrition, évaluer l'état des stocks de médicaments et intrants | Infirmier, Chargé VIH, Nutritionniste, Pharmacien |
| T4 | Octobre - Décembre | Évaluation annuelle et perspectives N+1 | Réaliser le bilan annuel des activités et indicateurs PEP/DBS, élaborer le rapport d'activités annuel, préparer la planification opérationnelle N+1, évaluer la performance du personnel et les besoins en formation, organiser la réunion COGES bilan | Chef de Centre, COGES, Toute l'équipe |

## III. OBJECTIFS STRATÉGIQUES PAR TRIMESTRE
Décrire en 2-3 phrases par trimestre les objectifs poursuivis.

## IV. MODALITÉS D'AFFICHAGE
- Lieu d'affichage: Tableau d'affichage du centre, Salle de réunion du personnel
- Période d'affichage: De janvier à décembre {periode}
- Responsable de l'affichage: Chef de Centre adjoint

Génère le tableau ci-dessus avec le contenu détaillé adapté au CSR NAGNENEFOUN.""",
    },
    "calendrier_nettoyage": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères un CALENDRIER DE NETTOYAGE conforme à la Norme 6.01 de la grille ESPC.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Infrastructure: Dispensaire, Maternité, Château d'eau, Hall d'attente, Toilettes, Cour
- Activités: Consultations, CPN, Accouchements, PEV, Paludisme, VIH, Nutrition

FORMAT OBLIGATOIRE:
- Génère le calendrier sous forme de TABLEAU avec minimum 5 colonnes
- Le séparateur doit être: |---|---|---|---|---|
- Chaque cellule doit contenir des informations précises, explicites et détaillées""",
        "user": """Génère CALENDRIER DE NETTOYAGE pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. PRÉSENTATION
Le présent calendrier de nettoyage est établi conformément à la Norme 6.01 de la grille ESPC. Il définit les tâches de nettoyage, les fréquences, les responsables et les produits à utiliser pour chaque zone du CSR {nom_etablissement}.

Établissement: {nom_etablissement}
Période: {periode}

## II. ZONES À NETTOYER
Les zones suivantes sont identifiées au CSR: {zones}

Fréquences applicables: {frequences}

## III. CALENDRIER DE NETTOYAGE (TABLEAU)

Génère un tableau avec 5 colonnes: Zone | Fréquence | Activités de nettoyage détaillées | Responsable | Produits/Matériel

Le tableau doit être complet et couvrir:

| Zone | Fréquence | Activités de nettoyage détaillées | Responsable | Produits/Matériel |
|---|---|---|---|---|
| Salle de consultation | Quotidien | Nettoyer et désinfecter la table d'examen après chaque patient, laver le sol à l'eau de javel diluée, dépoussiérer les surfaces de travail (bureau, armoires), vider et nettoyer la poubelle médicale, vérifier la disponibilité du savon et de l'eau | Infirmier / Aide-soignant | Eau de javel 0.5%, Savon liquide, Gants ménagers, Chiffons, Balai, Poubelle |
| Salle de consultation | Hebdomadaire | Laver les murs et les vitres, désinfecter les poignées de porte et interrupteurs, nettoyer les luminaires, vérifier l'état du matériel et signaler les dégradations, faire l'inventaire des produits d'entretien, aérer la salle | Aide-soignant | Détergent, Éponge, Seau, Brosse, Gants |
| Maternité | Quotidien | Nettoyer et désinfecter la table d'accouchement après chaque usage, laver le sol avec détergent puis eau de javel, nettoyer les incubateurs et berceaux, vider les poubelles DASRI, vérifier les kits d'accouchement, approvisionner en eau potable | Sage-femme / Aide-soignant | Eau de javel 0.5%, Savon antiseptique, Gants stériles, Poubelle DASRI, Chiffons |
| Maternité | Hebdomadaire | Nettoyer les murs et les plafonds, laver les rideaux et les linges, désinfecter le matériel d'accouchement, vérifier la stérilisation du matériel, nettoyer les rangements et tiroirs, contrôler la chaîne de froid des vaccins | Sage-femme | Autoclave, Détergent, Gants, Brosse, Savon |
| Hall d'attente | Quotidien | Balayer et laver le sol, nettoyer les bancs et chaises, vider les poubelles, vérifier la propreté des affiches et supports IEC/CCC, approvisionner le point d'eau | Aide-soignant / Planton | Balai, Savon, Seau, Eau de javel, Poubelle |
| Hall d'attente | Hebdomadaire | Laver les murs, vitres et portes, désinfecter les rampes et poignées, nettoyer les ventilateurs et luminaires, ranger et organiser la documentation IEC, inspecter l'état de la toiture | Aide-soignant | Détergent, Éponge, Savon, Gants |
| Toilettes | Quotidien | Nettoyer et désinfecter les cuvettes et urinoirs, laver le sol et les murs carrelés, approvisionner en eau et savon, vider les poubelles, vérifier le bon fonctionnement des chasses d'eau, désodoriser | Aide-soignant / Planton | Acide chlorhydrique dilué, Eau de javel, Balai-brosse, Gants, Désodorisant |
| Toilettes | Hebdomadaire | Nettoyer les plafonds et les aérations, détartrer les canalisations, inspecter l'état des installations sanitaires, signaler les fuites et réparations nécessaires, désinfecter les poignées de porte | Aide-soignant | Détartrant, Brosse métallique, Gants, Éponge |
| Cour / Extérieur | Quotidien | Balayer la cour et les abords, ramasser les déchets solides, vider les poubelles extérieures, arroser les espaces verts, vérifier l'état du château d'eau et de la clôture | Planton / Gardien | Balai, Brouette, Râteau, Tuyau d'arrosage, Poubelle |
| Cour / Extérieur | Hebdomadaire | Désherber les allées et abords, nettoyer le caniveau et les regards, laver les murs extérieurs, vérifier le système d'évacuation des eaux usées, contrôler l'état des fosses septiques | Planton / Gardien | Désherbant, Pelle, Pioche, Gants |

## IV. PLAN DE NETTOYAGE QUOTIDIEN DU SOIR (FILLES DE SALLE)

Personnel dédié: Filles de salle (agents d'entretien) - passages chaque soir

Génère un tableau avec 5 colonnes: Zone | Horaire | Tâches de nettoyage du soir | Responsable | Durée

| Zone | Horaire | Tâches de nettoyage du soir | Responsable | Durée |
|---|---|---|---|---|
| Salle de consultation | 17h00 - 17h30 | Balayer et laver le sol à grande eau, désinfecter la table d'examen, vider la poubelle médicale, nettoyer le bureau et les armoires, vérifier les stocks de savon et désinfectant, passer la serpillière avec eau de javel diluée | Fille de salle | 30 min |
| Maternité | 17h30 - 18h15 | Nettoyer et désinfecter la table d'accouchement, laver le sol avec détergent puis eau de javel, désinfecter les incubateurs et berceaux, vider les poubelles DASRI, nettoyer les toilettes de la maternité, approvisionner en eau potable et savon liquide, ranger le linge propre | Fille de salle | 45 min |
| Hall d'attente | 18h15 - 18h45 | Balayer et laver le sol, nettoyer les bancs et chaises avec désinfectant, vider les poubelles, dépoussiérer les affiches IEC/CCC, vérifier la propreté du point d'eau, nettoyer les portes et poignées | Fille de salle | 30 min |
| Toilettes (patients) | 18h45 - 19h15 | Nettoyer et désinfecter les cuvettes et urinoirs avec détergent et désinfectant, laver le sol et les murs carrelés, approvisionner en eau et savon, vider les poubelles, désodoriser, vérifier l'état des chasses d'eau et signaler les fuites | Fille de salle | 30 min |
| Toilettes (personnel) | 19h15 - 19h30 | Nettoyer et désinfecter la cuvette et le lavabo, laver le sol, approvisionner en savon et papier hygiénique, vider la poubelle, désodoriser | Fille de salle | 15 min |
| Bureau Chef de Centre | 19h30 - 19h45 | Dépoussiérer le bureau et les étagères, nettoyer le sol, vider la corbeille, aérer la pièce | Fille de salle | 15 min |
| Pharmacie / Magasin | 19h45 - 20h00 | Balayer le sol, dépoussiérer les étagères et rangements, vérifier l'état des murs et plafonds, signaler tout problème d'humidité ou d'infiltration | Fille de salle | 15 min |
| Cour / Abords | 20h00 - 20h30 | Balayer la cour et le parking, ramasser et évacuer les déchets solides, vérifier la propreté du château d'eau et des abords, fermer les portails et vérifier la clôture | Fille de salle / Gardien | 30 min |

## V. OBSERVATIONS
- Les produits d'entretien doivent être stockés dans un local fermé et identifié
- Le personnel doit porter des gants et des bottes pour les tâches de nettoyage
- Les fiches de suivi de nettoyage doivent être signées chaque jour
- La vérification mensuelle est assurée par le Major / Chef de Centre
- Les filles de salle sont tenues de signer la fiche de présence quotidienne chaque soir après le nettoyage

## VI. AFFICHAGE
- Lieu d'affichage: Dans chaque zone concernée et au tableau d'affichage central
- Responsable du suivi: Major du centre
- Période d'affichage: {periode} (renouvelé chaque année)

Génère le calendrier complet avec le tableau détaillé ci-dessus adapté au CSR NAGNENEFOUN.""",
    },
    "calendrier_reunions_mensuelles": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères un CALENDRIER DE RÉUNIONS MENSUELLES.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau
- Activités: Consultations, CPN, Accouchements, PEV, Paludisme, VIH, Nutrition, IEC/CCC
- Maladies principales: Paludisme (76%), IRA (19%), Diarrhées (5%)

FORMAT OBLIGATOIRE DU TABLEAU:
Génère le calendrier sous forme de tableau markdown avec 4 colonnes et 12 lignes (une par mois).
Le séparateur doit être: |---|---|---|---|
Chaque cellule doit contenir du texte EXPLICITE, précis et développé pour le mois concerné.""",
        "user": """Génère CALENDRIER RÉUNIONS MENSUELLES pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. INFORMATIONS GÉNÉRALES
- Établissement: {nom_etablissement}
- Année: {periode}

## II. CALENDRIER DES RÉUNIONS MENSUELLES

Génère un tableau avec 4 colonnes: Mois | Thème de la réunion | Activités clés à mener | Responsable

Pour chaque mois, le contenu doit être contextualisé au CSR NAGNENEFOUN (paludisme 76%, PEV, CPN/accouchements, VIH, nutrition, etc.).

Exemple de structure du tableau (à adapter avec le vrai contexte du CSR):
| Mois | Thème de la réunion | Activités clés à mener | Responsable |
|---|---|---|---|
| Janvier | Planification opérationnelle annuelle et stratégies PEV | Élaborer le plan d'action 2026, organiser les stratégies avancées PEV, planifier les CPN et accouchements du 1er trimestre, définir le calendrier des séances IEC/CCC | Chef de Centre, Major, Équipe |
| Février | Campagne paludisme et suivi des indicateurs | Analyser les indicateurs de morbidité paludisme (76% des cas), organiser la distribution des MII, renforcer le dépistage et traitement rapide du paludisme, superviser les ASC dans les villages | Infirmier, ASC, Chargé Paludisme |
| Mars | Gestion des stocks et approvisionnement en intrants | Évaluer l'état des stocks de médicaments (antipaludiques, ARV, vaccins), passer les commandes au district, vérifier la chaîne de froid et la conservation des vaccins | Pharmacien, Major |
| Avril | Santé maternelle et CPN/Accouchements | Évaluer les indicateurs CPN (consultations prénatales), analyser les données accouchements, renforcer la PTME et le dépistage VIH chez les femmes enceintes, organiser les séances de planification familiale | Sage-femme, Chargé VIH |
| Mai | Campagne CPS et activités PEV de rattrapage | Organiser la chimio-prévention du paludisme saisonnier (CPS), réaliser le PEV de rattrapage des enfants perdus de vue, intensifier les consultations curatives IRA (19%) et diarrhées (5%) | IDE, Chargé PEV, ASC |
| Juin | Surveillance épidémiologique et hygiène | Analyser les données épidémiologiques, renforcer la surveillance des maladies à potentiel épidémique, évaluer l'hygiène et l'assainissement du centre, planifier les activités de désinfection | Chef de Centre, Major, Équipe |
| Juillet | Nutrition et prise en charge des cas de malnutrition | Évaluer les indicateurs nutritionnels, organiser le dépistage de la malnutrition chez les enfants, renforcer les séances d'éducation nutritionnelle, coordonner avec le programme nutrition | Infirmier, Nutritionniste, ASC |
| Août | VIH/PTME et activités communautaires | Analyser les indicateurs VIH (dépistage, mise sous ARV), évaluer la PTME chez les femmes enceintes, superviser les activités des ASC dans les aires sanitaires, planifier les séances de sensibilisation | Chargé VIH, Sage-femme, ASC |
| Septembre | Gestion des ressources et formation du personnel | Évaluer les besoins en formation continue, organiser une séance de recyclage sur les protocoles, vérifier les équipements et infrastructures, planifier les réparations/maintenance | Chef de Centre, Major |
| Octobre | Campagne PEV intensifiée et IEC/CCC | Organiser la campagne de vaccination de masse, intensifier les séances IEC/CCC sur le paludisme et le VIH, réaliser les consultations mobiles dans les villages reculés | Chargé PEV, IDE, ASC |
| Novembre | Évaluation à mi-parcours et ajustements | Faire le point sur les indicateurs PEP/DBS, évaluer l'atteinte des objectifs, ajuster les stratégies pour le dernier trimestre, préparer les rapports d'activités | Chef de Centre, Toute l'équipe |
| Décembre | Bilan annuel et perspectives N+1 | Élaborer le rapport annuel d'activités, présenter les résultats au COGES, planifier les activités de l'année N+1, organiser la réunion bilan avec le personnel | Chef de Centre, COGES, Toute l'équipe |

## III. OBSERVATIONS
- Fréquence des réunions: 1 réunion par mois
- Durée recommandée: 2h
- Lieu: Salle de réunion du CSR NAGNENEFOUN
- Participants: Personnel du centre, ASC (selon ordre du jour), COGES (si nécessaire)

## IV. AFFICHAGE
- Lieu d'affichage: Tableau d'affichage du centre, Salle de réunion
- Période d'affichage: De janvier à décembre {periode}
- Responsable: Chef de Centre

Génère le tableau ci-dessus avec le contenu détaillé adapté au CSR NAGNENEFOUN.""",
    },
    "grille_supervision_asc": {
        "system": """Tu es un superviseur de centre de santé (CSR) en Côte d'Ivoire. Tu génères une GRILLE DE SUPERVISION ASC conforme à la grille ESPC.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Activités ASC: Sensibilisation communautaire, Dépistage paludisme, Référence des cas, Causeries éducatives, Suivi des patients perdus de vue
- Maladies principales: Paludisme (76%), IRA (19%), Diarrhées (5%)

RÈGLE ABSOLUE: Génère un TABLEAU structuré avec colonnes, notation sur 3 niveaux (A=Bon, B=Moyen, C=Faible) et observations.""",
        "user": """Génère GRILLE SUPERVISION ASC pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. INFORMATIONS GÉNÉRALES
- Établissement: {nom_etablissement}
- Période: {periode}
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- ASC supervisé: [Nom et prénom(s) de l'ASC]
- Village/Aire sanitaire: [Nom du village]
- Date de supervision: [Date]
- Superviseur: Chef de Centre / Major / Infirmier superviseur

## II. CRITÈRES DE SUPERVISION (TABLEAU)

Génère un tableau d'évaluation avec 5 colonnes: Critère | Indicateurs évalués | Note (A/B/C) | Observations | Actions recommandées

Le tableau doit couvrir les critères suivants avec des indicateurs précis:

| Critère | Indicateurs évalués | Note (A/B/C) | Observations | Actions recommandées |
|---|---|---|---|---|
| 1. Accueil et relation communautaire | Qualité de l'accueil des patients, Tenue et présentation de l'ASC, Discrétion et confidentialité, Respect des heures de travail, Relation avec les relais communautaires | A B C | (observations détaillées) | (actions concrètes) |
| 2. Sensibilisation et IEC/CCC | Organisation des causeries éducatives, Thèmes abordés (paludisme, VIH, nutrition, hygiène), Utilisation des supports IEC, Tenue de registre des séances, Participation communautaire | A B C | (observations détaillées) | (actions concrètes) |
| 3. Dépistage et prise en charge paludisme | Utilisation des TDR paludisme, Traitement correct des cas simples, Application du protocole national, Gestion des cas graves (référence), Tenue du registre paludisme | A B C | (observations détaillées) | (actions concrètes) |
| 4. Référence des cas vers le CSR | Identification des cas à référer (paludisme grave, malnutrition, complications grossesse), Remplissage de la fiche de référence, Suivi des patients référés, Contre-référence du CSR, Taux de référence | A B C | (observations détaillées) | (actions concrètes) |
| 5. Documentation et rapports | Tenue correcte des registres (consultations, TDR, références), Remplissage des rapports mensuels, Transmission des données au CSR, Archivage des documents, Utilisation des fiches de suivi | A B C | (observations détaillées) | (actions concrètes) |
| 6. Gestion des intrants et médicaments | Gestion des stocks TDR et antipaludiques, Conservation des médicaments (chaleur, humidité), Tenue de la fiche de stock, Gestion des périmés, Mouvement des intrants | A B C | (observations détaillées) | (actions concrètes) |
| 7. Hygiène et salubrité du poste | Propreté du poste de santé/ASC, Gestion des déchets biomédicaux, Hygiène des mains, État du matériel, Disponibilité de l'eau et du savon | A B C | (observations détaillées) | (actions concrètes) |
| 8. Participation aux activités du CSR | Présence aux réunions mensuelles, Participation aux campagnes (PEV, CPS), Collaboration avec le personnel du centre, Participation aux séances de démonstration, Implication dans les enquêtes | A B C | (observations détaillées) | (actions concrètes) |

## III. ÉVALUATION GLOBALE
- Total des A: /8
- Total des B: /8
- Total des C: /8
- Appréciation générale: Satisfaisante / Moyenne / Insuffisante
- Décision: Poursuite / Supervision renforcée / Avertissement / Remplacement

## IV. OBSERVATIONS ET RECOMMANDATIONS
- Points forts de l'ASC:
- Points à améliorer:
- Recommandations du superviseur:
- Date de la prochaine supervision:

## V. SIGNATURES
- Signature de l'ASC: ____________________
- Signature du Superviseur: ____________________
- Date: ____/____/______

Génère la grille complète avec le tableau d'évaluation détaillé, les notes et les rubriques à remplir.""",
    },
    "liste_coges": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères une LISTE DU PERSONNEL COGES conforme au format grille ESPC, prête à imprimer et afficher.

CONTEXTE CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants

FORMAT: Génère un tableau structuré avec 4 colonnes: N° | Nom & Prénoms | Fonction au COGES | Contacts/Téléphone""",
        "user": """Génère LISTE PERSONNEL COGES pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. EN-TÊTE OFFICIEL
RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA COUVERTURE MALADIE UNIVERSELLE
RÉGION SANITAIRE DU PORO
DISTRICT SANITAIRE DE KORHOGO 1
{nom_etablissement}

## II. LISTE DU PERSONNEL COGES (TABLEAU)

Génère un tableau avec 4 colonnes: N° | Nom & Prénoms | Fonction au COGES | Contacts/Téléphone

Le tableau doit reprendre les membres ci-dessous avec leurs informations complètes:
{membres}

Colonne par colonne, voici ce qui doit figurer:
- N°: Numérotation de 1 à N
- Nom & Prénoms: Nom complet du membre
- Fonction au COGES: Président, Vice-Président, Secrétaire, Trésorier, Commissaire aux Comptes, Membre
- Contacts/Téléphone: Numéro de téléphone mobile

## III. RÉCAPITULATIF
- Nombre total de membres: 
- Nombre d'hommes: 
- Nombre de femmes:
- Date de mise en place: Janvier {periode}
- Durée du mandat: 2 ans renouvelable

## IV. CONTACTS UTILES
- Président: [Nom] - Tél: [numéro]
- Vice-Président: [Nom] - Tél: [numéro]
- Secrétaire: [Nom] - Tél: [numéro]
- Trésorier: [Nom] - Tél: [numéro]
- Point focal CSR: Chef de Centre CSR NAGNENEFOUN - Tél: [numéro]

## V. OBSERVATIONS

## VI. APPROBATION
- Vu et approuvé par le Chef de Centre:
- Date: ____/____/{periode}
- Cachet et signature:

Génère la liste complète avec le tableau détaillé, prête à l'emploi, avec les mentions officielles de la République de Côte d'Ivoire.""",
    },
    "plan_action_infections_nosocomiales": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères un PLAN D'ACTION CONTRE LES INFECTIONS NOSOCOMIALES conforme à la grille ESPC.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Infrastructure: Dispensaire, Maternité, Château d'eau
- Activités médicales: Consultations curatives, CPN, Accouchements, PEV, Chirurgie mineure, Soins infirmiers

FORMAT OBLIGATOIRE:
- Génère le plan sous forme de TABLEAU avec 6 colonnes
- Le séparateur doit être: |---|---|---|---|---|---|
- Chaque cellule doit contenir du texte EXPLICITE, précis et contextualisé""",
        "user": """Génère PLAN D'ACTION INFECTIONS NOSOCOMIALES pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. CONTEXTE ET JUSTIFICATION
Les infections nosocomiales constituent un problème majeur de santé publique dans les centres de santé. Le présent plan d'action vise à prévenir et contrôler les infections associées aux soins au CSR {nom_etablissement}.

Établissement: {nom_etablissement}
Période: {periode}
District: KORHOGO 1 (PORO)
Population couverte: 34 055 habitants

Domaines d'action: {activites}

## II. OBJECTIFS
- Objectif général: Réduire l'incidence des infections nosocomiales de 80% au CSR NAGNENEFOUN
- Objectifs spécifiques:
  1. Renforcer les pratiques d'hygiène des mains du personnel
  2. Assurer la désinfection et la stérilisation du matériel médical
  3. Améliorer la gestion des déchets biomédicaux
  4. Former le personnel sur les protocoles de prévention des infections

## III. PLAN D'ACTION DÉTAILLÉ (TABLEAU)

Génère un tableau avec 6 colonnes: N° | Domaine d'action | Activités spécifiques | Responsable | Périodicité | Indicateur de suivi

Le tableau doit être complet et couvrir tous les aspects:

| N° | Domaine d'action | Activités spécifiques | Responsable | Périodicité | Indicateur de suivi |
|---|---|---|---|---|---|
| 1 | Hygiène des mains | Former tout le personnel sur les 5 moments du lavage des mains selon l'OMS, installer des points de lavage avec savon liquide et essuie-mains à usage unique dans chaque salle de soins, organiser des séances de démonstration pratique du lavage des mains au savon et à la solution hydro-alcoolique, afficher les protocoles de lavage des mains dans chaque zone de soins, évaluer mensuellement la conformité du lavage des mains par observation directe | Major, Infirmier hygiéniste | Trimestrielle | % de personnel formé, Nombre de points de lavage fonctionnels, Taux de conformité aux 5 moments |
| 2 | Désinfection et stérilisation | Établir un protocole écrit de désinfection du matériel médical (table d'examen, spéculums, pinces), organiser la stérilisation à l'autoclave des instruments réutilisables (boîtes d'accouchement, instruments chirurgicaux), contrôler mensuellement l'efficacité de la stérilisation avec des tests biologiques, désinfecter les surfaces et le matériel entre chaque patient, remplacer le matériel défectueux et les gants troués | Sage-femme, Infirmier major | Hebdomadaire | % de matériel stérilisé conforme, Nombre de protocoles affichés, Registre de stérilisation tenu à jour |
| 3 | Gestion des déchets biomédicaux | Mettre en place un système de tri des déchets (boîtes jaunes DASRI, poubelles noires ordinaires), former le personnel sur le tri et l'élimination des déchets biomédicaux conformément à la réglementation nationale, organiser l'incinération ou l'enlèvement des déchets DASRI par le district, tenir un registre de collecte et d'évacuation des déchets, approvisionner en boîtes de sécurité et sacs plastiques de couleur | Major, Aide-soignant, Planton | Quotidienne | % de déchets correctement triés, Nombre de boîtes de sécurité disponibles, Registre d'évacuation rempli |
| 4 | Hygiène de l'environnement | Assurer le nettoyage quotidien des salles de soins et de la maternité avec de l'eau de javel diluée à 0.5%, désinfecter les surfaces (poignées de porte, rampes, interrupteurs) avec un détergent-désinfectant, organiser la désinfection hebdomadaire des murs, plafonds et vitres, contrôler la qualité de l'eau du château d'eau et des points d'utilisation, entretenir les installations sanitaires et la plomberie | Aide-soignant, Fille de salle, Gardien | Quotidienne / Hebdomadaire | Fiche de suivi de nettoyage signée, Résultats analyse eau, Nombre de zones conformes |
| 5 | Précautions standard et isolement | Établir et diffuser les précautions standard pour tous les soins (port de gants, masque, blouse), organiser la mise en isolement des patients suspects d'infection contagieuse, approvisionner en équipements de protection individuelle (EPI) : gants, masques, charlottes, sur-blouses, tabliers, former le personnel à l'utilisation correcte des EPI, gérer les accidents d'exposition au sang (AES) | Chef de Centre, Major | Continue | % de soignants utilisant correctement les EPI, Nombre d'AES déclarés, Stock d'EPI disponible |
| 6 | Surveillance des infections | Mettre en place un registre de suivi des infections nosocomiales, notifier et investiguer tout cas suspect d'infection post-opératoire ou post-partum, analyser mensuellement les données de surveillance, organiser une réunion trimestrielle du comité de lutte contre les infections, élaborer un rapport trimestriel sur les infections nosocomiales | Chef de Centre, Major | Trimestrielle | Nombre d'infections notifiées, Taux d'incidence des infections, Rapports trimestriels produits |
| 7 | Formation et supervision | Organiser une séance de formation annuelle sur la prévention des infections nosocomiales pour tout le personnel, former les nouveaux agents dès leur arrivée, organiser des séances de recyclage semestrielles, superviser les pratiques d'hygiène lors des soins, évaluer les connaissances du personnel par un pré-test et post-test | Chef de Centre, Major | Semestrielle | % de personnel formé, Score moyen aux tests de connaissances, Nombre de séances de supervision effectuées |

## IV. SUIVI ET ÉVALUATION
- Le suivi est assuré par le Major du centre sous la supervision du Chef de Centre
- Une réunion trimestrielle du comité de lutte contre les infections est organisée
- Un rapport d'évaluation annuel est produit et transmis au District Sanitaire de KORHOGO 1

## V. AFFICHAGE
- Le plan est affiché dans la salle de réunion du personnel et dans chaque zone de soins
- Responsable: Chef de Centre / Major
- Période d'affichage: {periode} (renouvelé chaque année)

Génère le plan complet avec le tableau détaillé ci-dessus adapté au CSR NAGNENEFOUN.""",
    },
    "rapport_formation": {
        "system": """Tu es un responsable de centre de santé (CSR) en Côte d'Ivoire. Tu génères un RAPPORT DE FORMATION DU PERSONNEL conforme à la grille ESPC - Norme 2.01.

CONTEXTE DU CSR NAGNENEFOUN:
- District: KORHOGO 1 (PORO)
- Population: 34 055 habitants
- Personnel: Infirmiers, Sages-femmes, Aides-soignants, Filles de salle, ASC, Secrétaire, Gardiens

FORMAT OBLIGATOIRE:
- Rapport TRIMESTRIEL (T1: Jan-Mar, T2: Avr-Juin, T3: Juil-Sept, T4: Oct-Déc)
- Génère un tableau RÉCAPITULATIF DES SESSIONS avec 5 colonnes
- Le séparateur doit être: |---|---|---|---|---|
- Les listes de présence sont gérées manuellement (NE PAS inclure dans le rapport)
- Chaque cellule doit contenir des informations précises et contextualisées""",
        "user": """Génère RAPPORT DE FORMATION DU PERSONNEL pour {nom_etablissement} - {periode}.

CONTEXTE DU CENTRE:
{contexte}

SECTIONS À INCLURE:
{sections}

## I. INFORMATIONS GÉNÉRALES
Établissement: {nom_etablissement}
Période: {periode}
Trimestre concerné: {trimestre}
Domaine de formation: {domaine}
Date de la formation: {date_formation}
Durée: {duree}
Formateur(s): {formateur}
Lieu: Salle de réunion du CSR {nom_etablissement}

## II. OBJECTIFS DE LA FORMATION
- Objectif général: Renforcer les capacités du personnel du CSR {nom_etablissement} en {domaine}
- Objectifs spécifiques:
  1. Acquérir les connaissances théoriques sur {domaine}
  2. Maîtriser les techniques pratiques liées à {domaine}
  3. Améliorer la qualité des prestations dans le domaine
  4. Uniformiser les pratiques conformément aux protocoles nationaux

## III. RÉCAPITULATIF DES SESSIONS (TABLEAU)
| N° | Thème/Session | Contenu abordé | Méthode pédagogique | Durée |
|---|---|---|---|---|
| 1 | Introduction et contexte | Présentation du contexte, des objectifs, rappel des protocoles nationaux | Exposé interactif | 30 min |
| 2 | Module théorique | Définition, concepts clés, grands principes, protocoles et directives nationales | Exposé, Support PPT, Documentation | 1h 30 min |
| 3 | Démonstration pratique | Démonstration des gestes techniques, procédures, manipulation du matériel | Démonstration, Simulation | 1h |
| 4 | Travaux pratiques | Mise en situation, études de cas, exercices pratiques par groupe de 2-3 | Atelier pratique, Jeux de rôles | 1h 30 min |
| 5 | Discussion et échanges | Questions-réponses, partage d'expériences, clarification des difficultés | Table ronde interactive | 30 min |
| 6 | Évaluation | Pré-test et post-test, questionnaire de satisfaction | Questionnaire individuel | 30 min |
| Total | - | - | - | 5h 30 min |

## IV. ÉVALUATION DE LA FORMATION
- Pré-test: Score moyen de [X]%
- Post-test: Score moyen de [Y]%
- Progression: +[Z] points
- Taux de satisfaction: [W]%
- Points forts:
- Points à améliorer:

## V. DIFFICULTÉS RENCONTRÉES
- [À compléter]

## VI. RECOMMANDATIONS
1. Organiser une session de recyclage
2. Superviser la mise en pratique des acquis lors des soins quotidiens
3. Mettre à disposition les documents techniques de référence
4. Évaluer l'impact de la formation à 3 mois

## VII. SUIVI
- Responsable du suivi: Major du CSR et Chef de Centre
- Modalités: Supervision des pratiques, observation directe
- Prochaine session: [À définir]

Génère le rapport complet adapté au CSR {nom_etablissement} pour le trimestre {trimestre} en {domaine} avec {nb_participants} participants.""",
    },
    # =============================================================================
    # NOTES DE SERVICE (6 types exigés par la Grille ESPC)
    # =============================================================================
    "note_service": {
        "system": """Tu es un Chef de centre de santé (CSR) en Côte d'Ivoire. Tu rédiges des notes de service officielles.
Format officiel: République de Côte d'Ivoire - Union – Discipline – Travail / MSHP / Région Sanitaire / District / CSR.
Utilise le contexte réel du centre (population, personnel, infrastructures, activités).
N'inclus JAMAIS de zone de signature - signée manuellement.""",
        "user": """Génère une NOTE DE SERVICE pour {nom_etablissement}.

TYPE DE NOTE: {type_note}
DATE: {date_note}
NOM DU RESPONSABLE DÉSIGNÉ: {nom_responsable}
FONCTION: {fonction_responsable}

CONTEXTE:
{contexte}

SECTIONS:
{sections}

I. PRÉAMBULE / CONSIDÉRANTS
II. DÉSIGNATION
III. ATTRIBUTIONS / MISSIONS
IV. APPLICATION
V. DIFFUSION

Génère une note de service officielle et conforme à la réglementation.""",
    },
}

# =============================================================================
# LISTE DES DOCUMENTS
# =============================================================================

DOCUMENTS_LIST = [
    ("PV Réunion Mensuelle", "pv_reunion_mensuelle"),
    ("PV Réunion COGES", "pv_coges"),
    ("PV Assemblée Générale", "pv_ag"),
    ("Rapport Supervision ASC", "rapport_supervision_asc"),
    ("Rapport Plaintes/Suggestions", "rapport_plaintes"),
    ("Fiche de Poste", "fiche_poste"),
    ("Fiche de Nomination", "fiche_nomination"),
    ("Programme Réunions Trimestrielles", "programme_reunions_trimestrielles"),
    ("Calendrier Nettoyage Centre", "calendrier_nettoyage"),
    ("Calendrier Réunions Mensuelles", "calendrier_reunions_mensuelles"),
    ("Grille Supervision ASC", "grille_supervision_asc"),
    ("Liste Personnel COGES", "liste_coges"),
    ("Plan Action Infections Nosocomiales", "plan_action_infections_nosocomiales"),
    ("Plan Supervision ASC", "plan_supervision_asc"),
    ("Rapport Formation Personnel", "rapport_formation"),
    ("Liste Personnel Centre", "liste_personnel_centre"),
    ("Note de Service", "note_service"),
]

# =============================================================================
# FONCTIONS
# =============================================================================


def generer_avec_groq(system_prompt, user_prompt):
    """Génère du contenu avec l'API Groq en créant un client frais à chaque appel"""
    from dotenv import load_dotenv
    import os

    # Déterminer le bon dossier de base
    if "__file__" in dir():
        _base = os.path.dirname(os.path.abspath(__file__))
    else:
        _base = os.getcwd()

    # Charger la clé depuis .env
    load_dotenv(os.path.join(_base, ".env"))

    # Récupérer la clé (priorité: variable d'env existante > .env)
    api_key = os.environ.get("GROQ_API_KEY", "")

    if not api_key:
        return "⚠️ Clé API Groq non configurée."

    try:
        _client = Groq(api_key=api_key)
        response = _client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            max_tokens=4000,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Erreur: {str(e)}"


def creer_document_word(titre, contenu, meta=None):
    doc = Document()
    heading = doc.add_heading(titre, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta:
        for key, value in meta.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))

    doc.add_paragraph()

    # Parcours intelligent: détection des tableaux markdown et des paragraphes
    lignes = contenu.split("\n")
    i = 0

    while i < len(lignes):
        para = lignes[i]
        para_strip = para.strip()

        if not para_strip:
            i += 1
            continue

        # Détection d'un tableau markdown (ligne qui commence par |)
        if para_strip.startswith("|") and para_strip.endswith("|"):
            # Collecter toutes les lignes du tableau
            rows = []
            while i < len(lignes) and lignes[i].strip().startswith("|"):
                row_text = lignes[i].strip()
                # Ignorer la ligne de séparation (|---|---|)
                if "---" in row_text or "—" in row_text:
                    i += 1
                    continue
                # Extraire les cellules
                cells = [c.strip() for c in row_text.split("|")[1:-1]]
                rows.append(cells)
                i += 1

            # Créer le tableau Word
            if rows:
                nb_cols = max(len(row) for row in rows)
                word_table = doc.add_table(rows=len(rows), cols=nb_cols)
                word_table.style = "Light Grid Accent 1"

                for row_idx, row_data in enumerate(rows):
                    for col_idx in range(nb_cols):
                        cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                        cell = word_table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        # Mettre en gras la première ligne (en-tête)
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                doc.add_paragraph()
            continue

        # Titres et paragraphes
        if (
            any(
                x in para_strip.upper()
                for x in [
                    "I.",
                    "II.",
                    "III.",
                    "IV.",
                    "V.",
                    "VI.",
                    "VII.",
                    "VIII.",
                    "IX.",
                    "X.",
                ]
            )
            and len(para_strip) < 60
        ):
            doc.add_heading(para_strip, level=1)
        elif (
            any(
                x in para_strip.upper()
                for x in [
                    "CONTEXTE",
                    "INFORMATIONS",
                    "DÉLIBÉRATIONS",
                    "DÉCISIONS",
                    "SIGNATURES",
                    "LISTE",
                    "OBSERVATIONS",
                    "CALENDRIER",
                    "OBJECTIFS",
                    "ACTIVITÉS",
                    "AFFICHAGE",
                ]
            )
            and len(para_strip) < 50
        ):
            doc.add_heading(para_strip, level=2)
        elif para_strip.startswith("##"):
            doc.add_heading(para_strip.replace("#", "").strip(), level=2)
        elif len(para_strip) > 0:
            p = doc.add_paragraph(para_strip)

        i += 1

    return doc


# =============================================================================
# FORMULAIRES PRÉ-CONFORMES
# =============================================================================


def get_form_fields(doc_type):
    fields = {}

    if doc_type == "pv_reunion_mensuelle":
        st.markdown("### 📋 PV RÉUNION MENSUELLE (Conforme Grille ESPC)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["mois"] = st.selectbox(
            "Mois",
            [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ],
        )
        # L'IA génère automatiquement tout le contenu

    elif doc_type == "pv_coges":
        st.markdown("### 📋 PV COGES (Trimestriel - Conforme Norme 1.02)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["trimestre"] = st.selectbox(
            "Trimestre de réunion",
            [
                "T1 - Janvier-Mars",
                "T2 - Avril-Juin",
                "T3 - Juillet-Septembre",
                "T4 - Octobre-Décembre",
            ],
        )

    elif doc_type == "pv_ag":
        st.markdown("### 📋 PV ASSEMBLÉE GÉNÉRALE (Annuelle - Conforme Norme 1.03)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Année de l'AG", "2026")

    elif doc_type == "rapport_supervision_asc":
        st.markdown("### 📋 RAPPORT SUPERVISION ASC (Conforme Norme 14.01)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        # L'IA génère automatiquement

    elif doc_type == "rapport_plaintes":
        st.markdown("### 📋 RAPPORT BOÎTE À SUGGESTIONS")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["nb_suggestions"] = st.text_input("Nombre de suggestions reçues", "")
        fields["types_suggestions"] = st.text_area(
            "Types de suggestions (séparées par ;)",
            "Amélioration accueil ; Attente ; Hygiène ; Médicaments ; Autre",
        )
        fields["actions"] = st.text_area(
            "Actions menées (séparées par ;)",
            "Analyse des suggestions ; Réunion de réflexion ; Plan d'action",
        )

    elif doc_type == "fiche_poste":
        st.markdown("### 📋 FICHE DE POSTE (Template par catégorie)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")

        categories_poste = {
            "Chef de Centre": {
                "titre": "Chef de Centre de Santé",
                "superieur": "Médecin Chef du District Sanitaire de KORHOGO 1",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Assurer la coordination générale des activités du CSR ; Superviser l'équipe du personnel ; Garantir la qualité des soins ; Élaborer et exécuter le plan d'action annuel ; Assurer la gestion administrative et financière ; Rendre compte au District Sanitaire ; Présider les réunions de staff ; Assurer le suivi des indicateurs de performance",
                "qualifications": "Diplôme d'État de Sage-Femme ou Infirmier Diplômé d'État (IDE) ; Diplôme en Santé Publique (optionnel) ; Expérience minimale de 3 ans dans un centre de santé",
                "competences": "Leadership et management d'équipe ; Maîtrise des outils de planification ; Bonne connaissance du système de santé ivoirien",
                "responsabilites": "Responsable de la bonne marche du centre ; Responsable de la gestion du personnel ; Responsable de la qualité des soins",
                "moyens": "Bureau équipé (ordinateur, imprimante) ; Véhicule de service (si disponible) ; Budget de fonctionnement",
            },
            "Major / Sage-Femme": {
                "titre": "Major / Sage-Femme Diplômé(e) d'État",
                "superieur": "Chef de Centre",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Superviser les soins obstétricaux et néonatals ; Gérer la maternité et les activités de CPN ; Assurer les accouchements et soins post-partum ; Coordonner les activités de PF et PTME ; Gérer les urgences obstétricales",
                "qualifications": "Diplôme d'État de Sage-Femme (DESF) ; Inscription à l'Ordre national des Sages-Femmes ; Expérience minimale de 2 ans en maternité",
                "competences": "Maîtrise des techniques obstétricales ; Capacité à gérer les urgences vitales ; Leadership et encadrement d'équipe",
                "responsabilites": "Responsable des activités de la maternité ; Responsable de la qualité des soins maternels et infantiles",
                "moyens": "Salle d'accouchement équipée ; Matériel de réanimation néonatale ; Kits d'accouchement",
            },
            "Infirmier Diplômé d'État (IDE)": {
                "titre": "Infirmier Diplômé d'État (IDE)",
                "superieur": "Chef de Centre / Major",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Assurer les consultations curatives ; Administrer les médicaments et traitements ; Réaliser les soins infirmiers ; Participer aux activités de vaccination (PEV) ; Assurer la prise en charge du paludisme",
                "qualifications": "Diplôme d'État d'Infirmier (DEI) ; Inscription à l'Ordre national des Infirmiers ; Expérience minimale de 1 an",
                "competences": "Maîtrise des protocoles de soins infirmiers ; Capacité à diagnostiquer les pathologies courantes ; Bon relationnel avec les patients",
                "responsabilites": "Responsable des soins infirmiers quotidiens ; Responsable de la tenue des registres",
                "moyens": "Bocson médical complet ; Stéthoscope, tensiomètre ; Matériel de soins",
            },
            "Chargé(e) de Programme (PEV/Paludisme/VIH)": {
                "titre": "Chargé(e) de Programme",
                "superieur": "Chef de Centre",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Planifier et coordonner les activités du programme ; Gérer les intrants et vaccins ; Organiser les séances de vaccination ; Superviser les ASC ; Collecter et analyser les données du programme",
                "qualifications": "Diplôme d'État d'Infirmier ou Sage-Femme ; Formation spécifique au programme ; Expérience minimale de 2 ans",
                "competences": "Maîtrise des protocoles du programme ; Capacité à former et superviser ; Gestion de données",
                "responsabilites": "Responsable de l'atteinte des objectifs du programme ; Responsable de la gestion des intrants",
                "moyens": "Bureau partagé avec équipement informatique ; Matériel de supervision",
            },
            "Aide-Soignant(e)": {
                "titre": "Aide-Soignant(e)",
                "superieur": "Major / IDE de garde",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Assister l'infirmier dans les soins courants ; Préparer et entretenir le matériel de soins ; Assurer l'hygiène des salles de soins ; Accueillir et orienter les patients",
                "qualifications": "Certificat d'Aide-Soignant(e) ; Expérience en milieu de santé (1 an minimum)",
                "competences": "Sens de l'organisation et de la propreté ; Capacité à suivre les consignes ; Empathie et respect des patients",
                "responsabilites": "Responsable de la propreté des zones de soins ; Responsable de l'entretien du matériel",
                "moyens": "Matériel d'entretien et de nettoyage ; Équipements de protection",
            },
            "Secrétaire / Agent Administratif": {
                "titre": "Secrétaire / Agent Administratif",
                "superieur": "Chef de Centre",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Assurer la gestion administrative et le secrétariat ; Tenir les registres administratifs ; Accueillir et orienter les usagers ; Saisir les rapports et documents",
                "qualifications": "Diplôme de Secrétariat ou BTS Administration ; Maîtrise des outils bureautiques",
                "competences": "Excellente présentation ; Maîtrise de la bureautique ; Organisation et rigueur",
                "responsabilites": "Responsable de la tenue des archives ; Responsable de la gestion du courrier",
                "moyens": "Bureau équipé (ordinateur, imprimante, téléphone)",
            },
            "Agent d'Entretien / Fille de Salle": {
                "titre": "Agent d'Entretien",
                "superieur": "Major / IDE de garde",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Assurer le nettoyage et la désinfection des locaux ; Gérer les déchets biomédicaux ; Entretenir le linge ; Assurer la propreté des sanitaires",
                "qualifications": "Niveau primaire ou secondaire ; Formation aux règles d'hygiène",
                "competences": "Sens de la propreté ; Connaissance des techniques d'entretien ; Ponctualité",
                "responsabilites": "Responsable de la propreté des locaux ; Responsable du tri des déchets",
                "moyens": "Produits d'entretien et de désinfection ; Équipements de protection",
            },
            "Gardien / Planton": {
                "titre": "Gardien / Planton",
                "superieur": "Chef de Centre",
                "regime": "Temps plein - Garde 24h/24 (roulement)",
                "missions": "Assurer la sécurité du centre ; Contrôler les entrées et sorties ; Surveiller les locaux ; Ouvrir et fermer le centre selon les horaires",
                "qualifications": "Niveau primaire minimum ; Expérience en surveillance (optionnel)",
                "competences": "Vigilance ; Ponctualité et fiabilité ; Sens des responsabilités",
                "responsabilites": "Responsable de la sécurité du centre ; Responsable des clés",
                "moyens": "Local de garde ; Lampe torche ; Téléphone de service",
            },
            "ASC (Agent de Santé Communautaire)": {
                "titre": "Agent de Santé Communautaire (ASC)",
                "superieur": "Chargé de Programme / Chef de Centre",
                "regime": "Temps partiel - Communautaire",
                "missions": "Mener des séances de sensibilisation ; Dépister les cas suspects ; Faire la référence des patients ; Distribuer les MILD ; Animer les causeries éducatives",
                "qualifications": "Niveau secondaire (BEPC minimum) ; Formation ASC validée ; Parlant la langue locale",
                "competences": "Capacité à communiquer ; Connaissance de la communauté ; Dynamisme",
                "responsabilites": "Responsable des activités communautaires ; Responsable du matériel de sensibilisation",
                "moyens": "Kit ASC ; MILD et préservatifs ; Supports IEC/CCC ; Vélo (si disponible)",
            },
            "Pharmacien / Chargé Pharmacie": {
                "titre": "Pharmacien / Chargé de la Pharmacie",
                "superieur": "Chef de Centre",
                "regime": "Temps plein - 40h/semaine",
                "missions": "Gérer les stocks de médicaments ; Assurer la dispensation ; Tenir les fiches de stock ; Commander et réceptionner les médicaments ; Contrôler les périmés",
                "qualifications": "Diplôme de Technicien Supérieur en Pharmacie ou IDE formé ; Expérience de 2 ans",
                "competences": "Maîtrise de la gestion des stocks ; Connaissance des médicaments essentiels ; Rigueur et organisation",
                "responsabilites": "Responsable de la gestion des médicaments ; Responsable de la pharmacie",
                "moyens": "Pharmacie équipée ; Logiciel de gestion de stock ; Mobilier de rangement",
            },
        }

        cat_choice = st.selectbox(
            "Catégorie du personnel", list(categories_poste.keys())
        )
        cat_data = categories_poste[cat_choice]
        fields["categorie_poste"] = cat_choice
        fields["titre_poste"] = cat_data["titre"]
        nom_titulaire = st.text_input("Nom du titulaire du poste", "")
        fields["nom_titulaire"] = (
            nom_titulaire if nom_titulaire else "[Nom du titulaire]"
        )
        fields["superieur"] = cat_data["superieur"]
        fields["regime_travail"] = cat_data["regime"]
        fields["missions_poste"] = cat_data["missions"]
        fields["qualifications_poste"] = cat_data["qualifications"]
        fields["competences_poste"] = cat_data["competences"]
        fields["responsabilites_poste"] = cat_data["responsabilites"]
        fields["moyens_poste"] = cat_data["moyens"]
        fields["match_categorie"] = f"CATÉGORIE: {cat_choice}"

    elif doc_type == "fiche_nomination":
        st.markdown("### 📋 FICHE DE NOMINATION (12 types conformes Grille ESPC)")

        types_nomination = [
            (
                "Note de service",
                "Responsable de l'ESPC (Chef d'établissement)",
                "1.01 d",
                "Directeur Départemental (DD)",
            ),
            (
                "Note de service",
                "Responsable de chaque service (dispensaire, maternité...)",
                "1.01 d",
                "Responsable de l'ESPC",
            ),
            (
                "Arrêté préfectoral / sous-préfectoral",
                "Mise en place du COGES",
                "1.02 a",
                "Préfet / Sous-préfet",
            ),
            ("Note de service", "Point focal CMU", "4.03 c", "Responsable de l'ESPC"),
            (
                "Fiche de poste signée",
                "Agent d'accueil et d'orientation",
                "4.02 b",
                "Agent + Responsable ESPC",
            ),
            (
                "Note de service",
                "Responsable de l'hygiène hospitalière",
                "6.01 a",
                "Responsable ESPC",
            ),
            (
                "Note de service",
                "Responsable de la gestion des déchets biomédicaux",
                "6.05 a",
                "Responsable ESPC",
            ),
            (
                "Note de service + fiche de poste",
                "Gestionnaire des médicaments (pharmacie)",
                "11.01 g",
                "Agent + Responsable ESPC",
            ),
            (
                "Fiche de poste signée",
                "Personnel qualifié pour accouchements (SF/IDE/Maïeuticien)",
                "7.02 a",
                "Agent + Responsable ESPC",
            ),
            (
                "Fiche de poste signée",
                "Tout agent (vérifié sur 3 noms)",
                "2.01 b",
                "Agent + Responsable ESPC",
            ),
            (
                "Liste officielle nominative",
                "Agents de santé communautaire (ASC)",
                "14.01 a",
                "District sanitaire",
            ),
            (
                "Grille de supervision ASC",
                "Acte de supervision ASC",
                "14.01 c",
                "ASC + Superviseur",
            ),
        ]

        type_index = st.selectbox(
            "Type de document",
            range(len(types_nomination)),
            format_func=lambda i: f"{types_nomination[i][0]} - {types_nomination[i][1]} (Norme {types_nomination[i][2]})",
        )

        type_choisi = types_nomination[type_index]
        fields["type_nomination"] = type_choisi[0]
        fields["objet_nomination"] = type_choisi[1]
        fields["reference_norme"] = type_choisi[2]
        fields["signataire"] = type_choisi[3]
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["nom_beneficiaire"] = st.text_input("Nom du bénéficiaire", "")
        fields["fonction_beneficiaire"] = st.text_input(
            "Fonction du bénéficiaire",
            type_choisi[1].split(" (")[0] if " (" in type_choisi[1] else type_choisi[1],
        )
        fields["date_effet"] = st.text_input("Date de prise d'effet (JJ/MM/AAAA)", "")
        fields["numero_ordre"] = st.text_input(
            "Numéro d'ordre", "____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN"
        )
        st.caption(
            f"**Signataire requis :** {type_choisi[3]} | **Norme ESPC :** {type_choisi[2]}"
        )

        fields["match_categorie"] = (
            f"TYPE: {type_choisi[0]} | OBJET: {type_choisi[1]} | NORME: {type_choisi[2]}"
        )

    elif doc_type == "programme_reunions_trimestrielles":
        st.markdown("### 📋 PROGRAMME RÉUNIONS TRIMESTRIELLES")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")

    elif doc_type == "calendrier_nettoyage":
        st.markdown("### 📋 CALENDRIER NETTOYAGE (Conforme Norme 6.01 - À AFFICHER)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["zones"] = st.text_area(
            "Zones à nettoyer (séparées par ;)",
            "Salle de consultation ; Maternité ; Hall d'attente ; Toilettes ; Cour",
        )
        fields["frequences"] = st.text_area(
            "Fréquences (séparées par ;)", "Quotidien ; Hebdomadaire ; Mensuel"
        )

    elif doc_type == "calendrier_reunions_mensuelles":
        st.markdown("### 📋 CALENDRIER RÉUNIONS MENSUELLES")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")

    elif doc_type == "grille_supervision_asc":
        st.markdown("### 📋 GRILLE SUPERVISION ASC (À signer ASC + Superviseur)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["criteria"] = st.text_area(
            "Critères de supervision (séparés par ;)",
            "Accueil ; Sensibilisation ; Dépistage ; Référence ; Documentation",
        )

    elif doc_type == "liste_coges":
        st.markdown("### 📋 LISTE PERSONNEL COGES (Format grille ESPC)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["membres"] = st.text_area(
            "Membres COGES (Nom ; Fonction ; Téléphone)",
            "Koffi Kouassi ; Président ; 0102030405\n"
            "Ahoua N'Guessan ; Vice-Président ; 0102030406\n"
            "Konan Bertille ; Secrétaire ; 0102030407\n"
            "Kouamé Yao ; Trésorier ; 0102030408\n"
            "Soro Fatoumata ; Commissaire aux Comptes ; 0102030409\n"
            "Touré Mamadou ; Membre ; 0102030410",
        )

    elif doc_type == "liste_personnel_centre":
        st.markdown("### 📋 LISTE PERSONNEL CENTRE (Statique CSR NAGNENEFOUN)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")

    elif doc_type == "plan_action_infections_nosocomiales":
        st.markdown("### 📋 PLAN ACTION INFECTIONS NOSOCOMIALES")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["activites"] = st.text_area(
            "Activités principales (séparées par ;)",
            "Formation personnel ; Désinfection ; Lavage des mains ; Gestion des déchets ; Surveillance",
        )

    elif doc_type == "plan_supervision_asc":
        st.markdown("### 📋 PLAN SUPERVISION ASC (Plan annuel)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["activites"] = st.text_area(
            "Activités de supervision (séparées par ;)",
            "Inspection terrain ; Formation ; Dépistage communautaire ; Référence",
        )

    elif doc_type == "rapport_formation":
        st.markdown("### 📋 RAPPORT FORMATION (Trimestriel - Norme 2.01)")
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["periode"] = st.text_input("Période (Année)", "2026")
        fields["trimestre"] = st.selectbox(
            "Trimestre",
            [
                "T1 - Janvier-Mars",
                "T2 - Avril-Juin",
                "T3 - Juillet-Septembre",
                "T4 - Octobre-Décembre",
            ],
        )
        fields["domaine"] = st.text_input(
            "Domaine de formation", "ex: Paludisme, PEV, VIH, CPN, Hygiène..."
        )
        fields["date_formation"] = st.text_input("Date de la formation", "00/00/2026")
        fields["duree"] = st.text_input("Durée", "1 jour (5h30)")
        fields["formateur"] = st.text_input(
            "Formateur(s)", "Major / Infirmier superviseur / Chargé de programme"
        )
        fields["nb_participants"] = st.number_input(
            "Nombre de participants", min_value=1, max_value=50, value=12
        )

    elif doc_type == "note_service":
        st.markdown("### 📋 NOTE DE SERVICE (Conforme Grille ESPC)")

        types_notes = [
            "Désignation du Chef de Centre",
            "Désignation du Responsable de Service",
            "Désignation du Point Focal CMU",
            "Désignation du Responsable de l'Hygiène",
            "Désignation du Responsable de la Gestion des Déchets Biomédicaux",
            "Désignation du Gestionnaire des Médicaments en Pharmacie",
        ]
        fields["type_note"] = st.selectbox("Type de note de service", types_notes)
        fields["nom_etablissement"] = st.text_input("Établissement", "CSR NAGNENEFOUN")
        fields["date_note"] = st.text_input("Date (JJ/MM/AAAA)", "")
        fields["nom_responsable"] = st.text_input("Nom du responsable désigné", "")
        fields["fonction_responsable"] = st.text_input("Fonction du responsable", "")

    return fields


# =============================================================================
# GÉNÉRATEUR DE PLANNING - Cycle PG → R → P
# =============================================================================

PLANNING_CYCLE = ["PG", "R", "P"]

PLANNING_CATEGORIES = [
    {
        "id": "infirmier-dispensaire",
        "label": "Infirmiers (Dispensaire)",
        "prefixe": "Infirmier",
    },
    {"id": "aide-dispensaire", "label": "Aides (Dispensaire)", "prefixe": "Aide"},
    {
        "id": "sage-femme-maternite",
        "label": "Sages-femmes (Maternité)",
        "prefixe": "Sage-femme",
    },
    {"id": "aide-maternite", "label": "Aides (Maternité)", "prefixe": "Aide"},
    {"id": "fille-salle", "label": "Filles de salle", "prefixe": "Fille de salle"},
]

PLANNING_DATA_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "employes.json"
)


def charger_employes():
    try:
        if os.path.exists(PLANNING_DATA_FILE):
            with open(PLANNING_DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


def sauvegarder_employes(employes):
    with open(PLANNING_DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(employes, f, ensure_ascii=False, indent=2)


def generer_planning_employe(cycle_position, annee, mois):
    jours_dans_mois = calendar.monthrange(annee, mois + 1)[1]
    planning = []
    position = cycle_position
    for jour in range(1, jours_dans_mois + 1):
        planning.append({"jour": f"{jour:02d}", "shift": PLANNING_CYCLE[position]})
        position = (position + 1) % len(PLANNING_CYCLE)
    return planning, position


def set_shading(cell, color_hex):
    from docx.oxml import OxmlElement

    tc = cell._element.get_or_add_tcPr()
    for shd in tc.findall(qn("w:shd")):
        tc.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc.append(shd)


def exporter_planning_word(plannings_data, service_label, mois, annee, centre_sante=""):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    p = doc.add_paragraph("MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA CMU")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("RÉPUBLIQUE DE CÔTE D'IVOIRE - Union – Discipline – Travail")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if centre_sante:
        p = doc.add_paragraph(f"Centre: {centre_sante}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    titre = doc.add_paragraph()
    titre.add_run(f"PLANNING MENSUEL - {mois} {annee}").bold = True
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp = doc.add_paragraph(f"Service: {service_label}")
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    if plannings_data:
        jours = len(plannings_data[0]["planning"])
        table = doc.add_table(rows=len(plannings_data) + 1, cols=jours + 1)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "NOM & PRENOM"
        set_shading(header_cells[0], "3498DB")
        for i in range(jours):
            header_cells[i + 1].text = f"{i + 1:02d}"
            set_shading(header_cells[i + 1], "3498DB")
        for i, emp in enumerate(plannings_data):
            row = table.rows[i + 1].cells
            row[0].text = f"{emp['nom']}"
            for j, shift in enumerate(emp["planning"]):
                row[j + 1].text = shift["shift"]
    return doc


def afficher_page_planning():
    st.title("🏥 GÉNÉRATEUR DE PLANNING")
    st.caption(
        "Cycle automatique: **PG** (Permanence+Garde) → **P** (Permanence) → **R** (Repos)"
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**PG** = Permanence + Garde")
    with col2:
        st.markdown("**P** = Permanence")
    with col3:
        st.markdown("**R** = Repos")
    st.divider()

    employes_fresh = charger_employes()

    # Gestion du personnel
    st.header("👥 Gestion du Personnel")
    # Adapter le nombre de colonnes selon la largeur
    n_cols = min(5, max(2, len(PLANNING_CATEGORIES)))
    cols = st.columns(n_cols)
    for idx, cat in enumerate(PLANNING_CATEGORIES):
        with cols[idx]:
            st.subheader(cat["label"])
            employes_cat = [e for e in employes_fresh if e["service"] == cat["id"]]
            st.caption(f"{len(employes_cat)} employé(s)")
            with st.form(f"plan_form_{cat['id']}"):
                nouveau_nom = st.text_input(
                    "Nom", placeholder="Ex: YEO", key=f"plan_nom_{cat['id']}"
                )
                submitted = st.form_submit_button(
                    "➕ Ajouter", use_container_width=True
                )
                if submitted and nouveau_nom:
                    employes_meme_service = [
                        e for e in employes_fresh if e["service"] == cat["id"]
                    ]
                    nouvelle_pos = (
                        (max([e["cyclePosition"] for e in employes_meme_service]) + 1)
                        % len(PLANNING_CYCLE)
                        if employes_meme_service
                        else 0
                    )
                    nouveau = {
                        "id": len(employes_fresh) + 1,
                        "nom": nouveau_nom.upper(),
                        "prenom": cat["prefixe"],
                        "service": cat["id"],
                        "cyclePosition": nouvelle_pos,
                    }
                    employes_fresh.append(nouveau)
                    sauvegarder_employes(employes_fresh)
                    st.rerun()
            for emp in employes_cat:
                col_a, col_b = st.columns([3, 1])
                with col_a:
                    st.markdown(f"**{emp['nom']}**")
                with col_b:
                    if st.button("×", key=f"plan_del_{emp['id']}_{cat['id']}"):
                        employes_fresh = [
                            e for e in employes_fresh if e["id"] != emp["id"]
                        ]
                        sauvegarder_employes(employes_fresh)
                        st.rerun()
    st.divider()

    # Générateur
    st.header("📅 Générer le Planning")
    centre_sante = st.text_input(
        "🏥 Nom du Centre", placeholder="Ex: CSR NAGNENEFOUN", value="CSR NAGNENEFOUN"
    )
    col1, col2, col3 = st.columns(3)
    with col1:
        mois = st.selectbox(
            "Mois",
            [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ],
            index=datetime.now().month - 1,
        )
    with col2:
        annee = st.number_input(
            "Année", min_value=2020, max_value=2030, value=datetime.now().year
        )
    with col3:
        generer_tous = st.checkbox("Tous les services", value=False)
        service = (
            None
            if generer_tous
            else st.selectbox(
                "Service", PLANNING_CATEGORIES, format_func=lambda x: x["label"]
            )
        )

    if st.button("🔄 Générer le Planning", type="primary", use_container_width=True):
        employes_service = (
            employes_fresh
            if service is None
            else [e for e in employes_fresh if e["service"] == service["id"]]
        )
        if not employes_service:
            st.warning("Aucun employé dans ce service !")
        else:
            mois_num = [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ].index(mois)
            employes_fresh = charger_employes()
            employes_service = (
                employes_fresh
                if service is None
                else [e for e in employes_fresh if e["service"] == service["id"]]
            )
            plannings = []
            for emp in employes_service:
                planning, _ = generer_planning_employe(
                    emp["cyclePosition"], annee, mois_num
                )
                plannings.append(
                    {"nom": emp["nom"], "prenom": emp["prenom"], "planning": planning}
                )
            st.success(f"✅ Planning généré pour {len(plannings)} employé(s)")
            jours = len(plannings[0]["planning"])

            # Afficher le tableau avec Streamlit natif
            st.subheader("📅 Planning du mois")
            headers = ["Nom & Prénom"] + [f"{i + 1:02d}" for i in range(jours)]
            data = []
            for emp in plannings:
                row = [emp["nom"]]
                for shift in emp["planning"]:
                    row.append(shift["shift"])
                data.append(row)
            st.dataframe(pd.DataFrame(data, columns=headers), use_container_width=True)

            service_label = service["label"] if service else "Tous les services"
            doc = exporter_planning_word(
                plannings, service_label, mois, annee, centre_sante
            )
            temp_file = (
                f"planning_{service['id'] if service else 'tous'}_{mois}_{annee}.docx"
            )
            doc.save(temp_file)
            with open(temp_file, "rb") as f:
                st.download_button(
                    "📄 Exporter en Word",
                    f.read(),
                    file_name=temp_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


# GESTION DU PERSONNEL - MODULE SIMPLIFIÉ
# =============================================================================


def afficher_page_personnel():
    """Page de gestion du personnel CSR NAGNENEFOUN"""
    import personnel_db as pdb

    st.title("👥 GESTION DU PERSONNEL")
    st.caption("Enregistrez le personnel et attribuez leurs responsabilités")
    st.divider()

    # Initialiser la base de données
    pdb.init_personnel_db()

    # Onglets: Inscription / Liste / Responsabilités
    onglet1, onglet2, onglet3 = st.tabs(
        ["📝 Inscription", "📋 Liste Personnel", "🎯 Responsabilités"]
    )

    # =========================================================================
    # ONGLET 1: FORMULAIRE D'INSCRIPTION SIMPLIFIÉ
    # =========================================================================
    with onglet1:
        st.markdown("### ➕ Nouvelle inscription")
        st.markdown(
            "**Formulaire simplifié** - Remplissez les informations du membre du personnel"
        )

        with st.form("form_inscription"):
            col1, col2 = st.columns(2)
            with col1:
                nom = st.text_input("Nom *", placeholder="Ex: KONÉ")
            with col2:
                prenoms = st.text_input("Prénoms", placeholder="Ex: Awa")

            fonction = st.selectbox("Fonction *", pdb.get_fonctions_disponibles())
            telephone = st.text_input("Téléphone", placeholder="Ex: 0700000000")

            submitted = st.form_submit_button(
                "💾 Enregistrer", use_container_width=True
            )

            if submitted and nom:
                try:
                    personnel_id = pdb.ajouter_personnel(
                        nom=nom.upper(),
                        fonction=fonction,
                        prenoms=prenoms,
                        telephone=telephone,
                    )
                    st.success(f"✅ {nom.upper()} enregistré(e) avec succès!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Erreur: {str(e)}")
            elif submitted and not nom:
                st.warning("⚠️ Le nom est obligatoire")

    # =========================================================================
    # ONGLET 2: LISTE DU PERSONNEL
    # =========================================================================
    with onglet2:
        st.markdown("### 📋 Liste du Personnel")

        tout_personnel = pdb.get_tout_personnel()

        if tout_personnel:
            st.info(f"📊 Total: {len(tout_personnel)} membre(s)")

            # Filtre par fonction
            fonctions = ["Tous"] + pdb.get_fonctions_disponibles()
            filtre = st.selectbox("Filtrer par fonction", fonctions)

            if filtre != "Tous":
                tout_personnel = [p for p in tout_personnel if p["fonction"] == filtre]

            # Afficher la liste
            for p in tout_personnel:
                with st.container():
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.markdown(f"**{p['nom']}** {p['prenom'] or ''}")
                        st.caption(f"📞 {p['telephone'] or 'Non renseigné'}")
                    with col2:
                        st.markdown(f"🏥 {p['fonction']}")
                        st.caption(f"📅 Depuis: {p['date_ajout']}")
                    with col3:
                        if st.button("🗑️", key=f"del_{p['id']}"):
                            pdb.supprimer_personnel(p["id"])
                            st.rerun()
                    st.markdown("---")
        else:
            st.info(
                "Aucun personnel enregistré. Utilisez le formulaire ci-dessus pour commencer."
            )

    # =========================================================================
    # ONGLET 3: ATTRIBUTION DES RESPONSABILITÉS
    # =========================================================================
    with onglet3:
        st.markdown("### 🎯 Attribuer une responsabilité")
        st.markdown(
            "Sélectionnez un membre du personnel et attribuez-lui une responsabilité"
        )

        tout_personnel = pdb.get_tout_personnel()

        if tout_personnel:
            # Sélection du personnel
            options = [
                f"{p['nom']} {p['prenom'] or ''} - {p['fonction']}"
                for p in tout_personnel
            ]
            selected = st.selectbox("Membre du personnel *", options)
            selected_index = options.index(selected)
            selected_personnel = tout_personnel[selected_index]

            # Afficher les responsabilités actuelles
            resp_actuelles = pdb.get_responsabilites_personnel(selected_personnel["id"])
            if resp_actuelles:
                st.markdown("**Responsabilités actuelles:**")
                for r in resp_actuelles:
                    st.markdown(
                        f"- ✅ {r['responsabilite']} (depuis {r['date_affectation']})"
                    )
            else:
                st.info("Ce membre n'a pas encore de responsabilité attribuée.")

            st.markdown("---")

            # Nouvelle responsabilité
            st.markdown("**Nouvelle responsabilité:**")
            responsabilite = st.selectbox(
                "Responsabilité *", pdb.get_responsabilites_types()
            )
            notes = st.text_area(
                "Notes (optionnel)", placeholder="Ex: Prend effet immédiatement"
            )

            if st.button("➕ Ajouter la responsabilité", use_container_width=True):
                try:
                    pdb.ajouter_responsabilite(
                        personnel_id=selected_personnel["id"],
                        responsabilite=responsabilite,
                        notes=notes,
                    )
                    st.success(
                        f"✅ Responsabilité attribuée à {selected_personnel['nom']}!"
                    )
                    st.rerun()
                except Exception as e:
                    st.error(f"Erreur: {str(e)}")
        else:
            st.info(
                "Aucun personnel enregistré. Ajoutez d'abord du personnel dans l'onglet 'Inscription'."
            )

    st.divider()
    st.caption("**CSR NAGNENEFOUN** — District Sanitaire de KORHOGO 1 — Région du PORO")


# =============================================================================
# PAGES (fonctions séparées pour st.navigation)
# =============================================================================


def page_accueil():
    st.title("🏥 Générateur Documents ESPC")
    st.caption(
        "Conforme à la Grille d'Évaluation des Établissements Sanitaires de Premier Contact"
    )

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        ### 📄 17 documents prêts à l'emploi
        - PV Réunions (Mensuelle, COGES, AG)
        - Rapports (Supervision ASC, Plaintes, Formation)
        - Fiches (Poste, Nomination)
        - Calendriers (Nettoyage, Réunions)
        - Plans (Action, Supervision ASC)
        - Grilles, Listes et Notes de Service

        **Tous conformes à la grille ESPC**
        """)

    st.markdown("---")
    st.caption(
        "⚙️ Personnalisez vos modèles | 📥 Export Word | 👁️ Aperçu avant téléchargement"
    )
    st.caption("**CSR NAGNENEFOUN** — District Sanitaire de KORHOGO 1 — Région du PORO")


# =============================================================================
# INTERFACE PRINCIPALE
# =============================================================================


def main():
    sidebar_chat_widget()

    pg = st.navigation(
        {
            "🏠 Menu": [
                st.Page(page_accueil, title="Accueil", icon="🏠", default=True),
                st.Page(page_generateur, title="Générateur de documents", icon="📄"),
                st.Page(page_templates_rapides, title="Templates Rapides", icon="⚡"),
                st.Page(page_guide, title="Guide Cahiers / Registres", icon="📋"),
            ],
            "💬 Assistant": [
                st.Page(page_chatbot, title="Assistant ESPC", icon="💬"),
            ],
            "👥 Personnel": [
                st.Page(afficher_page_planning, title="Planning Personnel", icon="🏥"),
                st.Page(afficher_page_personnel, title="Gestion Personnel", icon="👥"),
            ],
        }
    )
    pg.run()


def page_templates_rapides():
    st.title("⚡ TEMPLATES RAPIDES")
    st.caption(
        "Documents pré-remplis CSR NAGNENEFOUN — Imprimez et complétez le nom + la date manuellement"
    )
    st.markdown("---")

    rapides = [
        (
            "👨‍⚕️",
            "Fiche de Poste - Chef de Centre",
            "Chef de Centre CSR - Missions, qualifications, responsabilités",
            """RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
MINISTÈRE DE LA SANTÉ, HYGIÈNE PUBLIQUE ET CMU
RÉGION SANITAIRE DU PORO - DISTRICT KORHOGO 1
CSR NAGNENEFOUN

FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION DU POSTE
Établissement: CSR NAGNENEFOUN
Poste: Chef de Centre de Santé
Supérieur: Médecin Chef District KORHOGO 1
Régime: Temps plein - 40h/semaine
Nom du titulaire: [À COMPLÉTER]
Date: ____/____/______

II. MISSIONS PRINCIPALES
1. Coordonner les activités cliniques, administratives et financières du CSR
2. Superviser l'équipe du personnel (IDE, SF, Aide-soignants, administratif)
3. Garantir la qualité des soins et le respect des protocoles nationaux
4. Élaborer et exécuter le plan d'action annuel
5. Gérer le budget et les ressources du centre
6. Présider les réunions de staff mensuelles et COGES trimestrielles
7. Assurer le suivi des indicateurs et la transmission des rapports au District
8. Représenter le centre auprès des autorités et partenaires

III. QUALIFICATIONS
- IDE ou Sage-Femme Diplômé d'État
- Expérience minimale de 3 ans en centre de santé
- Inscription à l'Ordre national

IV. SIGNATURE
Cachet et Signature du Chef de Centre:
Date: ____/____/______""",
        ),
        (
            "👩‍⚕️",
            "Fiche de Poste - Major / Sage-Femme",
            "Major/SF - Maternité, CPN, accouchements",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Major / Sage-Femme Diplômé(e) d'État
Supérieur: Chef de Centre
Régime: Temps plein
Nom: [À COMPLÉTER]

II. MISSIONS
1. Superviser les soins obstétricaux et néonatals
2. Assurer les CPN, CPoN et accouchements
3. Coordonner la PF et PTME
4. Gérer les urgences obstétricales
5. Tenir les registres de la maternité

III. QUALIFICATIONS
- DESF (Diplôme d'État de Sage-Femme)
- Inscription Ordre national
- Expérience 2 ans minimum

Signature: ____/____/______""",
        ),
        (
            "🩺",
            "Fiche de Poste - IDE",
            "Infirmier - Consultations, soins, PEV",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Infirmier Diplômé d'État (IDE)
Supérieur: Chef de Centre / Major
Régime: Temps plein
Nom: [À COMPLÉTER]

II. MISSIONS
1. Consultations curatives (paludisme, IRA, diarrhées)
2. Administration des médicaments et soins infirmiers
3. PEV et stratégies avancées
4. Prise en charge paludisme (TDR, ACT)
5. Dépistage VIH et lien PTME
6. Tenue des registres et rapports SIG

III. QUALIFICATIONS
- DEI (Diplôme d'État d'Infirmier)
- Inscription Ordre national

Signature: ____/____/______""",
        ),
        (
            "🩹",
            "Fiche de Poste - Aide-Soignant(e)",
            "Aide-Soignant - Assistance, hygiène",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Aide-Soignant(e)
Supérieur: Major / IDE de garde
Régime: Temps plein
Nom: [À COMPLÉTER]

II. MISSIONS
1. Assister l'infirmier dans les soins courants
2. Préparer et entretenir le matériel de soins
3. Assurer l'hygiène des salles de soins
4. Accueillir et orienter les patients

III. QUALIFICATIONS
- Certificat d'Aide-Soignant(e)
- Expérience 1 an minimum

Signature: ____/____/______""",
        ),
        (
            "📋",
            "Fiche de Poste - Agent Administratif",
            "Secrétaire - Gestion administrative",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Secrétaire / Agent Administratif
Supérieur: Chef de Centre
Régime: Temps plein
Nom: [À COMPLÉTER]

II. MISSIONS
1. Gestion administrative et secrétariat
2. Tenue des registres administratifs
3. Accueil et orientation des usagers
4. Saisie des rapports et documents

III. QUALIFICATIONS
- BTS Administration / Diplôme de Secrétariat
- Maîtrise bureautique

Signature: ____/____/______""",
        ),
        (
            "🧹",
            "Fiche de Poste - Agent d'Entretien",
            "Agent d'entretien - Nettoyage, hygiène",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Agent d'Entretien / Fille de Salle
Supérieur: Major / IDE de garde
Régime: Temps plein
Nom: [À COMPLÉTER]

II. MISSIONS
1. Nettoyage et désinfection des locaux
2. Gestion des déchets biomédicaux
3. Entretien du linge
4. Propreté des sanitaires

Signature: ____/____/______""",
        ),
        (
            "🔐",
            "Fiche de Poste - Gardien",
            "Gardien/Planton - Sécurité",
            """FICHE DE POSTE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

I. IDENTIFICATION
Poste: Gardien / Planton
Supérieur: Chef de Centre
Régime: Garde 24h (roulement)
Nom: [À COMPLÉTER]

II. MISSIONS
1. Assurer la sécurité du centre
2. Contrôler les entrées/sorties
3. Ouvrir et fermer le centre selon les horaires

Signature: ____/____/______""",
        ),
        (
            "📝",
            "Note de Service - Chef de Centre",
            "Désignation Chef de Centre (Norme 1.01 d)",
            """RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
MSHP-CMU / RÉGION PORO / DISTRICT KORHOGO 1
CSR NAGNENEFOUN

NOTE DE SERVICE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

OBJET: Désignation du Chef de Centre

Le Médecin Chef du District de KORHOGO 1,

VU les textes organiques;
VU les besoins de service;

DÉSIGNE

Nom: [À COMPLÉTER]
Fonction: Chef de Centre - CSR NAGNENEFOUN
Date d'effet: ____/____/______

Missions: Coordination, supervision, gestion, qualité des soins, rapports District.

Fait à KORHOGO 1, le ____/____/______
Le Médecin Chef du District

[Cachet et Signature]

Diffusion: Intéressé(e), District, Archives""",
        ),
        (
            "📝",
            "Note de Service - Point Focal CMU",
            "Point Focal CMU (Norme 4.03 c)",
            """NOTE DE SERVICE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

OBJET: Désignation Point Focal CMU

Le Chef du CSR NAGNENEFOUN,

DÉSIGNE

Nom: [À COMPLÉTER]
Fonction: Point Focal CMU
Date d'effet: ____/____/______

Missions: Accueil CMU, enregistrement, complétude registre, transmission SurveyCTO.

Fait à CSR NAGNENEFOUN, le ____/____/______
Le Chef de Centre

[Cachet et Signature]""",
        ),
        (
            "📝",
            "Note de Service - Responsable Hygiène",
            "Responsable Hygiène (Norme 6.01 a)",
            """NOTE DE SERVICE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

OBJET: Désignation Responsable Hygiène Hospitalière

Le Chef du CSR NAGNENEFOUN,

DÉSIGNE

Nom: [À COMPLÉTER]
Fonction: Responsable Hygiène
Date d'effet: ____/____/______

Missions: Supervision hygiène, déchets biomédicaux, formation personnel.

Fait à CSR NAGNENEFOUN, le ____/____/______
Le Chef de Centre

[Cachet et Signature]""",
        ),
        (
            "📝",
            "Note de Service - Gestionnaire Pharmacie",
            "Gestionnaire Médicaments (Norme 11.01 g)",
            """NOTE DE SERVICE N° ____/MS/RS-PORO/DS-K1/CSR NAGNENEFOUN

OBJET: Désignation Gestionnaire des Médicaments

Le Chef du CSR NAGNENEFOUN,

DÉSIGNE

Nom: [À COMPLÉTER]
Fonction: Gestionnaire Pharmacie
Date d'effet: ____/____/______

Missions: Gestion stocks, dispensation, commandes, contrôle périmés.

Fait à CSR NAGNENEFOUN, le ____/____/______
Le Chef de Centre

[Cachet et Signature]""",
        ),
        (
            "📜",
            "Arrêté - Mise en place COGES",
            "Arrêté COGES (Norme 1.02 a)",
            """RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
RÉGION PORO / DISTRICT KORHOGO 1

ARRÊTÉ N° ____/MS/RS-PORO/DS-K1/PORTANT MISE EN PLACE COGES

Le Préfet/Sous-Préfet de [À COMPLÉTER],

Considérant les textes organiques;
Considérant la grille ESPC (Norme 1.02 a);

ARRÊTE

Art.1: Le COGES du CSR NAGNENEFOUN est mis en place.
Président: [À COMPLÉTER]
Vice-Président: [À COMPLÉTER]
Secrétaire: [À COMPLÉTER]
Trésorier: [À COMPLÉTER]
Commissaire: [À COMPLÉTER]

Art.2: Le COGES assure la gestion participative, mobilisation des ressources, suivi des activités.

Art.3: Mandat de 3 ans renouvelable.

Fait à [Lieu], le ____/____/______
Le Préfet/Sous-Préfet

[Cachet et Signature]""",
        ),
        (
            "📜",
            "Arrêté - Liste Nominative ASC",
            "Liste officielle ASC (Norme 14.01 a)",
            """RÉPUBLIQUE DE CÔTE D'IVOIRE
Union – Discipline – Travail
RÉGION PORO / DISTRICT KORHOGO 1

ARRÊTÉ N° ____/MS/RS-PORO/DS-K1/PORTANT LISTE ASC

Le Médecin Chef du District de KORHOGO 1,

ARRÊTE

Liste des ASC rattachés au CSR NAGNENEFOUN:

N° | Nom & Prénoms | Village | Téléphone
1. | [À COMPLÉTER] | [À COMPLÉTER] | [À COMPLÉTER]
2. | [À COMPLÉTER] | [À COMPLÉTER] | [À COMPLÉTER]
3. | [À COMPLÉTER] | [À COMPLÉTER] | [À COMPLÉTER]

Fait à KORHOGO 1, le ____/____/______
Le Médecin Chef du District

[Cachet et Signature]""",
        ),
    ]

    for icone, titre, desc, contenu in rapides:
        with st.expander(f"{icone} **{titre}** — {desc}", expanded=False):
            st.text_area(
                "Contenu du template", contenu, height=200, key=f"rapide_{titre}"
            )
            doc = creer_document_word(titre, contenu)
            from io import BytesIO

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                f"📥 Télécharger {titre} (.docx)",
                buf,
                f"{titre}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


def page_guide():
    st.title("📋 Guide Stratégique ESPC")
    st.caption(
        "Tout ce que le Chef de Centre doit préparer — Conforme à la Grille 1600 pts"
    )
    st.divider()

    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "👤 Désignations",
        "📝 Notes de Service",
        "📒 Cahiers/Registres",
        "📊 Rapports",
        "🖼️ Affichage",
        "📅 Planification",
    ])

    # =========================================================================
    # TAB 1: DÉSIGNATIONS
    # =========================================================================
    with tab1:
        st.markdown("### 👤 Désignations et Nominations")
        st.info("Chaque désignation nécessite une **Note de Service** ou un **Arrêté**. Générables via l'app.")

        nominations = [
            ("1.01 d", "Chef de Centre", "Note de service", "DD KORHOGO 1"),
            ("1.01 d", "Responsable de chaque service", "Note de service", "Responsable ESPC"),
            ("1.02 a", "Membres du COGES (7 membres)", "Arrêté préfectoral", "Préfet/Sous-Préfet"),
            ("4.03 c", "Point Focal CMU", "Note de service", "Responsable ESPC"),
            ("4.02 b", "Agent d'accueil et orientation", "Fiche de poste signée", "Responsable ESPC"),
            ("6.01 a", "Responsable Hygiène Hospitalière", "Note de service", "Responsable ESPC"),
            ("6.05 a", "Responsable Déchets Biomédicaux", "Note de service", "Responsable ESPC"),
            ("11.01 g", "Gestionnaire Pharmacie", "Note + fiche de poste", "Responsable ESPC"),
            ("7.02 a", "Personnel qualifié accouchements", "Fiche de poste signée", "Responsable ESPC"),
            ("2.01 b", "Tout agent (fiche de poste)", "Fiche de poste signée", "Responsable ESPC"),
            ("14.01 a", "Liste ASC nominative", "Liste officielle", "District sanitaire"),
            ("14.01 c", "Grille supervision ASC", "Grille signée", "ASC + Superviseur"),
        ]

        for n in nominations:
            st.markdown(f"**`{n[0]}`** {n[1]}")
            st.caption(f"📄 {n[2]} — ✍️ {n[3]}")

        st.success(f"**{len(nominations)} désignations** — Utilisez l'onglet Fiche de Nomination ou Templates Rapides")

    # =========================================================================
    # TAB 2: NOTES DE SERVICE
    # =========================================================================
    with tab2:
        st.markdown("### 📝 Notes de Service à rédiger")
        st.info("Format officiel: **RÉPUBLIQUE DE CÔTE D'IVOIRE → MSHP → Région → District → CSR**")

        notes = [
            ("1.01 d", "Désignation du Chef de Centre", "DD KORHOGO 1"),
            ("1.01 d", "Désignation Responsables de service", "Responsable ESPC"),
            ("4.03 c", "Désignation Point Focal CMU", "Responsable ESPC"),
            ("6.01 a", "Désignation Responsable Hygiène", "Responsable ESPC"),
            ("6.05 a", "Désignation Responsable Déchets", "Responsable ESPC"),
            ("11.01 g", "Désignation Gestionnaire Pharmacie", "Responsable ESPC"),
            ("1.03 c", "Diffusion PV réunion aux chefs de service", "Responsable ESPC"),
            ("2.01 c", "Affichage planning mensuel", "Responsable ESPC"),
            ("2.01 d", "Affichage programme de gardes", "Responsable ESPC"),
            ("6.01 b", "Procédures d'hygiène", "Responsable Hygiène"),
        ]

        for n in notes:
            st.markdown(f"- **`{n[0]}`** {n[1]} → ✍️ {n[2]}")

        st.success(f"**{len(notes)} notes de service** — Générables via Générateur > Note de Service")

    # =========================================================================
    # TAB 3: CAHIERS/REGISTRES
    # =========================================================================
    with tab3:
        st.markdown("### 📒 Cahiers et Registres à tenir à jour")
        st.info("Le contrôleur vérifie: **existence**, **conformité** et **tenue à jour**")

        cahiers = [
            ("Gestion", [
                ("2.01 e", "Cahier de présence", "Date, nom, heure, signature, pas de saut"),
            ]),
            ("Finance", [
                ("3.01 a", "Journal de caisse (brouillard)", "Tenue à jour"),
                ("3.01 b", "Rapport financier trimestriel", "Toutes les ressources"),
                ("3.01 c", "États de redevances mensuels", "3 derniers mois"),
                ("3.01 d", "Ordres de paiement et liasses", "Factures, PV, bons"),
                ("11.02 a.iv", "Cahiers recettes journalières", "Tenue et archivage"),
            ]),
            ("Pharmacie", [
                ("11.02 a.ii", "Fiche de stock médicaments", "Concordance, ruptures"),
                ("11.02 a.iii", "Cahier inventaire pharmacie", "Inventaire régulier"),
                ("11.02 a.v", "Ordonnancier / Facture", "Disponibilité"),
            ]),
            ("Maternité", [
                ("7.02 a", "Registre d'accouchement", "Partogramme, médicaments, GATPA"),
                ("7.08 b", "Registre CPoN", "Tous les items complétés"),
                ("8.01 a", "Fiche notification décès maternel", "5 fiches vierges disponibles"),
            ]),
            ("Consultation", [
                ("7.03 e", "Registre consultation curative", "Paludisme, IRA, diarrhée"),
            ]),
            ("CMU", [
                ("4.03 g", "Registre prise en charge CMU", "Complétude, surveyCTO"),
                ("4.03 h", "Cahier transmission CMU", "Bordereaux déchargés"),
            ]),
            ("SIG", [
                ("1.04 d", "Rapport SIG mensuel", "Cohérence, corrections"),
                ("1.04 d", "Matrice de cohérence", "Triangulation SIG"),
            ]),
            ("Communautaire", [
                ("14.01 f", "Rapport mensuel communautaire ASC", "Par mois, concordance"),
                ("15.01 a", "Fiche stock ASC", "À jour"),
                ("15.01 b", "Rapport activités communautaires", "Par ASC"),
            ]),
            ("Hygiène", [
                ("6.06 b", "Cahier inventaire produits entretien", "3 derniers mois"),
            ]),
        ]

        for service, items in cahiers:
            st.markdown(f"#### {service}")
            for item in items:
                st.markdown(f"- **`{item[0]}`** {item[1]} — _{item[2]}_")

        total = sum(len(items) for _, items in cahiers)
        st.success(f"**{total} cahiers/registres** à vérifier")

    # =========================================================================
    # TAB 4: RAPPORTS
    # =========================================================================
    with tab4:
        st.markdown("### 📊 Rapports à produire")
        st.info("Ces rapports doivent être produits régulièrement et archivés")

        rapports = [
            ("1.01", "PAA budgétisé", "Annuel", "Chef de Centre"),
            ("1.01 g", "Revue trimestrielle PAA", "Trimestriel", "Chef de Centre"),
            ("1.02 b", "PV Réunion COGES", "Trimestriel", "Secrétaire COGES"),
            ("1.03", "PV Assemblée Générale", "Annuel", "Président COGES"),
            ("1.04 b", "Rapport Formation Personnel", "Trimestriel", "Major"),
            ("1.04 d", "Rapport SIG Mensuel", "Mensuel", "Point Focal SIG"),
            ("4.02", "Rapport Plaintes/Suggestions", "Trimestriel", "Agent Accueil"),
            ("6.05", "Rapport Infections Nosocomiales", "Trimestriel", "Resp. Hygiène"),
            ("8.01 b", "Rapport Communautaire ASC", "Mensuel", "Chargé Programme"),
            ("14.01 d", "Rapport Supervision ASC", "Mensuel", "Superviseur ASC"),
            ("15.01 b", "Rapport Activités Communautaires", "Mensuel", "Chargé Programme"),
        ]

        for r in rapports:
            st.markdown(f"**`{r[0]}`** {r[1]}")
            st.caption(f"📅 {r[2]} — ✍️ {r[3]}")

        st.success(f"**{len(rapports)} rapports** — La plupart générables via l'onglet Générateur (IA)")

    # =========================================================================
    # TAB 5: AFFICHAGE
    # =========================================================================
    with tab5:
        st.markdown("### 🖼️ Documents à afficher")
        st.info("Ces documents doivent être **visibles** dans le centre lors de l'évaluation")

        affichage = [
            ("1.01", "Plan d'Action Annuel (PAA)", "Tableau d'affichage"),
            ("1.02", "Arrêté COGES", "Tableau d'affichage"),
            ("1.02 b", "PV réunion COGES trimestre précédent", "Tableau d'affichage"),
            ("1.03", "PV réunion mensuelle personnel", "Tableau d'affichage"),
            ("2.01 c", "Planning mensuel de travail", "Tableau d'affichage"),
            ("2.01 d", "Programme gardes et astreintes", "Tableau d'affichage"),
            ("2.01 e", "Cahier de présence", "Point d'entrée"),
            ("4.03", "Liste tarifs et prestations CMU", "Accueil / Hall"),
            ("6.01", "Procédures d'hygiène", "Chaque salle de soins"),
            ("6.05", "Plan gestion déchets biomédicaux", "Zone de tri"),
            ("7.04", "Protocoles SONU", "Salle d'accouchement"),
            ("14.01", "Liste nominative ASC", "Bureau Chef Centre"),
        ]

        for a in affichage:
            st.markdown(f"- **`{a[0]}`** {a[1]} → 📍 {a[2]}")

        st.success(f"**{len(affichage)} documents** à afficher")

    # =========================================================================
    # TAB 6: PLANIFICATION
    # =========================================================================
    with tab6:
        st.markdown("### 📅 Planification annuelle")
        st.info("Documents de planification stratégique à produire")

        planification = [
            ("1.01", "Plan d'Action Annuel (PAA) budgétisé", "Annuel"),
            ("1.01 g", "Revue trimestrielle du PAA (×4)", "Trimestriel"),
            ("1.04 b", "Plan de Formation Continue", "Annuel"),
            ("6.01", "Calendrier de Nettoyage du centre", "Annuel"),
            ("14.01 b", "Plan de supervision annuel ASC", "Annuel"),
        ]

        for p in planification:
            st.markdown(f"- **`{p[0]}`** {p[1]} ({p[2]})")

        st.success(f"**{len(planification)} documents** de planification — Générables via l'IA")

    # =========================================================================
    # RÉCAPITULATIF
    # =========================================================================
    st.divider()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("👤 Désignations", "12")
    c2.metric("📝 Notes de Service", "10")
    c3.metric("📒 Cahiers/Registres", "25")
    c4.metric("📊 Rapports", "11")

    st.divider()
    st.caption("**CSR NAGNENEFOUN** — Guide stratégique ESPC — Conforme à la Grille 1600 pts")


def page_generateur():
    st.title("🏥 Générateur Documents ESPC")
    st.markdown("**Conforme à la Grille ESPC**")

    # =============================================================================
    # SECTION PERSONNALISATION DES TEMPLATES
    # =============================================================================
    with st.expander("⚙️ Personnaliser les templates"):
        st.markdown("### Modifier la structure des documents")

        # Choisir quel document modifier
        template_options = list(templates.keys()) if templates else []
        template_noms = (
            {k: templates[k]["nom"] for k in template_options} if templates else {}
        )
        template_choice = st.selectbox(
            "Choisir le document à modifier",
            template_options,
            format_func=lambda x: template_noms.get(x, x),
        )

        if template_choice and templates:
            st.markdown(f"#### 📄 {templates[template_choice]['nom']}")

            # Afficher les sections actuelles
            sections_actuelles = templates[template_choice]["sections"]

            # Modifier les sections
            sections_text = st.text_area(
                "Sections (une par ligne)",
                value="\n".join(sections_actuelles),
                height=150,
            )

            # Convertir en liste
            nouvelles_sections = [
                s.strip() for s in sections_text.split("\n") if s.strip()
            ]

            # Bouton pour sauvegarder
            if st.button("💾 Sauvegarder les modifications"):
                templates[template_choice]["sections"] = nouvelles_sections
                sauvegarder_templates(templates)
                st.success(
                    f"✅ Template '{templates[template_choice]['nom']}' mis à jour!"
                )
                st.rerun()

            # Bouton pour réinitialiser
            if st.button("↩️ Réinitialiser"):
                # Recréer le fichier original
                sauvegarder_templates(templates)
                st.success("Template réinitialisé!")
                st.rerun()

    st.markdown("---")

    doc_options = [d[0] for d in DOCUMENTS_LIST]
    type_doc = st.selectbox("📄 Choisir le document", doc_options)

    doc_key = None
    for name, key in DOCUMENTS_LIST:
        if name == type_doc:
            doc_key = key
            break

    st.markdown("---")

    # Sélecteur de thème principal (uniquement pour les PV et rapports d'activités)
    docs_avec_themes = [
        "pv_reunion_mensuelle",
        "pv_coges",
        "pv_ag",
        "rapport_supervision_asc",
        "rapport_plaintes",
        "rapport_formation",
    ]

    theme_principal = ""
    if doc_key in docs_avec_themes:
        st.markdown("### 🎯 Thème principal")
        themes_disponibles = [
            "Santé maternelle (CPN, accouchements, PF, PTME)",
            "Santé infantile (PEV, croissance, malnutrition)",
            "Paludisme (dépistage, traitement, prévention)",
            "Hygiène et infection",
            "Gouvernance (réunions, COGES)",
            "Surveillance épidémiologique",
            "Pharmacie et médicaments",
            "Nutrition",
            "IEC/CCC (sensibilisation)",
            "Activités communautaires (ASC)",
            "Gestion des équipements",
            "Rapports et données",
        ]
        theme_principal = st.selectbox(
            "Choisir le thème principal du document", themes_disponibles
        )

    st.markdown("---")

    if doc_key:
        donnees = get_form_fields(doc_key)
        # Ajouter le contexte du CSR aux données
        donnees["contexte"] = get_contexte_csr()
        # Ajouter les thèmes sélectionnés
        if theme_principal:
            donnees["themes"] = f"- {theme_principal}"
        else:
            donnees["themes"] = ""

        # Dossier de sauvegarde (répertoire de l'application)
        dossier_sortie = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "documents_generes"
        )
        if not os.path.exists(dossier_sortie):
            os.makedirs(dossier_sortie)

    st.markdown("---")

    if st.button("🚀 Générer le document", type="primary"):
        with st.spinner("Génération en cours..."):
            # Cas spécial: Liste Personnel Centre (statique, sans IA)
            if doc_key == "liste_personnel_centre":
                personnel_data = [
                    (
                        "1",
                        "Kouassi Yao",
                        "Chef de Centre",
                        "0102030401",
                        "Infirmerie",
                    ),
                    (
                        "2",
                        "Koné Abibata",
                        "Major / Sage-femme",
                        "0102030402",
                        "Maternité",
                    ),
                    ("3", "Touré Fatoumata", "Chargée PEV", "0102030403", "PEV"),
                    (
                        "4",
                        "Koffi Aka",
                        "Chargé Paludisme",
                        "0102030404",
                        "Consultation",
                    ),
                    (
                        "5",
                        "N'Guessan Kouamé",
                        "Chargé VIH/PTME",
                        "0102030405",
                        "VIH",
                    ),
                    (
                        "6",
                        "Kouakou Akissi",
                        "Chargée CPN",
                        "0102030406",
                        "Maternité",
                    ),
                    (
                        "7",
                        "Bamba Sékou",
                        "Chargé Pharmacie",
                        "0102030407",
                        "Pharmacie",
                    ),
                    (
                        "8",
                        "Kra Adjo",
                        "Agent d'entretien",
                        "0102030408",
                        "Nettoyage",
                    ),
                    ("9", "Dibi Franck", "Planton", "0102030409", "Accueil"),
                    (
                        "10",
                        "Kouamé Bertine",
                        "Secrétaire",
                        "0102030410",
                        "Secrétariat",
                    ),
                    ("11", "Gnahoua Olivier", "Gardien", "0102030411", "Sécurité"),
                    ("12", "Konan Blanche", "ASC", "0102030412", "Communautaire"),
                ]

                contenu = f"""LISTE DU PERSONNEL DU CSR NAGNENEFOUN

    I. EN-TÊTE OFFICIEL
    RÉPUBLIQUE DE CÔTE D'IVOIRE
    Union – Discipline – Travail
    MINISTÈRE DE LA SANTÉ
    RÉGION SANITAIRE DU PORO
    DISTRICT SANITAIRE DE KORHOGO 1
    {donnees.get("nom_etablissement", "CSR NAGNENEFOUN")}

    II. LISTE DU PERSONNEL (TABLEAU)
    N° | Nom & Prénoms | Fonction | Contact | Service"""

                for row in personnel_data:
                    contenu += f"\n| {' | '.join(row)} |"

                contenu += f"""

    III. RÉCAPITULATIF
    - Effectif total: 12 agents
    - Infirmiers: 4
    - Sages-femmes: 2
    - Personnel d'appui: 4
    - ASC: 1
    - Secrétaire: 1

    IV. OBSERVATIONS
    - Cette liste est établie pour l'année {donnees.get("periode", "2026")}
    - Tout changement de personnel doit être signalé au District Sanitaire de KORHOGO 1

    V. APPROBATION
    - Vu par le Chef de Centre:
    - Date: ____/____/{donnees.get("periode", "2026")}
    - Cachet et signature:
    """
                meta = {
                    "Établissement": donnees.get("nom_etablissement", ""),
                    "Période": donnees.get("periode", ""),
                }
                doc = creer_document_word(type_doc, contenu, meta)

                from io import BytesIO

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Sauvegarder dans le dossier
                nom_fichier = (
                    f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx"
                )
                chemin_fichier = os.path.join(dossier_sortie, nom_fichier)
                with open(chemin_fichier, "wb") as f:
                    f.write(buffer.getvalue())

                st.success(f"✅ Document généré! Sauvegardé dans: documents_generes/")

                st.download_button(
                    "📥 Télécharger",
                    buffer,
                    f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                with st.expander("👁️ Aperçu"):
                    st.text(contenu)

            elif doc_key == "plan_supervision_asc":
                contenu = f"""PLAN DE SUPERVISION DES ASC - CSR NAGNENEFOUN - {donnees.get("periode", "2026")}

    I. CONTEXTE
    Le présent plan de supervision est établi conformément à la grille ESPC pour encadrer et évaluer les Agents de Santé Communautaires (ASC) rattachés au CSR NAGNENEFOUN. Il couvre les activités de supervision sur le terrain et au centre pour l'année {donnees.get("periode", "2026")}.

    Établissement: {donnees.get("nom_etablissement", "CSR NAGNENEFOUN")}
    Période: {donnees.get("periode", "2026")}
    District: KORHOGO 1 (PORO)

    II. OBJECTIFS DE LA SUPERVISION
    - Objectif général: Assurer la qualité des prestations des ASC dans la communauté
    - Objectifs spécifiques:
      1. Évaluer les connaissances et compétences techniques des ASC
      2. Vérifier la qualité des données et des rapports
      3. Renforcer les capacités par des formations continues
      4. Assurer la disponibilité des intrants et médicaments
      5. Améliorer la référence des cas vers le CSR

    III. PLAN DE SUPERVISION ANNUEL (TABLEAU)
    Période | ASC/Village | Thème de la supervision | Activités détaillées | Superviseur
    |---|---|---|---|---
    | Janvier - Février | Tous les ASC | Planification annuelle et évaluation des connaissances | Organiser une réunion de planification avec tous les ASC, évaluer les connaissances sur le paludisme (TDR et traitement), vérifier les registres et rapports de l'année précédente, distribuer les nouveaux supports IEC/CCC, planifier le calendrier des stratégies avancées | Chef de Centre, Major |
    | Mars - Avril | ASC Villages Aire 1 | Supervision terrain paludisme et IEC/CCC | Accompagner l'ASC en stratégie avancée, observer une séance de sensibilisation IEC/CCC sur le paludisme, vérifier la réalisation des TDR et l'application du protocole, contrôler l'état des stocks d'intrants (TDR, ACT, MII), évaluer la qualité du remplissage des registres | Infirmier superviseur, Chargé paludisme |
    | Mai - Juin | ASC Villages Aire 2 | Campagne CPS et PEV communautaire | Superviser la campagne de chimio-prévention du paludisme saisonnier (CPS), vérifier l'administration correcte des doses, observer les séances de vaccination PEV en stratégie avancée, évaluer la gestion de la chaîne de froid, contrôler les formulaires de rapport CPS | IDE, Chargé PEV |
    | Juillet - Août | Tous les ASC | Nutrition et dépistage malnutrition | Organiser une formation sur le dépistage de la malnutrition, superviser les séances de démonstration nutritionnelle, vérifier l'utilisation du périmètre brachial (MUAC), évaluer la référence des cas de malnutrition vers le CSR, contrôler la tenue du registre nutrition | Infirmier, Nutritionniste |
    | Septembre - Octobre | ASC Villages Aire 3 | VIH/PTME et suivi des patients perdus de vue | Vérifier l'orientation des femmes enceintes pour la PTME, évaluer le dépistage VIH dans la communauté, superviser le suivi des patients sous ARV perdus de vue, contrôler la tenue du registre VIH, organiser une séance de sensibilisation sur le VIH | Chargé VIH, Sage-femme |
    | Novembre - Décembre | Tous les ASC | Évaluation annuelle et bilan N+1 | Réaliser l'évaluation annuelle des performances de chaque ASC, compiler les données de supervision de l'année, organiser une réunion de bilan avec tous les ASC, identifier les besoins en formation pour l'année N+1, élaborer le rapport annuel de supervision | Chef de Centre, Major |

    IV. GRILLE D'ÉVALUATION DES ASC
    Les ASC sont évalués selon les critères suivants:
    1. Accueil et relation communautaire - Note A/B/C
    2. Qualité des sensibilisations IEC/CCC - Note A/B/C
    3. Compétence en dépistage et TDR paludisme - Note A/B/C
    4. Gestion des intrants et médicaments - Note A/B/C
    5. Qualité des rapports et registres - Note A/B/C
    6. Taux de référence des cas vers le CSR - Note A/B/C
    7. Hygiène et tenue du poste - Note A/B/C
    8. Participation aux activités du CSR - Note A/B/C
    Légende: A = Bon (satisfaisant), B = Moyen (à améliorer), C = Faible (nécessite formation)

    V. RESPONSABILITÉS
    - Chef de Centre: Supervision globale et validation des rapports
    - Major: Coordination des activités de supervision
    - Infirmier superviseur: Supervisions terrain et formation des ASC
    - Chargé PEV: Supervision des activités PEV et CPS
    - Sage-femme: Supervision des activités PTME et santé maternelle

    VI. AFFICHAGE
    - Lieu d'affichage: Tableau d'affichage du CSR NAGNENEFOUN
    - Période d'affichage: {donnees.get("periode", "2026")}
    - Responsable de l'affichage: Chef de Centre adjoint
    """
                meta = {
                    "Établissement": donnees.get("nom_etablissement", ""),
                    "Période": donnees.get("periode", ""),
                }
                doc = creer_document_word(type_doc, contenu, meta)

                from io import BytesIO

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                nom_fichier = (
                    f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx"
                )
                chemin_fichier = os.path.join(dossier_sortie, nom_fichier)
                with open(chemin_fichier, "wb") as f:
                    f.write(buffer.getvalue())

                st.success(f"✅ Document généré! Sauvegardé dans: documents_generes/")

                st.download_button(
                    "📥 Télécharger",
                    buffer,
                    f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                with st.expander("👁️ Aperçu"):
                    st.text(contenu)

            else:
                prompts = PROMPTS.get(doc_key)

                if prompts:
                    sections = get_sections_template(doc_key)
                    sections_str = "\n".join([f"I. {s}" for s in sections])

                    # Ajouter les sections aux données
                    donnees["sections"] = sections_str

                user_prompt = prompts["user"].format(**donnees)
                contenu = generer_avec_groq(prompts["system"], user_prompt)

                if "Erreur" in contenu:
                    st.error(contenu)
                else:
                    meta = {
                        "Établissement": donnees.get("nom_etablissement", ""),
                        "Période": donnees.get("periode", ""),
                    }

                    doc = creer_document_word(type_doc, contenu, meta)

                    from io import BytesIO

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    # Sauvegarder dans le dossier
                    nom_fichier = f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx"
                    chemin_fichier = os.path.join(dossier_sortie, nom_fichier)
                    with open(chemin_fichier, "wb") as f:
                        f.write(buffer.getvalue())

                    st.success(
                        f"✅ Document généré! Sauvegardé dans: documents_generes/"
                    )

                    st.download_button(
                        "📥 Télécharger",
                        buffer,
                        f"{type_doc}_{donnees.get('nom_etablissement', 'document')}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                    with st.expander("👁️ Aperçu"):
                        st.text(contenu)


# =============================================================================
# PAGE: CAHIER DE PRÉSENCE (Norme 2.01e - 20pts)
# =============================================================================

PRESENCE_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "presence_data.json"
)


def charger_presence():
    try:
        if os.path.exists(PRESENCE_FILE):
            with open(PRESENCE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


def sauvegarder_presence(data):
    with open(PRESENCE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def page_cahier_presence():
    st.title("📝 Cahier de Présence")
    st.caption("Norme ESPC 2.01e - Pointage quotidien du personnel (20 pts)")
    st.divider()

    personnel_liste = [
        {"nom": "KOUASSI Yao", "fonction": "Chef de Centre"},
        {"nom": "KONÉ Abibata", "fonction": "Major / Sage-femme"},
        {"nom": "TOURÉ Fatoumata", "fonction": "Chargée PEV"},
        {"nom": "KOFFI Aka", "fonction": "Chargé Paludisme"},
        {"nom": "N'GUESSAN Kouamé", "fonction": "Chargé VIH/PTME"},
        {"nom": "KOUAKOU Akissi", "fonction": "Chargée CPN"},
        {"nom": "BAMBA Sékou", "fonction": "Chargé Pharmacie"},
        {"nom": "KRA Adjo", "fonction": "Agent d'entretien"},
        {"nom": "DIBI Franck", "fonction": "Planton"},
        {"nom": "KOUAMÉ Bertine", "fonction": "Secrétaire"},
        {"nom": "GNAHOUA Olivier", "fonction": "Gardien"},
        {"nom": "KONAN Blanche", "fonction": "ASC"},
    ]

    presence_data = charger_presence()
    today = datetime.now().strftime("%Y-%m-%d")
    jour_label = datetime.now().strftime("%A %d/%m/%Y")

    tab1, tab2, tab3 = st.tabs(
        ["✅ Pointage du jour", "📊 Historique", "📥 Export Word"]
    )

    with tab1:
        st.markdown(f"### 📅 {jour_label}")
        st.info("Cochez les agents présents aujourd'hui, puis enregistrez le pointage.")

        avec_retard = []
        absents = []

        for p in personnel_liste:
            col1, col2, col3, col4 = st.columns([3, 2, 1, 1])
            col1.markdown(f"**{p['nom']}**")
            col2.caption(p["fonction"])
            present = col3.checkbox("Présent", key=f"pres_{p['nom']}_{today}")
            retard = col4.checkbox("Retard", key=f"retard_{p['nom']}_{today}")
            if present and retard:
                avec_retard.append(p["nom"])
            if not present:
                absents.append(p["nom"])

        st.divider()
        st.markdown(
            f"**Résumé:** {len(personnel_liste) - len(absents)}/{len(personnel_liste)} présent(s) | {len(avec_retard)} en retard | {len(absents)} absent(s)"
        )

        if absents:
            st.warning(f"Absents: {', '.join(absents)}")
        if avec_retard:
            st.info(f"En retard: {', '.join(avec_retard)}")

        heure_arrivee = st.time_input("Heure d'arrivée commune", key="heure_arr")
        notes_jour = st.text_area("Notes / Observations du jour", key="notes_jour")

        if st.button(
            "💾 Enregistrer le pointage du jour",
            type="primary",
            use_container_width=True,
        ):
            entry = {
                "date": today,
                "heure_arrivee": str(heure_arrivee),
                "personnel": [],
                "notes": notes_jour,
            }
            for p in personnel_liste:
                present_key = f"pres_{p['nom']}_{today}"
                retard_key = f"retard_{p['nom']}_{today}"
                entry["personnel"].append(
                    {
                        "nom": p["nom"],
                        "fonction": p["fonction"],
                        "present": st.session_state.get(present_key, False),
                        "retard": st.session_state.get(retard_key, False),
                    }
                )

            # Remove existing entry for today if any
            presence_data = [d for d in presence_data if d["date"] != today]
            presence_data.append(entry)
            sauvegarder_presence(presence_data)
            st.success("✅ Pointage enregistré!")

    with tab2:
        st.markdown("### 📊 Historique des pointages")

        if presence_data:
            sorted_data = sorted(presence_data, key=lambda x: x["date"], reverse=True)

            for entry in sorted_data[:30]:
                presents = [p for p in entry["personnel"] if p["present"]]
                absents_jour = [p for p in entry["personnel"] if not p["present"]]
                retardataires = [p for p in entry["personnel"] if p.get("retard")]

                with st.expander(
                    f"📅 {entry['date']} — {len(presents)}/{len(entry['personnel'])} présent(s)"
                ):
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Présents", len(presents))
                    c2.metric("Retards", len(retardataires))
                    c3.metric("Absents", len(absents_jour))

                    if absents_jour:
                        st.warning(
                            f"**Absents:** {', '.join([a['nom'] for a in absents_jour])}"
                        )
                    if retardataires:
                        st.info(
                            f"**En retard:** {', '.join([r['nom'] for r in retardataires])}"
                        )
                    if entry.get("notes"):
                        st.text(f"Notes: {entry['notes']}")
        else:
            st.info(
                "Aucun pointage enregistré. Commencez par enregistrer le pointage du jour."
            )

    with tab3:
        st.markdown("### 📥 Exporter le cahier de présence en Word")

        mois_export = st.selectbox(
            "Mois",
            [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ],
            index=datetime.now().month - 1,
        )
        annee_export = st.number_input(
            "Année", min_value=2020, max_value=2030, value=datetime.now().year
        )

        if st.button("📄 Générer le cahier de présence Word"):
            mois_num = [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ].index(mois_export) + 1

            jours_dans_mois = calendar.monthrange(annee_export, mois_num)[1]

            doc = Document()
            heading = doc.add_heading("CAHIER DE PRÉSENCE", level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                "MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA CMU"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph("RÉGION SANITAIRE DU PORO - DISTRICT KORHOGO 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph("CSR NAGNENEFOUN")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            p = doc.add_paragraph()
            run = p.add_run(f"Mois de {mois_export} {annee_export}")
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            table = doc.add_table(
                rows=len(personnel_liste) + 1, cols=jours_dans_mois + 2
            )
            table.style = "Table Grid"

            table.rows[0].cells[0].text = "NOM & PRÉNOMS"
            table.rows[0].cells[1].text = "FONCTION"
            for j in range(jours_dans_mois):
                table.rows[0].cells[j + 2].text = f"{j + 1}"

            for i, p_data in enumerate(personnel_liste):
                table.rows[i + 1].cells[0].text = p_data["nom"]
                table.rows[i + 1].cells[1].text = p_data["fonction"]

                for j in range(jours_dans_mois):
                    jour_str = f"{annee_export}-{mois_num:02d}-{j + 1:02d}"
                    day_entry = next(
                        (d for d in presence_data if d["date"] == jour_str), None
                    )
                    if day_entry:
                        emp = next(
                            (
                                e
                                for e in day_entry["personnel"]
                                if e["nom"] == p_data["nom"]
                            ),
                            None,
                        )
                        if emp:
                            if emp["present"]:
                                table.rows[i + 1].cells[j + 2].text = (
                                    "P" if not emp.get("retard") else "R"
                                )
                            else:
                                table.rows[i + 1].cells[j + 2].text = "A"

            doc.add_paragraph()
            doc.add_paragraph("Légende: P = Présent, R = Retard, A = Absent")
            doc.add_paragraph()

            doc.add_paragraph("Signatures:")
            doc.add_paragraph()
            p = doc.add_paragraph("Le Chef de Centre: ____________________")
            p = doc.add_paragraph("Le Major: ____________________")
            p = doc.add_paragraph("Date: ____/____/______")

            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Télécharger le Cahier de Présence (.docx)",
                buffer,
                f"cahier_presence_{mois_export}_{annee_export}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# =============================================================================
# PAGE: SUIVI STOCK MÉDICAMENTS (Norme 12.01 - 75pts)
# =============================================================================

STOCK_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "stock_medicaments.json"
)

MEDICAMENTS_DEFAUT = [
    {
        "nom": "Artemether-Luméfantrine 20/120mg (Adulte)",
        "categorie": "Antipaludique",
        "seuil_alerte": 100,
        "unite": "comprimés",
    },
    {
        "nom": "Artemether-Luméfantrine 15/120mg (Enfant)",
        "categorie": "Antipaludique",
        "seuil_alerte": 200,
        "unite": "comprimés",
    },
    {
        "nom": "Artesunate injectable 60mg",
        "categorie": "Antipaludique",
        "seuil_alerte": 50,
        "unite": "ampoules",
    },
    {
        "nom": "Sulfadoxine-Pyriméthamine (SP)",
        "categorie": "Antipaludique",
        "seuil_alerte": 150,
        "unite": "comprimés",
    },
    {
        "nom": "Amoxicilline 500mg",
        "categorie": "Antibiotique",
        "seuil_alerte": 200,
        "unite": "gélules",
    },
    {
        "nom": "Amoxicilline 250mg (Enfant)",
        "categorie": "Antibiotique",
        "seuil_alerte": 300,
        "unite": "gélules",
    },
    {
        "nom": "Metronidazole 250mg",
        "categorie": "Antibiotique",
        "seuil_alerte": 150,
        "unite": "comprimés",
    },
    {
        "nom": "ORS (Sels de réhydratation)",
        "categorie": "Réhydratation",
        "seuil_alerte": 100,
        "unite": "sachets",
    },
    {
        "nom": "Zinc 20mg",
        "categorie": "Nutrition",
        "seuil_alerte": 200,
        "unite": "comprimés",
    },
    {
        "nom": "Paracétamol 500mg",
        "categorie": "Antidouleur",
        "seuil_alerte": 300,
        "unite": "comprimés",
    },
    {
        "nom": "Paracétamol 100mg (Enfant)",
        "categorie": "Antidouleur",
        "seuil_alerte": 200,
        "unite": "comprimés",
    },
    {
        "nom": "Mébendazole 500mg",
        "categorie": "Antihelminthique",
        "seuil_alerte": 100,
        "unite": "comprimés",
    },
    {
        "nom": "Vitamine A",
        "categorie": "Nutrition",
        "seuil_alerte": 100,
        "unite": "gélules",
    },
    {
        "nom": "Fer + Acide folique",
        "categorie": "Nutrition",
        "seuil_alerte": 200,
        "unite": "comprimés",
    },
    {
        "nom": "TDR Paludisme (boîte 25)",
        "categorie": "Diagnostic",
        "seuil_alerte": 20,
        "unite": "boîtes",
    },
    {
        "nom": "Gants stériles (boîte 100)",
        "categorie": "Consommable",
        "seuil_alerte": 10,
        "unite": "boîtes",
    },
    {
        "nom": "Seringues 5ml",
        "categorie": "Consommable",
        "seuil_alerte": 50,
        "unite": "unités",
    },
    {
        "nom": "Eau pour injection 10ml",
        "categorie": "Consommable",
        "seuil_alerte": 30,
        "unite": "ampoules",
    },
]

STOCK_CATEGORIES = [
    "Antipaludique",
    "Antibiotique",
    "Réhydratation",
    "Nutrition",
    "Antidouleur",
    "Antihelminthique",
    "Diagnostic",
    "Consommable",
]


def charger_stock():
    try:
        if os.path.exists(STOCK_FILE):
            with open(STOCK_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return MEDICAMENTS_DEFAUT


def sauvegarder_stock(data):
    with open(STOCK_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def page_suivi_stock():
    st.title("💊 Suivi Stock Médicaments")
    st.caption(
        "Norme ESPC 12.01 — Disponibilité continue des médicaments traceurs (75 pts)"
    )
    st.divider()

    stock_data = charger_stock()
    today = datetime.now().strftime("%Y-%m-%d")

    tab1, tab2, tab3, tab4 = st.tabs(
        [
            "📊 État des stocks",
            "➕ Mouvement de stock",
            "🚨 Alertes rupture",
            "📥 Export Word",
        ]
    )

    with tab1:
        st.markdown("### 📊 État actuel des stocks")

        cat_filter = st.selectbox(
            "Filtrer par catégorie", ["Toutes"] + STOCK_CATEGORIES
        )
        affichage = stock_data
        if cat_filter != "Toutes":
            affichage = [s for s in stock_data if s.get("categorie") == cat_filter]

        for i, med in enumerate(affichage):
            stock_actuel = med.get("stock_actuel", 0)
            seuil = med.get("seuil_alerte", 0)
            unite = med.get("unite", "unités")

            if stock_actuel <= 0:
                status = "🔴 RUPTURE"
            elif stock_actuel <= seuil:
                status = "🟡 ALERTE"
            else:
                status = "🟢 OK"

            col1, col2, col3, col4 = st.columns([4, 2, 2, 1])
            col1.markdown(f"**{med['nom']}**")
            col2.markdown(f"Stock: **{stock_actuel}** {unite}")
            col3.markdown(status)
            col4.caption(f"Seuil: {seuil}")

    with tab2:
        st.markdown("### ➕ Enregistrer un mouvement de stock")

        with st.form("mouvement_stock"):
            med_options = [m["nom"] for m in stock_data]
            med_select = st.selectbox("Médicament", med_options)
            type_mouvement = st.selectbox(
                "Type de mouvement",
                ["Entrée (réception)", "Sortie (consommation)", "Périmé/Détruit"],
            )
            quantite = st.number_input("Quantité", min_value=1, value=1)
            source = st.text_input(
                "Source / Fournisseur", "District / Pharmacie centrale"
            )
            observation = st.text_area("Observation", "")

            submitted = st.form_submit_button(
                "💾 Enregistrer le mouvement", use_container_width=True
            )

            if submitted:
                idx = next(
                    (i for i, m in enumerate(stock_data) if m["nom"] == med_select),
                    None,
                )
                if idx is not None:
                    if "Entrée" in type_mouvement:
                        stock_data[idx]["stock_actuel"] = (
                            stock_data[idx].get("stock_actuel", 0) + quantite
                        )
                    else:
                        stock_data[idx]["stock_actuel"] = max(
                            0, stock_data[idx].get("stock_actuel", 0) - quantite
                        )

                    sauvegarder_stock(stock_data)
                    st.success(
                        f"✅ Mouvement enregistré: {type_mouvement} de {quantite} {stock_data[idx].get('unite', 'unités')}"
                    )

    with tab3:
        st.markdown("### 🚨 Alertes de rupture")

        ruptures = [s for s in stock_data if s.get("stock_actuel", 0) <= 0]
        alertes = [
            s
            for s in stock_data
            if 0 < s.get("stock_actuel", 0) <= s.get("seuil_alerte", 0)
        ]

        if ruptures:
            st.error(f"**{len(ruptures)} médicament(s) en RUPTURE de stock:**")
            for r in ruptures:
                st.markdown(f"- 🔴 **{r['nom']}** — Stock: 0 {r.get('unite', '')}")
        else:
            st.success("Aucune rupture de stock.")

        if alertes:
            st.warning(f"**{len(alertes)} médicament(s) en stock d'ALERTE:**")
            for a in alertes:
                st.markdown(
                    f"- 🟡 **{a['nom']}** — Stock: {a.get('stock_actuel', 0)} {a.get('unite', '')} (seuil: {a.get('seuil_alerte', 0)})"
                )
        else:
            st.success("Tous les stocks sont au-dessus du seuil d'alerte.")

        st.divider()
        total = len(stock_data)
        ok = len(
            [
                s
                for s in stock_data
                if s.get("stock_actuel", 0) > s.get("seuil_alerte", 0)
            ]
        )
        st.metric(
            "Taux de disponibilité",
            f"{ok}/{total} ({ok * 100 // total if total else 0}%)",
        )

    with tab4:
        st.markdown("### 📥 Exporter le suivi des stocks en Word")

        if st.button("📄 Générer le rapport de stock Word", key="export_stock"):
            doc = Document()
            heading = doc.add_heading(
                "SUIVI DE LA DISPONIBILITÉ DES MÉDICAMENTS TRACEURS", level=0
            )
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                "CSR NAGNENEFOUN — District KORHOGO 1 — Région du PORO"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            p = doc.add_paragraph()
            run = p.add_run(f"État des stocks au {today}")
            run.bold = True
            doc.add_paragraph()

            table = doc.add_table(rows=len(stock_data) + 1, cols=5)
            table.style = "Table Grid"

            headers = [
                "Désignation",
                "Catégorie",
                "Stock actuel",
                "Seuil alerte",
                "Statut",
            ]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for i, med in enumerate(stock_data):
                stock_actuel = med.get("stock_actuel", 0)
                seuil = med.get("seuil_alerte", 0)
                if stock_actuel <= 0:
                    statut = "RUPTURE"
                elif stock_actuel <= seuil:
                    statut = "ALERTE"
                else:
                    statut = "OK"

                table.rows[i + 1].cells[0].text = med["nom"]
                table.rows[i + 1].cells[1].text = med.get("categorie", "")
                table.rows[i + 1].cells[
                    2
                ].text = f"{stock_actuel} {med.get('unite', '')}"
                table.rows[i + 1].cells[3].text = f"{seuil} {med.get('unite', '')}"
                table.rows[i + 1].cells[4].text = statut

            doc.add_paragraph()
            doc.add_paragraph(f"Rapport établi le {today} par le Chargé de Pharmacie")
            doc.add_paragraph("Signatures:")
            doc.add_paragraph("Le Chargé de Pharmacie: ____________________")
            doc.add_paragraph("Le Chef de Centre: ____________________")

            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Télécharger le rapport de stock (.docx)",
                buffer,
                f"suivi_stock_{today}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# =============================================================================
# PAGE: ENQUÊTE SATISFACTION USAGERS (Norme 13.01 - 150pts)
# =============================================================================

ENQUETE_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "enquete_satisfaction.json"
)

QUESTIONS_ENQUETE = [
    {"id": 1, "texte": "Accueil à l'entrée du centre", "categorie": "Accueil"},
    {"id": 2, "texte": "Temps d'attente avant consultation", "categorie": "Accueil"},
    {"id": 3, "texte": "Écoute et attention du soignant", "categorie": "Soins"},
    {"id": 4, "texte": "Clarté des explications reçues", "categorie": "Soins"},
    {"id": 5, "texte": "Propreté des locaux", "categorie": "Environnement"},
    {"id": 6, "texte": "Disponibilité des médicaments", "categorie": "Pharmacie"},
    {"id": 7, "texte": "Respect de la vie privée", "categorie": "Soins"},
    {"id": 8, "texte": "Coût des soins acceptable", "categorie": "Coût"},
    {
        "id": 9,
        "texte": "Facilité de trouver le service souhaité",
        "categorie": "Accueil",
    },
    {"id": 10, "texte": "Satisfaction globale", "categorie": "Général"},
]


def charger_enquete():
    try:
        if os.path.exists(ENQUETE_FILE):
            with open(ENQUETE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


def sauvegarder_enquete(data):
    with open(ENQUETE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def page_satisfaction():
    st.title("⭐ Enquête Satisfaction Usagers")
    st.caption("Norme ESPC 13.01 — Enquête de satisfaction des patients (150 pts)")
    st.divider()

    enquete_data = charger_enquete()
    today = datetime.now().strftime("%Y-%m-%d")

    tab1, tab2, tab3 = st.tabs(
        ["📝 Nouvelle enquête", "📊 Résultats", "📥 Export Word"]
    )

    with tab1:
        st.markdown("### 📝 Formulaire de satisfaction")
        st.markdown(
            "Demandez au patient d'évaluer chaque critère de **1 (Très mécontent) à 5 (Très content)**."
        )

        with st.form("form_enquete"):
            col1, col2 = st.columns(2)
            with col1:
                sexe = st.selectbox("Sexe", ["Homme", "Femme"])
                age = st.number_input("Âge", min_value=0, max_value=120, value=30)
            with col2:
                service_consulte = st.selectbox(
                    "Service consulté",
                    [
                        "Consultation générale",
                        "Maternité/CPN",
                        "PEV/Vaccination",
                        "Pharmacie",
                        "Urgence",
                        "Autre",
                    ],
                )
                date_enquete = st.date_input("Date de l'enquête")

            st.markdown("---")
            reponses = {}
            for q in QUESTIONS_ENQUETE:
                reponses[str(q["id"])] = st.slider(
                    f"{q['texte']} ({q['categorie']})",
                    min_value=1,
                    max_value=5,
                    value=3,
                    key=f"q_{q['id']}",
                )

            commentaire = st.text_area("Commentaire / Suggestion du patient")

            submitted = st.form_submit_button(
                "💾 Enregistrer l'enquête", use_container_width=True
            )

            if submitted:
                entry = {
                    "date": str(date_enquete),
                    "sexe": sexe,
                    "age": age,
                    "service": service_consulte,
                    "reponses": reponses,
                    "commentaire": commentaire,
                    "score_moyen": round(sum(reponses.values()) / len(reponses), 1),
                }
                enquete_data.append(entry)
                sauvegarder_enquete(enquete_data)
                st.success(
                    f"✅ Enquête enregistrée! Score moyen: {entry['score_moyen']}/5"
                )

    with tab2:
        st.markdown("### 📊 Résultats des enquêtes")

        if enquete_data:
            total_enquetes = len(enquete_data)
            score_global = round(
                sum(e["score_moyen"] for e in enquete_data) / total_enquetes, 1
            )

            col1, col2, col3 = st.columns(3)
            col1.metric("Total enquêtes", total_enquetes)
            col2.metric("Score moyen global", f"{score_global}/5")
            col3.metric("Taux satisfaction", f"{score_global * 20:.0f}%")

            st.divider()
            st.markdown("### 📈 Scores par critère")
            for q in QUESTIONS_ENQUETE:
                scores = [int(e["reponses"].get(str(q["id"]), 3)) for e in enquete_data]
                moyenne = round(sum(scores) / len(scores), 1) if scores else 0
                st.markdown(f"- **{q['texte']}**: {moyenne}/5")

            st.divider()
            st.markdown("### 📊 Répartition par service")
            services = {}
            for e in enquete_data:
                s = e.get("service", "Non précisé")
                if s not in services:
                    services[s] = []
                services[s].append(e["score_moyen"])
            for s, scores in services.items():
                avg = round(sum(scores) / len(scores), 1)
                st.markdown(f"- **{s}**: {avg}/5 ({len(scores)} enquêtes)")

            st.divider()
            st.markdown("### 💬 Derniers commentaires")
            for e in reversed(enquete_data[-10:]):
                if e.get("commentaire"):
                    st.markdown(
                        f"- _{e['date']}_ ({e.get('service', '')}): {e['commentaire']}"
                    )
        else:
            st.info("Aucune enquête enregistrée. Commencez par remplir le formulaire.")

    with tab3:
        st.markdown("### 📥 Exporter les résultats en Word")

        mois_export = st.selectbox(
            "Mois",
            [
                "Janvier",
                "Février",
                "Mars",
                "Avril",
                "Mai",
                "Juin",
                "Juillet",
                "Août",
                "Septembre",
                "Octobre",
                "Novembre",
                "Décembre",
            ],
            index=datetime.now().month - 1,
            key="mois_enquete",
        )
        annee_export = st.number_input(
            "Année",
            min_value=2020,
            max_value=2030,
            value=datetime.now().year,
            key="annee_enquete",
        )

        if st.button("📄 Générer le rapport de satisfaction Word"):
            doc = Document()
            heading = doc.add_heading(
                "RAPPORT D'ENQUÊTE DE SATISFACTION DES USAGERS", level=0
            )
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                "CSR NAGNENEFOUN — District KORHOGO 1 — Région du PORO"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            total_enquetes = len(enquete_data)
            score_global = round(
                sum(e["score_moyen"] for e in enquete_data) / max(total_enquetes, 1), 1
            )

            doc.add_paragraph(f"Période: {mois_export} {annee_export}")
            doc.add_paragraph(f"Nombre d'enquêtes: {total_enquetes}")
            doc.add_paragraph(
                f"Score moyen global: {score_global}/5 ({score_global * 20:.0f}%)"
            )
            doc.add_paragraph()

            doc.add_heading("Scores par critère", level=1)
            table = doc.add_table(rows=len(QUESTIONS_ENQUETE) + 1, cols=4)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Critère"
            table.rows[0].cells[1].text = "Catégorie"
            table.rows[0].cells[2].text = "Score moyen"
            table.rows[0].cells[3].text = "Appréciation"

            for i, q in enumerate(QUESTIONS_ENQUETE):
                scores = [int(e["reponses"].get(str(q["id"]), 3)) for e in enquete_data]
                moyenne = round(sum(scores) / len(scores), 1) if scores else 0
                if moyenne >= 4:
                    appreciation = "Satisfaisant"
                elif moyenne >= 3:
                    appreciation = "Moyen"
                else:
                    appreciation = "Insuffisant"
                table.rows[i + 1].cells[0].text = q["texte"]
                table.rows[i + 1].cells[1].text = q["categorie"]
                table.rows[i + 1].cells[2].text = f"{moyenne}/5"
                table.rows[i + 1].cells[3].text = appreciation

            doc.add_paragraph()
            doc.add_heading("Recommandations", level=1)
            doc.add_paragraph(
                "1. Améliorer les points identifiés avec un score inférieur à 3/5"
            )
            doc.add_paragraph(
                "2. Maintenir les points forts identifiés par les usagers"
            )
            doc.add_paragraph("3. Diffuser les résultats à l'ensemble du personnel")
            doc.add_paragraph(
                "4. Mettre en place un plan d'action pour les améliorations"
            )

            doc.add_paragraph()
            doc.add_paragraph("Rapport établi le " + today)
            doc.add_paragraph("Le Chef de Centre: ____________________")

            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Télécharger le rapport de satisfaction (.docx)",
                buffer,
                f"rapport_satisfaction_{mois_export}_{annee_export}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# =============================================================================
# PAGE: FICHES STOCK ASC (Norme 15.01 - 50pts)
# =============================================================================

STOCK_ASC_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "stock_asc.json"
)

INTRANTS_ASC = [
    {"nom": "TDR Paludisme (boîte 25)", "unite": "boîtes"},
    {"nom": "ACT Adulte (cure)", "unite": "cures"},
    {"nom": "ACT Enfant (cure)", "unite": "cures"},
    {"nom": "MII (Moustiquaires Imprégnées)", "unite": "moustiquaires"},
    {"nom": "Préservatifs (paquet de 3)", "unite": "paquets"},
    {"nom": "ORS (sachets)", "unite": "sachets"},
    {"nom": "Zinc 20mg", "unite": "comprimés"},
    {"nom": "Vitamine A", "unite": "gélules"},
    {"nom": "MUAC (Bracelet malnutrition)", "unite": "unités"},
    {"nom": "Fiches de sensibilisation", "unite": "exemplaires"},
]

ASC_LISTE = [
    {"nom": "KONAN Blanche", "village": "Nagnenefoun Centre"},
    {"nom": "YEO Aminata", "village": "Korhogo Nord"},
    {"nom": "DIALLO Moussa", "village": "Ferké"},
]


def charger_stock_asc():
    try:
        if os.path.exists(STOCK_ASC_FILE):
            with open(STOCK_ASC_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


def sauvegarder_stock_asc(data):
    with open(STOCK_ASC_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def page_stock_asc():
    st.title("📦 Fiches Stock ASC")
    st.caption("Norme ESPC 15.01 — Suivi des intrants et médicaments des ASC (50 pts)")
    st.divider()

    stock_asc = charger_stock_asc()
    today = datetime.now().strftime("%Y-%m-%d")

    tab1, tab2, tab3 = st.tabs(
        ["📊 État des stocks ASC", "➕ Mouvement", "📥 Export Word"]
    )

    with tab1:
        st.markdown("### 📊 Suivi des stocks par ASC")

        asc_select = st.selectbox("Choisir un ASC", [a["nom"] for a in ASC_LISTE])
        asc_info = next((a for a in ASC_LISTE if a["nom"] == asc_select), None)

        st.info(f"Village/Aire: {asc_info['village'] if asc_info else ''}")

        # Get or init stock for this ASC
        asc_stock = next((s for s in stock_asc if s["asc"] == asc_select), None)
        if not asc_stock:
            asc_stock = {
                "asc": asc_select,
                "village": asc_info["village"] if asc_info else "",
                "intrants": {},
            }
            stock_asc.append(asc_stock)
            sauvegarder_stock_asc(stock_asc)

        for intrant in INTRANTS_ASC:
            stock_val = asc_stock["intrants"].get(intrant["nom"], 0)
            col1, col2, col3 = st.columns([4, 2, 2])
            col1.markdown(f"**{intrant['nom']}**")
            col2.markdown(f"Stock: **{stock_val}** {intrant['unite']}")
            if stock_val == 0:
                col3.markdown("🔴 RUPTURE")
            else:
                col3.markdown("🟢 OK")

    with tab2:
        st.markdown("### ➕ Enregistrer un mouvement")

        with st.form("mouvement_asc"):
            asc_mv = st.selectbox("ASC", [a["nom"] for a in ASC_LISTE], key="asc_mv")
            intrant_mv = st.selectbox("Intrant", [i["nom"] for i in INTRANTS_ASC])
            type_mv = st.selectbox(
                "Type",
                [
                    "Entrée (réception du centre)",
                    "Sortie (distribution communautaire)",
                    "Perte/Expiration",
                ],
            )
            quantite_mv = st.number_input("Quantité", min_value=1, value=1)
            date_mv = st.date_input("Date")
            obs_mv = st.text_area("Observation", "")

            if st.form_submit_button("💾 Enregistrer", use_container_width=True):
                asc_stock = next((s for s in stock_asc if s["asc"] == asc_mv), None)
                if asc_stock:
                    current = asc_stock["intrants"].get(intrant_mv, 0)
                    if "Entrée" in type_mv:
                        asc_stock["intrants"][intrant_mv] = current + quantite_mv
                    else:
                        asc_stock["intrants"][intrant_mv] = max(
                            0, current - quantite_mv
                        )
                    sauvegarder_stock_asc(stock_asc)
                    st.success("✅ Mouvement enregistré!")

    with tab3:
        st.markdown("### 📥 Exporter les fiches stock ASC en Word")

        if st.button("📄 Générer les fiches stock ASC Word", key="export_asc"):
            doc = Document()
            heading = doc.add_heading(
                "FICHES DE STOCK DES AGENTS DE SANTÉ COMMUNAUTAIRE", level=0
            )
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                "CSR NAGNENEFOUN — District KORHOGO 1 — Région du PORO"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            for asc in ASC_LISTE:
                asc_stock = next((s for s in stock_asc if s["asc"] == asc["nom"]), None)

                doc.add_heading(
                    f"ASC: {asc['nom']} — Village: {asc['village']}", level=1
                )

                table = doc.add_table(rows=len(INTRANTS_ASC) + 1, cols=3)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = "Intrant"
                table.rows[0].cells[1].text = "Stock disponible"
                table.rows[0].cells[2].text = "Statut"

                for i, intrant in enumerate(INTRANTS_ASC):
                    stock_val = (
                        asc_stock["intrants"].get(intrant["nom"], 0) if asc_stock else 0
                    )
                    statut = "RUPTURE" if stock_val == 0 else "OK"
                    table.rows[i + 1].cells[
                        0
                    ].text = f"{intrant['nom']} ({intrant['unite']})"
                    table.rows[i + 1].cells[1].text = str(stock_val)
                    table.rows[i + 1].cells[2].text = statut

                doc.add_paragraph()

            doc.add_paragraph(f"Fiche établie le {today}")
            doc.add_paragraph("Le Chargé de Programme: ____________________")

            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Télécharger les fiches stock ASC (.docx)",
                buffer,
                f"fiches_stock_asc_{today}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# =============================================================================
# PAGE: FICHE NOTIFICATION DÉCÈS MATERNEL (Norme 8.01 - 20pts)
# =============================================================================

DECES_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "notifications_deces.json"
)


def charger_deces():
    try:
        if os.path.exists(DECES_FILE):
            with open(DECES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


def sauvegarder_deces(data):
    with open(DECES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def page_notification_deces():
    st.title("📋 Notification Décès Maternel")
    st.caption("Norme ESPC 8.01 — Outils de notification des décès maternels (20 pts)")
    st.divider()

    deces_data = charger_deces()

    tab1, tab2, tab3 = st.tabs(
        ["📝 Nouvelle notification", "📋 Registre", "📥 Export Word"]
    )

    with tab1:
        st.markdown("### 📝 Fiche de notification de décès maternel")
        st.warning(
            "Norme 8.01 a: Tenir à disposition 5 fiches vierges de notification de décès maternel"
        )

        with st.form("form_deces"):
            col1, col2 = st.columns(2)
            with col1:
                nom_defunt = st.text_input("Nom et prénoms de la défunte")
                age = st.number_input("Âge", min_value=12, max_value=60, value=25)
                adresse = st.text_input("Adresse / Village")
            with col2:
                date_deces = st.date_input("Date du décès")
                lieu = st.selectbox(
                    "Lieu du décès",
                    [
                        "Centre de santé",
                        "Domicile",
                        "Chemin",
                        "Hôpital de référence",
                        "Autre",
                    ],
                )
                date_accouchement = st.date_input(
                    "Date de l'accouchement (si applicable)"
                )

            col3, col4 = st.columns(2)
            with col3:
                nb_grossesses = st.number_input(
                    "Nombre de grossesses", min_value=1, max_value=15, value=2
                )
                nb_enfants = st.number_input(
                    "Nombre d'enfants vivants", min_value=0, max_value=15, value=1
                )
            with col4:
                type_accouchement = st.selectbox(
                    "Type d'accouchement",
                    ["Spontané", "Césarienne", "Autre", "Non applicable"],
                )
                duree_grossesse = st.number_input(
                    "Durée de grossesse (semaines)",
                    min_value=20,
                    max_value=45,
                    value=38,
                )

            st.markdown("**Circonstances du décès:**")
            circonstances = st.selectbox(
                "Circonstance principale",
                [
                    "Hémorragie du post-partum",
                    "Éclampsie / Préeclampsie",
                    "Infection / Sepse",
                    "Obstruction mécanique",
                    "Rupture utérine",
                    "Embolie",
                    "Autre",
                    "Inconnue",
                ],
            )

            delai_prendre_charge = st.selectbox(
                "Délai avant prise en charge",
                [
                    "Moins de 1 heure",
                    "1 à 3 heures",
                    "3 à 6 heures",
                    "Plus de 6 heures",
                    "Non pris en charge",
                ],
            )

            observation = st.text_area("Observations complémentaires")

            if st.form_submit_button(
                "💾 Enregistrer la notification", use_container_width=True
            ):
                entry = {
                    "date_notification": datetime.now().strftime("%Y-%m-%d"),
                    "nom_defunt": nom_defunt,
                    "age": age,
                    "adresse": adresse,
                    "date_deces": str(date_deces),
                    "lieu_deces": lieu,
                    "date_accouchement": str(date_accouchement),
                    "nb_grossesses": nb_grossesses,
                    "nb_enfants": nb_enfants,
                    "type_accouchement": type_accouchement,
                    "duree_grossesse": duree_grossesse,
                    "circonstances": circonstances,
                    "delai_prendre_charge": delai_prendre_charge,
                    "observation": observation,
                }
                deces_data.append(entry)
                sauvegarder_deces(deces_data)
                st.success("✅ Notification enregistrée!")

    with tab2:
        st.markdown("### 📋 Registre des notifications de décès maternels")

        if deces_data:
            for i, d in enumerate(reversed(deces_data)):
                with st.expander(f"📋 {d['nom_defunt']} — Décès le {d['date_deces']}"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**Âge:** {d['age']} ans")
                        st.markdown(f"**Adresse:** {d['adresse']}")
                        st.markdown(f"**Lieu:** {d['lieu_deces']}")
                        st.markdown(f"**Circonstances:** {d['circonstances']}")
                    with col2:
                        st.markdown(f"**Grossesses:** {d['nb_grossesses']}")
                        st.markdown(f"**Enfants vivants:** {d['nb_enfants']}")
                        st.markdown(f"**Type accouchement:** {d['type_accouchement']}")
                        st.markdown(
                            f"**Délai prise en charge:** {d['delai_prendre_charge']}"
                        )
                    if d.get("observation"):
                        st.text(f"Observation: {d['observation']}")
        else:
            st.info("Aucune notification enregistrée.")

    with tab3:
        st.markdown("### 📥 Exporter les notifications en Word")

        if st.button(
            "📄 Générer le registre de notifications Word", key="export_deces"
        ):
            doc = Document()
            heading = doc.add_heading(
                "FICHE DE NOTIFICATION DE DÉCÈS MATERNEL", level=0
            )
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                "CSR NAGNENEFOUN — District KORHOGO 1 — Région du PORO"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph("Conforme à la Norme ESPC 8.01 a")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            if deces_data:
                for d in deces_data:
                    doc.add_heading(
                        f"Notification N° {deces_data.index(d) + 1}", level=2
                    )

                    table = doc.add_table(rows=8, cols=2)
                    table.style = "Table Grid"
                    table.rows[0].cells[0].text = "Nom et prénoms"
                    table.rows[0].cells[1].text = d["nom_defunt"]
                    table.rows[1].cells[0].text = "Âge"
                    table.rows[1].cells[1].text = f"{d['age']} ans"
                    table.rows[2].cells[0].text = "Adresse"
                    table.rows[2].cells[1].text = d["adresse"]
                    table.rows[3].cells[0].text = "Date du décès"
                    table.rows[3].cells[1].text = d["date_deces"]
                    table.rows[4].cells[0].text = "Lieu du décès"
                    table.rows[4].cells[1].text = d["lieu_deces"]
                    table.rows[5].cells[0].text = "Circonstances"
                    table.rows[5].cells[1].text = d["circonstances"]
                    table.rows[6].cells[0].text = "Délai prise en charge"
                    table.rows[6].cells[1].text = d["delai_prendre_charge"]
                    table.rows[7].cells[0].text = "Observations"
                    table.rows[7].cells[1].text = d.get("observation", "")

                    doc.add_paragraph()
            else:
                doc.add_paragraph("Aucune notification enregistrée.")
                doc.add_paragraph()
                doc.add_paragraph("(Fiches vierges prêtes à imprimer pour le terrain)")

                for i in range(5):
                    doc.add_heading(f"Fiche vierge N° {i + 1}", level=2)
                    table = doc.add_table(rows=8, cols=2)
                    table.style = "Table Grid"
                    table.rows[0].cells[0].text = "Nom et prénoms"
                    table.rows[1].cells[0].text = "Âge"
                    table.rows[2].cells[0].text = "Adresse"
                    table.rows[3].cells[0].text = "Date du décès"
                    table.rows[4].cells[0].text = "Lieu du décès"
                    table.rows[5].cells[0].text = "Circonstances"
                    table.rows[6].cells[0].text = "Délai prise en charge"
                    table.rows[7].cells[0].text = "Observations"
                    doc.add_paragraph()

            doc.add_paragraph()
            doc.add_paragraph("Registre tenu par le Chef de Centre")
            doc.add_paragraph(
                "Transmission au District Sanitaire de KORHOGO 1: ____________________"
            )

            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Télécharger le registre de notifications (.docx)",
                buffer,
                f"registre_notifications_deces_{today}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# =============================================================================
# ASSISTANT ESPC - CHATBOT
# =============================================================================


def get_espc_knowledge():
    return """GRILLE ESPC - Évaluation des ESPC (1600 points)

A. MANAGEMENT (600 pts)
1. Gouvernance (200 pts): PAA budgétisé (50), COGES (20), Transparence (35), Assurance qualité (75), Données (20)
2. RH (50 pts): Liste personnel (5), Fiches de poste (5), Planning (5), Gardes (5), Présence (20), Évaluations (10)
3. Finance (350 pts): Documents comptables (175), Primes (100), Flux financiers (75)

B. QUALITÉ DES SOINS (750 pts)
4. Accueil (175 pts): Signalétique (30), Dispositif accueil (40), CMU (75), Attente (30)
5. Sécurité (25 pts): Sécurité/éclairage/environnement
6. Hygiène (150 pts): Cadre hygiène (8), Procédures (5), Prévention infections (55), Stérilisation (5), Déchets (40), Entretien (15), Toilettes (22)
7. SONU (125 pts): Infrastructures (40), Personnel (7), Fonctions (24), Protocoles (3), Urgences (3), Médicaments (8), Nouveau-né (9), Post-partum (32)
8. Décès maternels (50 pts): Notification (20), Transmission district (30)
9. PF (50 pts): Local (8), Matériels (12), Personnel DIU (15), Personnel Implant (15)
10. Pathologies (50 pts): Paludisme (30), IRA-Diarrhée (20)
11. Pharmacie (50 pts): Local stockage (15), Outils gestion (20), Gestion qualité (15)
12. Médicaments traceurs (75 pts): Disponibilité continue (75)

C. SATISFACTION (150 pts)
13. Enquête satisfaction (150 pts)

D. COMMUNAUTAIRE (100 pts)
14. Supervision ASC (50 pts): Liste (3), Plan (2), Grille (3), Rapports (23), Transmission (15)
15. Médicaments ASC (50 pts): Fiches stock (10), Rapport activités (10), Ruptures (30)"""


def get_templates_knowledge():
    return """DOCUMENTS DISPONIBLES (17 documents + 5 modules):
1. PV Réunion Mensuelle - Réunion mensuelle personnel (Norme 1.01)
2. PV Réunion COGES - Réunion trimestrielle comité gestion (Norme 1.02)
3. PV Assemblée Générale - Assemblée annuelle (Norme 1.03)
4. Rapport Supervision ASC - Rapport mensuel supervision (Norme 14.01)
5. Rapport Plaintes/Suggestions - Boîte à suggestions (Norme 4.02)
6. Fiche de Poste - Par catégorie (Norme 2.01b)
7. Fiche de Nomination - 12 types (Note service, Arrêté, etc.)
8. Programme Réunions Trimestrielles - Calendrier stratégique
9. Calendrier Nettoyage Centre - Plan nettoyage (Norme 6.01)
10. Calendrier Réunions Mensuelles - 12 mois
11. Grille Supervision ASC - Évaluation ASC (Norme 14.01c)
12. Liste Personnel COGES - Liste officielle (Norme 1.02)
13. Plan Action Infections Nosocomiales - PCI (Norme 6.05)
14. Plan Supervision ASC - Plan annuel (Norme 14.01b)
15. Rapport Formation Personnel - Rapport formation (Norme 1.04b)
16. Liste Personnel Centre - Liste officielle (Norme 2.01a)
17. Note de Service - 6 types de notes de service

MODULES: Présence (2.01e), Stock Médicaments (12.01), Satisfaction (13.01), Stock ASC (15.01), Décès Maternel (8.01)"""


def get_saved_data_summary():
    summary = []
    try:
        stock = charger_stock()
        if stock:
            ruptures = [s for s in stock if s.get("stock_actuel", 0) <= 0]
            alertes = [
                s
                for s in stock
                if 0 < s.get("stock_actuel", 0) <= s.get("seuil_alerte", 0)
            ]
            summary.append(
                f"Stock médicaments: {len(stock)} produits, {len(ruptures)} ruptures, {len(alertes)} alertes"
            )
    except:
        pass
    try:
        presence = charger_presence()
        if presence:
            summary.append(f"Cahier de présence: {len(presence)} pointages")
    except:
        pass
    try:
        enquete = charger_enquete()
        if enquete:
            score = round(sum(e["score_moyen"] for e in enquete) / len(enquete), 1)
            summary.append(f"Satisfaction: {len(enquete)} enquêtes, score {score}/5")
    except:
        pass
    try:
        stock_asc = charger_stock_asc()
        if stock_asc:
            summary.append(f"Stock ASC: {len(stock_asc)} ASC suivis")
    except:
        pass
    try:
        deces = charger_deces()
        if deces:
            summary.append(f"Décès maternels: {len(deces)} notifications")
    except:
        pass
    return "\n".join(summary) if summary else "Aucun enregistrement."


def get_chatbot_response(user_message, chat_history):
    api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return "Clé API Groq non configurée."

    system_prompt = f"""Tu es l'assistant ESPC du CSR NAGNENEFOUN. Tu aides avec:
1. La grille d'évaluation ESPC (1600 points)
2. Le choix et la génération de documents
3. Les questions sur les normes et critères
4. La consultation des données enregistrées

{get_espc_knowledge()}

{get_templates_knowledge()}

ENREGISTREMENTS:
{get_saved_data_summary()}

RÈGLES:
1. Français, précis, concis
2. Quand on demande un document, identifie le type exact et les paramètres nécessaires
3. Cite les normes ESPC pertinentes
4. Pour la génération, fournis le document complet en markdown prêt à convertir en Word
5. Pour les données, utilise les enregistrements disponibles
6. Si l'utilisateur demande un document, génère-le directement en markdown avec en-tête officiel"""

    try:
        client = Groq(api_key=api_key)
        messages = [{"role": "system", "content": system_prompt}]
        for msg in chat_history[-10:]:
            messages.append({"role": msg["role"], "content": msg["content"]})
        messages.append({"role": "user", "content": user_message})

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            max_tokens=2000,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Erreur: {str(e)}"


def page_chatbot():
    st.title("💬 Assistant ESPC")
    st.caption(
        "Questions sur la grille ESPC, choix de documents, consultation des données"
    )
    st.divider()

    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []
    if "chat_doc_counter" not in st.session_state:
        st.session_state.chat_doc_counter = 0

    for i, msg in enumerate(st.session_state.chat_messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

            if msg.get("is_doc"):
                doc_name = msg.get("doc_name", "Document_ESPC")
                doc_content = msg.get("doc_content", msg["content"])
                doc = creer_document_word(doc_name.replace("_", " "), doc_content)
                from io import BytesIO

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    f"📥 Télécharger {doc_name.replace('_', ' ')} (.docx)",
                    buf,
                    f"{doc_name}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"chat_dl_{i}",
                )

    if prompt := st.chat_input("Posez votre question sur l'ESPC..."):
        st.session_state.chat_messages.append({"role": "user", "content": prompt})

        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Réflexion..."):
                response = get_chatbot_response(prompt, st.session_state.chat_messages)
            st.markdown(response)

            is_doc = False
            doc_name = None
            doc_content = response

            doc_keywords = {
                "pv réunion mensuelle": "PV_Réunion_Mensuelle",
                "pv réunion coges": "PV_Réunion_COGES",
                "pv assemblée": "PV_Assemblée_Générale",
                "rapport supervision": "Rapport_Supervision_ASC",
                "rapport plaintes": "Rapport_Plaintes",
                "fiche de poste": "Fiche_de_Poste",
                "fiche de nomination": "Fiche_de_Nomination",
                "calendrier nettoyage": "Calendrier_Nettoyage",
                "plan action infections": "Plan_Action_Infections",
                "rapport formation": "Rapport_Formation",
                "note de service": "Note_de_Service",
                "liste personnel": "Liste_Personnel",
                "liste coges": "Liste_COGES",
                "grille supervision": "Grille_Supervision_ASC",
                "plan supervision": "Plan_Supervision_ASC",
            }

            for key, name in doc_keywords.items():
                if key in prompt.lower():
                    is_doc = True
                    doc_name = name
                    break

            if not is_doc:
                doc_markers = [
                    "| N°",
                    "## I.",
                    "RÉPUBLIQUE DE CÔTE",
                    "MINISTÈRE",
                    "ARTICLE",
                    "Le Chef du",
                    "Le Médecin Chef",
                ]
                for marker in doc_markers:
                    if marker in response:
                        is_doc = True
                        doc_name = f"Document_ESPC_{st.session_state.chat_doc_counter}"
                        st.session_state.chat_doc_counter += 1
                        break

            if is_doc and doc_name:
                st.info(f"Document généré: **{doc_name.replace('_', ' ')}**")
                doc = creer_document_word(doc_name.replace("_", " "), doc_content)
                from io import BytesIO

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    f"📥 Télécharger {doc_name.replace('_', ' ')} (.docx)",
                    buf,
                    f"{doc_name}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="chat_dl_new",
                )

        st.session_state.chat_messages.append(
            {
                "role": "assistant",
                "content": response,
                "is_doc": is_doc,
                "doc_name": doc_name,
                "doc_content": doc_content if is_doc else None,
            }
        )


# =============================================================================
# SIDEBAR CHAT WIDGET (accessible depuis toutes les pages)
# =============================================================================


def sidebar_chat_widget():
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 💬 Assistant ESPC")

        if "sidebar_chat_history" not in st.session_state:
            st.session_state.sidebar_chat_history = []

        for msg in st.session_state.sidebar_chat_history[-2:]:
            role = "Vous" if msg["role"] == "user" else "Assistant"
            st.caption(f"**{role}:** {msg['content'][:80]}...")

        with st.form("sidebar_chat_form", clear_on_submit=True):
            user_msg = st.text_input("Question rapide...", key="sidebar_chat_input")
            submitted = st.form_submit_button("Envoyer", use_container_width=True)

            if submitted and user_msg:
                st.session_state.sidebar_chat_history.append(
                    {"role": "user", "content": user_msg}
                )
                with st.spinner("..."):
                    response = get_chatbot_response(
                        user_msg, st.session_state.sidebar_chat_history
                    )
                st.session_state.sidebar_chat_history.append(
                    {"role": "assistant", "content": response}
                )
                st.rerun()

        if st.button("💬 Ouvrir l'assistant complet", use_container_width=True):
            st.session_state["open_chatbot"] = True


if __name__ == "__main__":
    main()
